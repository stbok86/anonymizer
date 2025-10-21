#!/usr/bin/env python3
"""
Скрипт для создания DOCX инструкции по сохранению проекта Anonymizer на GitHub
"""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

def create_github_instructions():
    """Создает DOCX файл с инструкцией по GitHub"""
    
    # Создаем новый документ
    doc = Document()
    
    # Заголовок документа
    title = doc.add_heading('ИНСТРУКЦИЯ: Сохранение проекта Anonymizer на GitHub', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Добавляем информацию о проекте
    doc.add_paragraph('Дата создания: 21 октября 2025 г.')
    doc.add_paragraph('Версия: 1.0.0')
    doc.add_paragraph('Статус: BlockBuilder с поддержкой header/footer извлечения завершен')
    
    # Этап 1
    doc.add_heading('Этап 1: Подготовка репозитория', level=1)
    
    doc.add_heading('1.1 Проверка текущего состояния Git', level=2)
    p = doc.add_paragraph('Выполните в PowerShell:')
    code = doc.add_paragraph('cd C:\\Projects\\Anonymizer\ngit status')
    code.style = 'Intense Quote'
    
    doc.add_heading('1.2 Инициализация Git (если нужно)', level=2)
    p = doc.add_paragraph('Если Git не инициализирован:')
    code = doc.add_paragraph('git init')
    code.style = 'Intense Quote'
    
    doc.add_heading('1.3 Создание .gitignore файла', level=2)
    doc.add_paragraph('Создайте файл .gitignore со следующим содержимым:')
    
    gitignore_content = '''# Python
__pycache__/
*.py[cod]
*$py.class
*.so
.Python
env/
venv/
venv_*/
ENV/
env.bak/
venv.bak/

# FastAPI
.env
.env.local
.env.production

# IDE
.vscode/
.idea/
*.swp
*.swo
*~

# OS
.DS_Store
Thumbs.db

# Logs
*.log
logs/

# Test results
test_result.json
test_*.json

# Temporary files
temp/
tmp/
*.tmp

# Virtual environments
venv_gateway/
venv_orchestrator/
venv_unified_document_service/
venv_nlp_service/
venv_rule_engine/
venv_frontend/'''
    
    code = doc.add_paragraph(gitignore_content)
    code.style = 'Intense Quote'
    
    # Этап 2
    doc.add_heading('Этап 2: Подготовка коммита', level=1)
    
    doc.add_heading('2.1 Добавление файлов в индекс', level=2)
    p = doc.add_paragraph('Добавляем все изменения (кроме файлов из .gitignore):')
    code = doc.add_paragraph('git add .')
    code.style = 'Intense Quote'
    
    doc.add_heading('2.2 Создание коммита', level=2)
    p = doc.add_paragraph('Создаем коммит с описанием изменений:')
    commit_msg = '''git commit -m "feat: Complete BlockBuilder with header/footer extraction for anonymization

- Implement robust DOCX parsing with SDT support
- Add header/footer text extraction with XPath
- Include text normalization for consistent matching
- Add sensitive_matches metadata for anonymization
- Support all 6 microservices with virtual environments
- Complete startup scripts for Windows environment
- Fix critical issue with missing header data in blocks"'''
    
    code = doc.add_paragraph(commit_msg)
    code.style = 'Intense Quote'
    
    # Этап 3
    doc.add_heading('Этап 3: Создание GitHub репозитория', level=1)
    
    doc.add_heading('3.1 Создание README.md', level=2)
    doc.add_paragraph('Создайте файл README.md с описанием проекта (см. содержимое ниже)')
    
    doc.add_heading('3.2 Подключение к GitHub', level=2)
    doc.add_paragraph('Замените [YOUR_USERNAME] на ваш GitHub username:')
    code = doc.add_paragraph('git remote add origin https://github.com/[YOUR_USERNAME]/anonymizer.git')
    code.style = 'Intense Quote'
    
    doc.add_heading('3.3 Пуш в GitHub', level=2)
    p = doc.add_paragraph('Отправляем код в GitHub:')
    push_commands = '''git branch -M main
git push -u origin main'''
    code = doc.add_paragraph(push_commands)
    code.style = 'Intense Quote'
    
    # Этап 4
    doc.add_heading('Этап 4: Создание релиза', level=1)
    
    doc.add_heading('4.1 Создание тега версии', level=2)
    p = doc.add_paragraph('Создаем тег версии для релиза:')
    tag_commands = '''git tag -a v1.0.0 -m "Release v1.0.0: Complete BlockBuilder with header extraction"
git push origin v1.0.0'''
    code = doc.add_paragraph(tag_commands)
    code.style = 'Intense Quote'
    
    # Дополнительная информация
    doc.add_heading('Структура проекта', level=1)
    
    structure = '''Anonymizer/
├── frontend/                    # Streamlit frontend (порт 8501)
├── gateway/                     # API Gateway (порт 8000)
├── orchestrator/               # Orchestrator service (порт 8002)
├── unified_document_service/   # Document processing (порт 8001)
│   ├── app/
│   │   ├── main.py
│   │   ├── block_builder.py    # КРИТИЧНЫЙ: извлечение header/footer
│   │   └── document_io.py
│   ├── test_docs/              # Тестовые документы
│   └── test_block_builder.py   # Тесты
├── nlp_service/                # NLP processing (порт 8003)
├── rule_engine/                # Rule engine (порт 8004)
├── start_*.bat                 # Скрипты запуска сервисов
├── venv_*/                     # Virtual environments
└── *.md                        # Документация'''
    
    code = doc.add_paragraph(structure)
    code.style = 'Intense Quote'
    
    # Важные замечания
    doc.add_heading('Важные замечания', level=1)
    
    doc.add_paragraph('🔴 КРИТИЧНО: BlockBuilder исправлен для анонимизации')
    doc.add_paragraph('• Извлечение header/footer данных через SDT элементы')
    doc.add_paragraph('• Нормализация текста (неразрывные пробелы)')
    doc.add_paragraph('• Метаданные applies_to: "section"')
    doc.add_paragraph('• Поддержка всех типов колонтитулов')
    
    doc.add_paragraph('')
    doc.add_paragraph('✅ Готово к продакшн:')
    doc.add_paragraph('• 6 микросервисов с виртуальными окружениями')
    doc.add_paragraph('• Windows .bat скрипты для запуска')
    doc.add_paragraph('• Полное извлечение чувствительных данных')
    doc.add_paragraph('• Тесты и валидация BlockBuilder')
    
    # Сохраняем документ
    doc.save('C:\\Projects\\Anonymizer\\GitHub_Setup_Instructions.docx')
    print("✅ Инструкция сохранена в GitHub_Setup_Instructions.docx")

if __name__ == "__main__":
    create_github_instructions()