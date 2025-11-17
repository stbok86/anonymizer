#!/usr/bin/env python3
"""
Отладка проблемы селективной анонимизации: 53 → 40 → 26
Проверяем каждый этап обработки данных
"""

import requests
import json
import tempfile
from docx import Document

def create_test_document():
    """Создает тестовый документ"""
    doc = Document()
    
    # Добавляем параграф с известными данными
    paragraph = doc.add_paragraph()
    paragraph.add_run("МИНИСТЕРСТВО ИНФОРМАЦИОННОГО РАЗВИТИЯ И СВЯЗИ ПЕРМСКОГО КРАЯ\n")
    paragraph.add_run("Email: test@example.com\n")
    paragraph.add_run("Телефон: +7 923 123-45-67\n")
    paragraph.add_run("Организация: ООО «Технические решения»\n")
    
    # Сохраняем во временный файл
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        return tmp.name

def debug_selective_anonymization():
    """Отлаживаем селективную анонимизацию"""
    print("🐛 Отладка селективной анонимизации...")
    
    # Создаем тестовый документ
    test_file = create_test_document()
    print(f"📄 Создан тестовый документ: {test_file}")
    
    # Шаг 1: Анализируем документ для получения всех найденных элементов
    print("\n📊 Шаг 1: Анализ документа...")
    
    try:
        with open(test_file, 'rb') as f:
            files = {'file': (f'test_document.docx', f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
            data = {'patterns_file': 'patterns/sensitive_patterns.xlsx'}
            
            # Анализируем документ
            response = requests.post(
                "http://localhost:8002/analyze_document",
                files=files,
                data=data,
                timeout=30
            )
        
        if response.status_code == 200:
            result = response.json()
            found_items = result.get('items', [])
            print(f"✅ Анализ завершен: найдено {len(found_items)} элементов")
            
            # Показываем первые 3 найденных элемента
            for i, item in enumerate(found_items[:3]):
                print(f"   {i+1}. '{item.get('original_value', 'N/A')}' ({item.get('category', 'N/A')})")
            
            if len(found_items) > 3:
                print(f"   ... и еще {len(found_items) - 3} элементов")
                
            # Шаг 2: Попробуем селективную анонимизацию ВСЕХ найденных элементов
            print(f"\n🔒 Шаг 2: Селективная анонимизация {len(found_items)} элементов...")
            
            # Подготавливаем данные для селективной анонимизации
            selected_items = []
            for item in found_items:
                selected_item = {
                    'block_id': item.get('block_id', ''),
                    'original_value': item.get('original_value', ''),
                    'uuid': item.get('uuid', ''),
                    'position': item.get('position', {}),
                    'category': item.get('category', ''),
                    'confidence': item.get('confidence', 1.0)
                }
                selected_items.append(selected_item)
            
            print(f"📝 Подготовлено для анонимизации: {len(selected_items)} элементов")
            
            # Отправляем на селективную анонимизацию
            with open(test_file, 'rb') as f:
                files = {'file': (f'test_document.docx', f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
                data = {
                    'patterns_file': 'patterns/sensitive_patterns.xlsx',
                    'selected_items': json.dumps(selected_items)
                }
                
                response = requests.post(
                    "http://localhost:8002/anonymize_selected",
                    files=files,
                    data=data,
                    timeout=30
                )
            
            if response.status_code == 200:
                result = response.json()
                print(f"✅ Селективная анонимизация завершена")
                print(f"📊 Результат:")
                print(f"   • Статус: {result.get('status', 'N/A')}")
                print(f"   • Обработано элементов: {result.get('selected_items_count', 'N/A')}")
                print(f"   • Выполнено замен: {result.get('replacements_applied', 'N/A')}")
                
                # Анализируем потери
                original_count = len(found_items)
                processed_count = result.get('selected_items_count', 0)
                applied_count = result.get('replacements_applied', 0)
                
                print(f"\n📈 Анализ потерь:")
                print(f"   🔍 Найдено: {original_count}")
                print(f"   📝 Отправлено на обработку: {len(selected_items)}")
                print(f"   ⚙️  Обработано системой: {processed_count}")
                print(f"   ✅ Фактически заменено: {applied_count}")
                
                if original_count != applied_count:
                    loss_1 = original_count - processed_count
                    loss_2 = processed_count - applied_count
                    print(f"\n⚠️  ПОТЕРИ ДАННЫХ:")
                    if loss_1 > 0:
                        print(f"   📤 Потеря на этапе отправки: -{loss_1} элементов")
                    if loss_2 > 0:
                        print(f"   🔧 Потеря на этапе применения: -{loss_2} элементов")
                else:
                    print(f"\n✅ Потерь данных нет - система работает корректно!")
                    
            else:
                print(f"❌ Ошибка селективной анонимизации: {response.status_code}")
                print(f"📋 Ответ: {response.text}")
                
        else:
            print(f"❌ Ошибка анализа: {response.status_code}")
            print(f"📋 Ответ: {response.text}")
            
    except Exception as e:
        print(f"❌ Ошибка: {str(e)}")
    finally:
        # Удаляем временный файл
        import os
        if os.path.exists(test_file):
            os.remove(test_file)

if __name__ == "__main__":
    debug_selective_anonymization()