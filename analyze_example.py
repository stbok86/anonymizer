#!/usr/bin/env python3
"""
Скрипт для демонстрации процесса анонимизации на конкретном примере
Документ: test_01_1_4_S.docx
Цель: "Общество с ограниченной ответственностью «КАМА Технологии»" из paragraph_82
"""

import os
import sys
import json
import uuid
from pathlib import Path

# Добавляем пути к модулям
unified_service_path = os.path.join(os.path.dirname(__file__), 'unified_document_service', 'app')
sys.path.append(unified_service_path)

try:
    from docx import Document
    from block_builder import BlockBuilder
    from rule_adapter import RuleEngineAdapter
    from formatter_applier import FormatterApplier
    from full_anonymizer import FullAnonymizer
except ImportError as e:
    print(f"❌ Ошибка импорта модулей: {e}")
    print("Убедитесь, что все зависимости установлены")
    sys.exit(1)

def main():
    # Путь к тестовому документу
    doc_path = "unified_document_service/test_docs/test_01_1_4_S.docx"
    
    if not os.path.exists(doc_path):
        print(f"❌ Файл не найден: {doc_path}")
        return
    
    print("🔍 ДЕМОНСТРАЦИЯ ПРОЦЕССА АНОНИМИЗАЦИИ")
    print("=" * 60)
    print(f"📄 Анализируемый файл: {doc_path}")
    print(f"🎯 Искомый текст: 'Общество с ограниченной ответственностью «КАМА Технологии»'")
    print(f"📍 Ожидаемое расположение: paragraph_82")
    print()
    
    # ЭТАП 1: Загрузка и анализ структуры документа
    print("📊 ЭТАП 1: АНАЛИЗ СТРУКТУРЫ ДОКУМЕНТА")
    print("-" * 40)
    
    doc = Document(doc_path)
    block_builder = BlockBuilder()
    
    print(f"📋 Загружен документ с {len(doc.paragraphs)} параграфами и {len(doc.tables)} таблицами")
    
    # Построение блоков
    blocks = block_builder.build_blocks(doc)
    print(f"🧩 Извлечено блоков: {len(blocks)}")
    
    # Поиск параграфа с целевым текстом
    target_text = "Общество с ограниченной ответственностью «КАМА Технологии»"
    target_block = None
    target_paragraph = None
    
    print(f"\n🔍 Поиск целевого текста: '{target_text}'")
    
    for block in blocks:
        block_content = block.get('content', '')
        if target_text in block_content:
            target_block = block
            target_paragraph = block.get('element')
            print(f"✅ Найден в блоке: {block['block_id']}")
            print(f"📝 Содержимое блока: '{block_content[:100]}{'...' if len(block_content) > 100 else ''}'")
            break
    
    if not target_block:
        print("❌ Целевой текст не найден в документе")
        
        # Покажем содержимое некоторых блоков для отладки
        print("\n🔍 Доступные блоки (первые 10):")
        for i, block in enumerate(blocks[:10]):
            content = block.get('content', '')[:80]
            print(f"  {block['block_id']}: '{content}{'...' if len(content) >= 80 else ''}'")
        
        return
    
    print()
    
    # ЭТАП 2: Моделирование обнаружения чувствительных данных
    print("🕵️ ЭТАП 2: ОБНАРУЖЕНИЕ ЧУВСТВИТЕЛЬНЫХ ДАННЫХ")
    print("-" * 50)
    
    # Симуляция результата анализа (как если бы это пришло из rule_engine или nlp_service)
    detected_item = {
        'block_id': target_block['block_id'],
        'original_value': target_text,
        'uuid': str(uuid.uuid4()),
        'position': {
            'start': target_block['content'].find(target_text),
            'end': target_block['content'].find(target_text) + len(target_text)
        },
        'category': 'organization',
        'confidence': 0.95,
        'method': 'regex_pattern',
        'approved': True  # Пользователь одобрил для анонимизации
    }
    
    print(f"🎯 Обнаружен элемент:")
    print(f"   📍 Блок ID: {detected_item['block_id']}")
    print(f"   📝 Значение: '{detected_item['original_value']}'")
    print(f"   🔑 UUID: {detected_item['uuid']}")
    print(f"   📊 Позиция: {detected_item['position']['start']}-{detected_item['position']['end']}")
    print(f"   🏷️  Категория: {detected_item['category']}")
    print(f"   ✅ Уверенность: {detected_item['confidence']}")
    print()
    
    # ЭТАП 3: Подготовка к анонимизации
    print("🔧 ЭТАП 3: ПОДГОТОВКА К АНОНИМИЗАЦИИ")
    print("-" * 40)
    
    # Создаем карту блоков
    blocks_map = {block['block_id']: block for block in blocks}
    print(f"🗺️  Создана карта блоков: {len(blocks_map)} элементов")
    
    # Подготавливаем замену
    replacement = {
        'block_id': detected_item['block_id'],
        'original_value': detected_item['original_value'],
        'uuid': detected_item['uuid'],
        'position': detected_item['position'],
        'element': target_block.get('element'),
        'category': detected_item['category']
    }
    
    print(f"🔄 Подготовлена замена:")
    print(f"   🎯 Заменяем: '{replacement['original_value']}'")
    print(f"   ➡️  На UUID: '{replacement['uuid']}'")
    print(f"   📍 В элементе: {type(replacement['element'])}")
    print()
    
    # ЭТАП 4: Применение замены
    print("✨ ЭТАП 4: ПРИМЕНЕНИЕ ЗАМЕНЫ")
    print("-" * 30)
    
    formatter = FormatterApplier(highlight_replacements=True)
    
    # Показываем содержимое до замены
    original_content = target_paragraph.text if target_paragraph else "N/A"
    print(f"📝 Содержимое ДО замены:")
    print(f"   '{original_content}'")
    print()
    
    # Применяем замену
    result = formatter._apply_single_replacement(replacement)
    
    # Показываем содержимое после замены
    if target_paragraph:
        modified_content = target_paragraph.text
        print(f"✅ Содержимое ПОСЛЕ замены:")
        print(f"   '{modified_content}'")
        print()
        
        print(f"🔄 Замена {'УСПЕШНА' if result else 'НЕ УДАЛАСЬ'}")
        
        if result:
            print(f"   🔹 Исходный текст: '{target_text}'")
            print(f"   🔹 Заменен на: '{detected_item['uuid']}'")
            print(f"   🔹 Длина исходного: {len(target_text)} символов")
            print(f"   🔹 Длина замены: {len(detected_item['uuid'])} символов")
    
    # ЭТАП 5: Статистика
    print()
    print("📊 ЭТАП 5: СТАТИСТИКА АНОНИМИЗАЦИИ")
    print("-" * 35)
    
    stats = {
        'total_replacements': 1 if result else 0,
        'categories': {'organization': 1} if result else {},
        'blocks_processed': 1,
        'replacement_details': [{
            'uuid': detected_item['uuid'],
            'category': detected_item['category'],
            'original_value': detected_item['original_value'],
            'success': result
        }]
    }
    
    print(f"🎯 Всего замен: {stats['total_replacements']}")
    print(f"📦 Блоков обработано: {stats['blocks_processed']}")
    print(f"🏷️  По категориям: {stats['categories']}")
    print()
    
    # Сохранение результата (опционально)
    output_path = "test_anonymized_demo.docx"
    if result:
        doc.save(output_path)
        print(f"💾 Результат сохранен в: {output_path}")
    
    print()
    print("🎉 ДЕМОНСТРАЦИЯ ЗАВЕРШЕНА!")
    print("=" * 60)

if __name__ == "__main__":
    main()