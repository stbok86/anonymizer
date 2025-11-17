#!/usr/bin/env python3
"""
Скрипт для создания DOCX документа с требованиями к виртуальной машине
для анонимайзера документов
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

def add_heading_with_emoji(doc, text, level=1):
    """Добавляет заголовок с эмодзи"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_code_block(doc, code_text, language="bash"):
    """Добавляет блок кода с форматированием"""
    p = doc.add_paragraph()
    p.style = 'Normal'
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    # Добавляем отступ
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return p

def create_vm_requirements_document():
    """Создает документ с требованиями к ВМ"""
    
    # Создаем новый документ
    doc = Document()
    
    # Заголовок документа
    title = doc.add_heading('ТЕХНИЧЕСКИЕ ТРЕБОВАНИЯ К ВИРТУАЛЬНОЙ МАШИНЕ', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Система анонимизации документов')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(14)
    subtitle_run.bold = True
    
    # Дата документа
    date_p = doc.add_paragraph(f'Дата: 13 ноября 2025')
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_page_break()
    
    # 1. Операционная система
    add_heading_with_emoji(doc, '1. ОПЕРАЦИОННАЯ СИСТЕМА', level=1)
    doc.add_paragraph('Ubuntu Server 20.04 LTS или Ubuntu Server 22.04 LTS', style='Intense Quote')
    
    doc.add_paragraph('Обоснование выбора:')
    benefits = [
        'Стабильная поддержка Python 3.8+',
        'Легкая установка spaCy моделей',
        'Поддержка systemd для управления сервисами',
        'Оптимальная совместимость с FastAPI/uvicorn',
        'Долгосрочная поддержка (LTS)'
    ]
    for benefit in benefits:
        doc.add_paragraph(f'• {benefit}', style='List Bullet')
    
    # 2. Ресурсы процессора
    add_heading_with_emoji(doc, '2. РЕСУРСЫ ПРОЦЕССОРА', level=1)
    doc.add_paragraph('Минимум: 8 vCPU | Рекомендовано: 12 vCPU', style='Intense Quote')
    
    doc.add_paragraph('Требования по компонентам:')
    cpu_requirements = [
        'NLP Service: требует значительных вычислительных ресурсов для spaCy модели ru_core_news_lg (1-2 CPU на процесс)',
        '5 микросервисов одновременно (Gateway, Orchestrator, Rule Engine, Unified Document, NLP)',
        'Параллельная обработка: до 30 пользователей одновременно',
        'Обработка документов до 100 страниц за ≤2 минуты'
    ]
    for req in cpu_requirements:
        doc.add_paragraph(f'• {req}', style='List Bullet')
    
    # 3. Оперативная память
    add_heading_with_emoji(doc, '3. ОПЕРАТИВНАЯ ПАМЯТЬ', level=1)
    doc.add_paragraph('Минимум: 16 GB | Рекомендовано: 24 GB', style='Intense Quote')
    
    doc.add_paragraph('Распределение по сервисам:')
    memory_distribution = [
        'NLP Service: 4-6 GB (spaCy модель ru_core_news_lg ~500MB + pymorphy3 + обработка)',
        'Unified Document Service: 2-3 GB (python-docx, обработка больших документов)',
        'Frontend (Streamlit): 1-2 GB',
        'Rule Engine: 1 GB',
        'Gateway + Orchestrator: 1 GB суммарно',
        'ОС + буферы: 2-3 GB',
        'Резерв для пиковых нагрузок: 4-6 GB'
    ]
    for mem in memory_distribution:
        doc.add_paragraph(f'• {mem}', style='List Bullet')
    
    # 4. Дисковое пространство
    add_heading_with_emoji(doc, '4. ДИСКОВОЕ ПРОСТРАНСТВО', level=1)
    doc.add_paragraph('Минимум: 50 GB | Рекомендовано: 100 GB SSD', style='Intense Quote')
    
    doc.add_paragraph('Распределение дискового пространства:')
    disk_distribution = [
        'Система и зависимости: 15 GB',
        'Python окружения: 5 GB (spaCy модели ~1.5GB)',
        'Логи сервисов: 5 GB',
        'Временные файлы документов: 10 GB',
        'Резервные копии конфигураций: 5 GB',
        'Резерв для роста: 20-60 GB'
    ]
    for disk in disk_distribution:
        doc.add_paragraph(f'• {disk}', style='List Bullet')
    
    # 5. Сетевые требования
    add_heading_with_emoji(doc, '5. СЕТЕВЫЕ ТРЕБОВАНИЯ', level=1)
    
    doc.add_paragraph('Входящие порты:')
    external_ports = ['80, 443 (веб-интерфейс)', '8501 (Streamlit frontend)']
    for port in external_ports:
        doc.add_paragraph(f'• {port}', style='List Bullet')
    
    doc.add_paragraph('Внутренние порты микросервисов:')
    internal_ports = [
        '8002 (Gateway)',
        '8003 (Rule Engine)', 
        '8004 (Orchestrator)',
        '8006 (NLP Service)',
        '8009 (Unified Document Service)'
    ]
    for port in internal_ports:
        doc.add_paragraph(f'• {port}', style='List Bullet')
        
    doc.add_paragraph('Пропускная способность: минимум 100 Mbps для загрузки/выгрузки документов')
    
    # 6. Специфические зависимости
    add_heading_with_emoji(doc, '6. СПЕЦИФИЧЕСКИЕ ЗАВИСИМОСТИ', level=1)
    
    doc.add_paragraph('Python 3.8+ с основными пакетами:')
    add_code_block(doc, '''# Основные зависимости
fastapi uvicorn streamlit python-docx
spacy pymorphy3 pandas openpyxl
regex lxml requests pydantic
python-multipart''')
    
    doc.add_paragraph('Установка spaCy модели для русского языка:')
    add_code_block(doc, '''# Рекомендуемая большая модель
python -m spacy download ru_core_news_lg

# Альтернативы при недоступности
python -m spacy download ru_core_news_md
python -m spacy download ru_core_news_sm''')
    
    # 7. Производительность для 20-30 пользователей
    add_heading_with_emoji(doc, '7. ПРОИЗВОДИТЕЛЬНОСТЬ ДЛЯ 20-30 ПОЛЬЗОВАТЕЛЕЙ', level=1)
    
    doc.add_paragraph('Расчетная нагрузка:')
    performance_metrics = [
        'Одновременных анализов документов: 5-8',
        'Пиковая нагрузка: обработка 100-страничного документа за 2 минуты',
        'NLP обработка: ~13 детекций на документ с уверенностью 0.7-1.0',
        'Средний размер документа: 10-20 страниц = обработка за 10-30 секунд'
    ]
    for metric in performance_metrics:
        doc.add_paragraph(f'• {metric}', style='List Bullet')
    
    doc.add_paragraph('Архитектурные требования:')
    arch_requirements = [
        'Микросервисная архитектура: 6 независимых сервисов',
        'Load balancing: встроенный в FastAPI/uvicorn',
        'Масштабирование: горизонтальное через дополнительные экземпляры сервисов'
    ]
    for req in arch_requirements:
        doc.add_paragraph(f'• {req}', style='List Bullet')
    
    # 8. Мониторинг и логирование
    add_heading_with_emoji(doc, '8. МОНИТОРИНГ И ЛОГИРОВАНИЕ', level=1)
    monitoring_features = [
        'Логи всех сервисов с ротацией',
        'Мониторинг ресурсов (CPU, RAM, disk)',
        'Health checks для всех микросервисов',
        'Alerting при недоступности сервисов'
    ]
    for feature in monitoring_features:
        doc.add_paragraph(f'• {feature}', style='List Bullet')
    
    # 9. Безопасность
    add_heading_with_emoji(doc, '9. БЕЗОПАСНОСТЬ', level=1)
    security_features = [
        'Firewall настройка для портов приложений',
        'Regular updates ОС и зависимостей',
        'Backup конфигураций и паттернов',
        'Изоляция сервисов через systemd'
    ]
    for feature in security_features:
        doc.add_paragraph(f'• {feature}', style='List Bullet')
    
    # 10. Критерии приемки MAX
    add_heading_with_emoji(doc, '10. СООТВЕТСТВИЕ КРИТЕРИЯМ ПРИЕМКИ (MAX)', level=1)
    
    doc.add_paragraph('Данная конфигурация ВМ обеспечивает выполнение всех максимальных критериев приемки:')
    
    acceptance_criteria = [
        '✅ Система поддерживает все заявленные категории данных (ФИО, адреса, паспорта, системы, контракты и пр.)',
        '✅ Система находит все виды чувствительных данных в документе, включая заголовки, таблицы, сноски, подписи',
        '✅ Замены выполняются с учётом падежей и регистра (например, «Единая/Единой/ЕДИНАЯ система» → [SYS-001])',
        '✅ Обрабатываются все элементы документа: текст, заголовки, таблицы, сноски, подписи к рисункам',
        '✅ Номера документов вида «№ …» заменяются автоматически регулярными правилами',
        '✅ Таблица замен формируется без дубликатов, с уникальными ID',
        '✅ Поддержка документов до 100 страниц за ≤ 2 минуты',
        '✅ В логах указывается статистика: количество замен, ошибки, предупреждения',
        '✅ Возможность расширять категории данных без изменения кода (конфигурационный файл)',
        '✅ Количество одновременно работающих пользователей - 20-30 сотрудников'
    ]
    
    for criterion in acceptance_criteria:
        doc.add_paragraph(criterion, style='List Bullet')
    
    # Рекомендуемые характеристики
    doc.add_page_break()
    add_heading_with_emoji(doc, 'РЕКОМЕНДУЕМАЯ КОНФИГУРАЦИЯ ДЛЯ PRODUCTION', level=1)
    
    # Создаем таблицу с характеристиками
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Компонент'
    hdr_cells[1].text = 'Спецификация'
    
    # Данные для таблицы
    specs = [
        ('CPU', '12 vCPU (Intel/AMD x64)'),
        ('RAM', '24 GB'),
        ('Disk', '100 GB SSD'),
        ('OS', 'Ubuntu Server 22.04 LTS'),
        ('Network', '1 Gbps'),
        ('Архитектура', 'Микросервисная (6 сервисов)'),
        ('Python', '3.8+ с зависимостями'),
        ('NLP модель', 'spaCy ru_core_news_lg')
    ]
    
    for component, specification in specs:
        row_cells = table.add_row().cells
        row_cells[0].text = component
        row_cells[1].text = specification
    
    # Заключение
    doc.add_paragraph()
    conclusion = doc.add_paragraph('Данная конфигурация обеспечивает стабильную работу системы анонимизации документов с возможностью обслуживания 20-30 пользователей одновременно и обработкой документов до 100 страниц за время не более 2 минут.')
    conclusion.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Сохраняем документ
    output_path = 'C:\\Projects\\Anonymizer\\VM_Requirements_Anonymizer_System.docx'
    doc.save(output_path)
    
    return output_path

if __name__ == "__main__":
    try:
        output_file = create_vm_requirements_document()
        print(f"✅ Документ успешно создан: {output_file}")
    except Exception as e:
        print(f"❌ Ошибка создания документа: {str(e)}")