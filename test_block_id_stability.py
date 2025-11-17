#!/usr/bin/env python3
"""
Проверка стабильности block_id
Проверяем, что block_id генерируются одинаково при повторной загрузке документа
"""

import tempfile
from docx import Document
import sys
import os

# Добавляем путь к модулям
sys.path.append(os.path.join(os.path.dirname(__file__), 'unified_document_service', 'app'))

from block_builder import BlockBuilder

def create_test_document():
    """Создает простой тестовый документ"""
    doc = Document()
    
    # Добавляем параграфы
    doc.add_paragraph("Первый параграф с текстом")
    doc.add_paragraph("Второй параграф с данными")
    doc.add_paragraph("МИНИСТЕРСТВО ИНФОРМАЦИОННОГО РАЗВИТИЯ И СВЯЗИ ПЕРМСКОГО КРАЯ")
    
    # Добавляем таблицу
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Email: test@example.com"
    table.cell(0, 1).text = "Телефон: +7 123 456-78-90"
    table.cell(1, 0).text = "Организация: ООО «Тест»"
    table.cell(1, 1).text = "ИНН: 1234567890"
    
    # Сохраняем во временный файл
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        return tmp.name

def test_block_id_stability():
    """Тестирует стабильность генерации block_id"""
    print("🧪 Тест стабильности block_id...")
    
    # Создаем тестовый документ
    test_file = create_test_document()
    print(f"📄 Создан тестовый документ: {test_file}")
    
    try:
        # Первая загрузка
        print("\n🔄 Первая загрузка документа...")
        doc1 = Document(test_file)
        builder1 = BlockBuilder()
        blocks1 = builder1.build_blocks(doc1)
        
        print(f"📊 Найдено блоков: {len(blocks1)}")
        for block in blocks1:
            print(f"   • {block['block_id']}: '{block['text'][:50]}...' ({block['type']})")
        
        # Вторая загрузка (симуляция селективной анонимизации)
        print("\n🔄 Вторая загрузка документа...")
        doc2 = Document(test_file)
        builder2 = BlockBuilder()
        blocks2 = builder2.build_blocks(doc2)
        
        print(f"📊 Найдено блоков: {len(blocks2)}")
        for block in blocks2:
            print(f"   • {block['block_id']}: '{block['text'][:50]}...' ({block['type']})")
        
        # Сравнение block_id
        print(f"\n🔍 Сравнение block_id...")
        
        block_ids_1 = {block['block_id'] for block in blocks1}
        block_ids_2 = {block['block_id'] for block in blocks2}
        
        print(f"📋 Первая загрузка: {sorted(block_ids_1)}")
        print(f"📋 Вторая загрузка: {sorted(block_ids_2)}")
        
        if block_ids_1 == block_ids_2:
            print(f"✅ block_id стабильны - генерируются одинаково!")
        else:
            print(f"❌ block_id НЕ стабильны!")
            missing_in_2 = block_ids_1 - block_ids_2
            missing_in_1 = block_ids_2 - block_ids_1
            
            if missing_in_2:
                print(f"   🚫 Отсутствуют во второй загрузке: {missing_in_2}")
            if missing_in_1:
                print(f"   ➕ Новые во второй загрузке: {missing_in_1}")
                
        # Дополнительная проверка содержимого
        print(f"\n📝 Проверка содержимого блоков...")
        content_map_1 = {block['block_id']: block['text'] for block in blocks1}
        content_map_2 = {block['block_id']: block['text'] for block in blocks2}
        
        content_matches = 0
        for block_id in block_ids_1:
            if block_id in content_map_2:
                if content_map_1[block_id] == content_map_2[block_id]:
                    content_matches += 1
                else:
                    print(f"   ⚠️ Содержимое изменилось для {block_id}:")
                    print(f"      Первая: '{content_map_1[block_id]}'")
                    print(f"      Вторая: '{content_map_2[block_id]}'")
        
        print(f"📈 Совпадений содержимого: {content_matches}/{len(block_ids_1)}")
        
        if len(block_ids_1) == len(block_ids_2) and block_ids_1 == block_ids_2 and content_matches == len(block_ids_1):
            print(f"\n🎉 ВЫВОД: BlockBuilder работает детерминированно!")
        else:
            print(f"\n⚠️ ВЫВОД: BlockBuilder НЕ детерминированный - это причина потерь данных!")
            print(f"   📊 Ожидаемый результат: одинаковые block_id при повторной загрузке")
            print(f"   📊 Фактический результат: различающиеся block_id")
            print(f"   💡 РЕШЕНИЕ: Нужно стабилизировать логику генерации block_id")
        
    except Exception as e:
        print(f"❌ Ошибка: {str(e)}")
        import traceback
        traceback.print_exc()
    finally:
        # Удаляем временный файл
        if os.path.exists(test_file):
            os.remove(test_file)

if __name__ == "__main__":
    test_block_id_stability()