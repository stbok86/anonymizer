#!/usr/bin/env python3
"""
АНАЛИЗ ПРОБЛЕМЫ ДУБЛИРОВАНИЯ UUID
===============================

Анализируем документ test_01_1_4_S_anonymized (3).docx 
для выявления причин повторного использования UUID
"""

import os
from docx import Document
import re

def analyze_uuid_duplication():
    """Анализирует дублирование UUID в документе"""
    
    doc_path = "unified_document_service/test_docs/test_01_1_4_S_anonymized (3).docx"
    
    if not os.path.exists(doc_path):
        print(f"❌ Файл не найден: {doc_path}")
        return
    
    print("🔍 АНАЛИЗ ДУБЛИРОВАНИЯ UUID")
    print("=" * 60)
    
    # Загружаем документ
    doc = Document(doc_path)
    
    # UUID паттерн
    uuid_pattern = re.compile(r'[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}')
    
    # Собираем все UUID из документа
    uuid_occurrences = {}
    
    # Анализируем параграфы
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text
        if text.strip():
            uuids = uuid_pattern.findall(text.lower())
            for uuid in uuids:
                if uuid not in uuid_occurrences:
                    uuid_occurrences[uuid] = []
                uuid_occurrences[uuid].append(f"paragraph_{i}: '{text.strip()}'")
    
    # Анализируем таблицы
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                text = cell.text
                if text.strip():
                    uuids = uuid_pattern.findall(text.lower())
                    for uuid in uuids:
                        if uuid not in uuid_occurrences:
                            uuid_occurrences[uuid] = []
                        uuid_occurrences[uuid].append(f"table_{table_idx}_cell_{row_idx}_{cell_idx}: '{text.strip()}'")
    
    print(f"📊 Найдено уникальных UUID: {len(uuid_occurrences)}")
    
    # Ищем дублированные UUID
    duplicated_uuids = {uuid: locations for uuid, locations in uuid_occurrences.items() if len(locations) > 1}
    
    if duplicated_uuids:
        print(f"\n🚨 НАЙДЕНЫ ДУБЛИРОВАННЫЕ UUID: {len(duplicated_uuids)}")
        print("-" * 50)
        
        for uuid, locations in duplicated_uuids.items():
            print(f"\n🔄 UUID: {uuid}")
            print(f"   Использован {len(locations)} раз(а):")
            for location in locations:
                print(f"   • {location}")
    else:
        print("\n✅ Дублированных UUID не найдено")
    
    # Ищем конкретно проблемный UUID
    target_uuid = "545094b7-602f-4e1d-9e95-95142918f380"
    if target_uuid in uuid_occurrences:
        print(f"\n🎯 АНАЛИЗ ПРОБЛЕМНОГО UUID: {target_uuid}")
        print("-" * 50)
        locations = uuid_occurrences[target_uuid]
        print(f"Найдено вхождений: {len(locations)}")
        for location in locations:
            print(f"• {location}")
    
    # Анализируем table_2 детально
    print(f"\n📋 ДЕТАЛЬНЫЙ АНАЛИЗ TABLE_2:")
    print("-" * 30)
    
    if len(doc.tables) > 2:
        table_2 = doc.tables[2]
        print(f"Размер таблицы: {len(table_2.rows)} строк x {len(table_2.columns)} столбцов")
        
        for row_idx, row in enumerate(table_2.rows):
            for cell_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                if text and (target_uuid in text.lower() or "14 августа 2023" in text):
                    print(f"Строка {row_idx}, Ячейка {cell_idx}: '{text}'")
    else:
        print("❌ table_2 не найдена или документ содержит менее 3 таблиц")

if __name__ == "__main__":
    analyze_uuid_duplication()