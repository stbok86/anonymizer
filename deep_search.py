#!/usr/bin/env python3
"""
УГЛУБЛЕННАЯ ДИАГНОСТИКА
=====================

Ищем где именно находится текст "Общество с ограниченной ответственностью «КАМА Технологии»"
в реальном документе
"""

import os
import sys

unified_service_path = os.path.join(os.path.dirname(__file__), 'unified_document_service', 'app')
sys.path.append(unified_service_path)

def deep_search():
    print("🔍 УГЛУБЛЕННАЯ ДИАГНОСТИКА ПОИСКА ТЕКСТА")
    print("=" * 80)
    
    try:
        from docx import Document
        from block_builder import BlockBuilder
    except ImportError as e:
        print(f"❌ Ошибка импорта: {e}")
        return
    
    doc_path = "unified_document_service/test_docs/test_01_1_4_S.docx"
    
    if not os.path.exists(doc_path):
        print(f"❌ Файл не найден: {doc_path}")
        return
    
    target_text = "Общество с ограниченной ответственностью «КАМА Технологии»"
    target_short = "КАМА Технологии"
    
    print(f"📄 Анализ документа: {doc_path}")
    print(f"🎯 Ищем: '{target_text}'")
    print(f"🎯 Короткий поиск: '{target_short}'")
    print()
    
    doc = Document(doc_path)
    
    # ПОИСК 1: Прямой поиск по всем параграфам
    print("🔍 ПОИСК 1: ПРЯМОЙ ПОИСК ПО ПАРАГРАФАМ")
    print("-" * 40)
    
    found_paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        if target_text in text:
            found_paragraphs.append((i, para, 'full'))
            print(f"✅ НАЙДЕН полный текст в параграфе {i}: '{text[:100]}...'")
        elif target_short in text:
            found_paragraphs.append((i, para, 'partial'))
            print(f"⚠️  НАЙДЕН частично в параграфе {i}: '{text[:100]}...'")
    
    if not found_paragraphs:
        print("❌ НЕ найдено в параграфах!")
    
    # ПОИСК 2: Поиск в таблицах
    print("\\n🔍 ПОИСК 2: ПОИСК В ТАБЛИЦАХ")
    print("-" * 30)
    
    found_tables = []
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text
                if target_text in cell_text:
                    found_tables.append((table_idx, row_idx, cell_idx, cell, 'full'))
                    print(f"✅ НАЙДЕН полный текст в таблице {table_idx}, строка {row_idx}, ячейка {cell_idx}")
                elif target_short in cell_text:
                    found_tables.append((table_idx, row_idx, cell_idx, cell, 'partial'))
                    print(f"⚠️  НАЙДЕН частично в таблице {table_idx}, строка {row_idx}, ячейка {cell_idx}")
                    print(f"    Текст ячейки: '{cell_text[:150]}...'")
    
    if not found_tables:
        print("❌ НЕ найдено в таблицах!")
    
    # ПОИСК 3: Поиск через BlockBuilder
    print("\\n🔍 ПОИСК 3: АНАЛИЗ ЧЕРЕЗ BLOCKBUILDER")
    print("-" * 35)
    
    block_builder = BlockBuilder()
    blocks = block_builder.build_blocks(doc)
    
    found_blocks = []
    for block in blocks:
        content = block.get('content', '')
        if target_text in content:
            found_blocks.append((block, 'full'))
            print(f"✅ НАЙДЕН полный текст в блоке {block['block_id']}")
        elif target_short in content:
            found_blocks.append((block, 'partial'))
            print(f"⚠️  НАЙДЕН частично в блоке {block['block_id']}: '{content[:100]}...'")
    
    if not found_blocks:
        print("❌ НЕ найдено через BlockBuilder!")
        print("\\n📋 Показываем содержимое всех блоков:")
        for i, block in enumerate(blocks):
            content = block.get('content', '')
            if content.strip():  # Только непустые блоки
                print(f"  {block['block_id']}: '{content[:80]}{'...' if len(content) > 80 else ''}'")
            if i > 20:  # Ограничиваем вывод
                print(f"  ... и еще {len(blocks) - 20} блоков")
                break
    
    # ПОИСК 4: Поиск вариаций текста
    print("\\n🔍 ПОИСК 4: ПОИСК ВАРИАЦИЙ ТЕКСТА")
    print("-" * 35)
    
    variations = [
        "Общество с ограниченной ответственностью",
        "КАМА Технологии", 
        "«КАМА Технологии»",
        "КАМА",
        "Технологии",
        "ограниченной ответственностью"
    ]
    
    all_text = "\\n".join([para.text for para in doc.paragraphs])
    all_text += "\\n" + "\\n".join([
        cell.text 
        for table in doc.tables 
        for row in table.rows 
        for cell in row.cells
    ])
    
    print("🔍 Ищем вариации в полном тексте документа:")
    for variation in variations:
        if variation in all_text:
            count = all_text.count(variation)
            print(f"✅ '{variation}' найдено {count} раз")
        else:
            print(f"❌ '{variation}' НЕ найдено")
    
    # ПОИСК 5: Побайтовый анализ
    print("\\n🔬 ПОИСК 5: АНАЛИЗ КОДИРОВКИ")
    print("-" * 30)
    
    print(f"🎯 Искомый текст (bytes): {target_text.encode('utf-8')}")
    print(f"🎯 Длина в символах: {len(target_text)}")
    print(f"🎯 Длина в байтах: {len(target_text.encode('utf-8'))}")
    
    # Анализ специальных символов
    special_chars = []
    for i, char in enumerate(target_text):
        if ord(char) > 127:
            special_chars.append((i, char, ord(char)))
    
    print(f"🔤 Специальные символы в искомом тексте:")
    for pos, char, code in special_chars:
        print(f"   Позиция {pos}: '{char}' (код {code})")
    
    # ИТОГОВЫЙ ОТЧЕТ
    print("\\n📊 ИТОГОВЫЙ ОТЧЕТ")
    print("-" * 17)
    
    total_found = len(found_paragraphs) + len(found_tables) + len(found_blocks)
    
    if total_found == 0:
        print("🚨 КРИТИЧЕСКАЯ ПРОБЛЕМА:")
        print("   Текст НЕ найден НИГДЕ в документе!")
        print("   Возможные причины:")
        print("   1. Текст отсутствует в данном документе")
        print("   2. Текст имеет другую кодировку")
        print("   3. Текст разбит между элементами")
        print("   4. Текст в другом формате (изображение, объект)")
    else:
        print(f"✅ Текст найден в {total_found} местах:")
        print(f"   Параграфы: {len(found_paragraphs)}")
        print(f"   Таблицы: {len(found_tables)}")
        print(f"   Блоки: {len(found_blocks)}")
    
    print("\\n" + "=" * 80)

if __name__ == "__main__":
    deep_search()