#!/usr/bin/env python3
"""
Симуляция полного пользовательского workflow:
1. Анализ документа
2. Выбор всех найденных элементов  
3. Селективная анонимизация
Цель: найти где теряются данные между 53 → 26
"""

import tempfile
import json
from docx import Document
import sys
import os

# Добавляем путь к модулям
sys.path.append(os.path.join(os.path.dirname(__file__), 'unified_document_service', 'app'))

from full_anonymizer import FullAnonymizer
from block_builder import BlockBuilder
from rule_adapter import RuleEngineAdapter

def create_complex_test_document():
    """Создает сложный тестовый документ с различными типами данных"""
    doc = Document()
    
    # Добавляем множество параграфов с разными типами данных
    doc.add_paragraph("МИНИСТЕРСТВО ИНФОРМАЦИОННОГО РАЗВИТИЯ И СВЯЗИ ПЕРМСКОГО КРАЯ")
    doc.add_paragraph("Контактная информация:")
    doc.add_paragraph("Email: admin@ministry.gov.ru")
    doc.add_paragraph("Телефон: +7 (342) 123-45-67")
    doc.add_paragraph("Дополнительный email: support@ministry.gov.ru")
    doc.add_paragraph("Организация: ООО «Технические решения»")
    doc.add_paragraph("Адрес: г. Пермь, ул. Ленина, д. 10")
    doc.add_paragraph("ИНН: 1234567890")
    doc.add_paragraph("ОГРН: 1234567890123")
    doc.add_paragraph("КПП: 123456789")
    
    # Добавляем таблицу с данными
    table = doc.add_table(rows=3, cols=3)
    table.cell(0, 0).text = "Ответственный"
    table.cell(0, 1).text = "Должность"  
    table.cell(0, 2).text = "Контакты"
    table.cell(1, 0).text = "И.И. Иванов"
    table.cell(1, 1).text = "Директор"
    table.cell(1, 2).text = "ivanov@company.ru"
    table.cell(2, 0).text = "П.П. Петров"
    table.cell(2, 1).text = "Заместитель"
    table.cell(2, 2).text = "+7 123 456-78-90"
    
    # Еще несколько параграфов
    doc.add_paragraph("Документ создан 15.11.2023 года.")
    doc.add_paragraph("СНИЛС: 123-456-789-12")
    doc.add_paragraph("Паспорт: 1234 567890")
    doc.add_paragraph("Банковские реквизиты: 12345678901234567890")
    doc.add_paragraph("БИК: 044525225")
    
    # Сохраняем во временный файл
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        return tmp.name

def simulate_full_workflow():
    """Симулирует полный workflow пользователя"""
    print("🎯 Симуляция полного пользовательского workflow...")
    
    # Создаем документ
    test_file = create_complex_test_document()
    print(f"📄 Создан тестовый документ: {test_file}")
    
    try:
        # ШАГ 1: Анализ документа (как в /analyze_document)
        print(f"\n📊 ШАГ 1: Анализ документа...")
        
        doc = Document(test_file)
        builder = BlockBuilder()
        blocks = builder.build_blocks(doc)
        
        # Применяем правила поиска
        rule_engine = RuleEngineAdapter("patterns/sensitive_patterns.xlsx")
        processed_blocks = rule_engine.apply_rules_to_blocks(blocks)
        
        # Собираем найденные элементы
        found_items = []
        for block in processed_blocks:
            if 'sensitive_patterns' in block:
                for pattern in block['sensitive_patterns']:
                    item = {
                        'block_id': block['block_id'],
                        'original_value': pattern['original_value'],
                        'uuid': pattern['uuid'],
                        'position': pattern['position'],
                        'category': pattern['category'],
                        'confidence': pattern.get('confidence', 1.0)
                    }
                    found_items.append(item)
        
        print(f"✅ Анализ завершен: найдено {len(found_items)} элементов")
        
        # Показываем первые 10 найденных элементов
        print(f"🔍 Найденные элементы:")
        for i, item in enumerate(found_items[:10]):
            print(f"   {i+1}. '{item['original_value']}' в {item['block_id']} ({item['category']})")
        if len(found_items) > 10:
            print(f"   ... и еще {len(found_items) - 10} элементов")
        
        # ШАГ 2: Пользователь выбирает ВСЕ элементы для анонимизации
        print(f"\n👤 ШАГ 2: Пользователь выбирает ВСЕ {len(found_items)} элементов")
        selected_items = found_items  # Пользователь выбрал все
        
        # ШАГ 3: Селективная анонимизация (как в /anonymize_selected)
        print(f"\n🔒 ШАГ 3: Селективная анонимизация...")
        
        # Создаем выходной файл
        output_file = test_file.replace(".docx", "_anonymized.docx")
        
        # Используем FullAnonymizer для селективной анонимизации
        anonymizer = FullAnonymizer()
        result = anonymizer.anonymize_selected_items(
            input_path=test_file,
            output_path=output_file,
            selected_items=selected_items
        )
        
        print(f"✅ Селективная анонимизация завершена")
        print(f"📊 Результат:")
        print(f"   • Статус: {result.get('status', 'N/A')}")
        print(f"   • Обработано элементов: {result.get('selected_items_count', 'N/A')}")
        print(f"   • Выполнено замен: {result.get('replacements_applied', 'N/A')}")
        
        # ШАГ 4: Анализ потерь
        print(f"\n📈 ШАГ 4: Анализ потерь...")
        
        original_count = len(found_items)
        selected_count = len(selected_items)
        processed_count = result.get('selected_items_count', 0)
        applied_count = result.get('replacements_applied', 0)
        
        print(f"🔍 Этапы обработки:")
        print(f"   1️⃣ Найдено при анализе: {original_count}")
        print(f"   2️⃣ Выбрано пользователем: {selected_count}")
        print(f"   3️⃣ Получено FullAnonymizer: {processed_count}")  
        print(f"   4️⃣ Фактически заменено: {applied_count}")
        
        # Рассчитываем потери
        loss_1 = original_count - selected_count  # Потери при выборе пользователя
        loss_2 = selected_count - processed_count  # Потери при передаче в FullAnonymizer
        loss_3 = processed_count - applied_count   # Потери при применении FormatterApplier
        
        print(f"\n⚠️ Анализ потерь:")
        if loss_1 != 0:
            print(f"   📤 Потеря на этапе выбора пользователя: {loss_1} элементов")
        if loss_2 != 0:
            print(f"   📥 Потеря при передаче в FullAnonymizer: {loss_2} элементов")
        if loss_3 != 0:
            print(f"   🔧 Потеря при применении FormatterApplier: {loss_3} элементов")
            
        if applied_count == original_count:
            print(f"✅ Потерь данных нет - система работает идеально!")
        else:
            total_loss = original_count - applied_count
            print(f"❌ Общая потеря данных: {total_loss} элементов ({total_loss/original_count*100:.1f}%)")
            
            if loss_3 > 0:
                print(f"\n💡 Основная проблема: FormatterApplier не может заменить {loss_3} элементов")
                print(f"   Возможные причины:")
                print(f"   • Элемент не найден в документе")
                print(f"   • Позиционирование неточное")
                print(f"   • Проблемы с типом элемента")
            
        print(f"\n🎯 ЗАКЛЮЧЕНИЕ:")
        if applied_count >= original_count * 0.8:  # 80%+ успешности
            print(f"✅ Система работает приемлемо ({applied_count/original_count*100:.1f}% успешности)")
        else:
            print(f"❌ Система требует доработки ({applied_count/original_count*100:.1f}% успешности)")
            print(f"   Цель: довести до 90%+ успешности")
        
    except Exception as e:
        print(f"❌ Ошибка: {str(e)}")
        import traceback
        traceback.print_exc()
    finally:
        # Удаляем временные файлы
        for file_path in [test_file, test_file.replace(".docx", "_anonymized.docx")]:
            if os.path.exists(file_path):
                os.remove(file_path)

if __name__ == "__main__":
    simulate_full_workflow()