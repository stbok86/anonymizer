#!/usr/bin/env python3
"""
ПОИСК РЕАЛЬНОЙ ПРИЧИНЫ ДУБЛИРОВАНИЯ UUID
========================================

Проблема: в анонимизированном документе UUID "545094b7-602f-4e1d-9e95-95142918f380"
появляется 3 раза, хотя каждое вхождение должно иметь свой уникальный UUID.

Анализируем где именно происходит неправильное переиспользование UUID.
"""

import requests
import json
import re
from docx import Document

def find_real_uuid_duplication_cause():
    """Находит реальную причину дублирования UUID"""
    
    print("🔍 ПОИСК РЕАЛЬНОЙ ПРИЧИНЫ ДУБЛИРОВАНИЯ UUID")
    print("=" * 60)
    
    test_file = "unified_document_service/test_docs/test_01_1_4_S.docx"
    target_text = "14 августа 2023"
    
    # ЭТАП 1: Анализируем что приходит из анализа
    print("📊 ЭТАП 1: АНАЛИЗ ИСХОДНЫХ ДАННЫХ")
    print("-" * 40)
    
    try:
        with open(test_file, 'rb') as f:
            files = {'file': ('test_document.docx', f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
            data = {
                'patterns_file': 'patterns/sensitive_patterns.xlsx',
                'include_nlp': 'false'
            }
            
            response = requests.post(
                "http://localhost:8002/analyze_document",
                files=files,
                data=data,
                timeout=30
            )
        
        if response.status_code == 200:
            result = response.json()
            found_items = result.get('found_items', [])
            
            # Фильтруем только элементы с целевым текстом в table_2
            table_2_items = [
                item for item in found_items 
                if item.get('original_value') == target_text and item.get('block_id') == 'table_2'
            ]
            
            print(f"✅ Найдено элементов '{target_text}' в table_2: {len(table_2_items)}")
            
            for i, item in enumerate(table_2_items):
                print(f"📄 table_2 элемент {i+1}:")
                print(f"   block_id: {item.get('block_id')}")
                print(f"   original_value: '{item.get('original_value')}'")
                print(f"   uuid: {item.get('uuid')}")
                print(f"   position: {item.get('position')}")
                print()
            
            # ЭТАП 2: Отправляем ТОЛЬКО table_2 элементы на анонимизацию
            print("🔧 ЭТАП 2: АНОНИМИЗАЦИЯ ТОЛЬКО TABLE_2 ЭЛЕМЕНТОВ")
            print("-" * 50)
            
            selected_items = []
            for item in table_2_items:
                selected_item = {
                    'block_id': item.get('block_id', ''),
                    'original_value': item.get('original_value', ''),
                    'uuid': item.get('uuid', ''),
                    'position': item.get('position', {}),
                    'category': item.get('category', ''),
                    'confidence': item.get('confidence', 1.0)
                }
                selected_items.append(selected_item)
            
            print(f"📝 Отправляем на анонимизацию: {len(selected_items)} элементов")
            print("   UUID элементов:")
            for i, item in enumerate(selected_items):
                print(f"     {i+1}. {item['uuid']}")
            
            # Создаем копию файла для анонимизации
            import tempfile
            import shutil
            
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_input:
                shutil.copy2(test_file, temp_input.name)
                temp_input_path = temp_input.name
            
            try:
                with open(temp_input_path, 'rb') as f:
                    files = {'file': ('test_document.docx', f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
                    data = {
                        'patterns_file': 'patterns/sensitive_patterns.xlsx',
                        'selected_items': json.dumps(selected_items)
                    }
                    
                    response = requests.post(
                        "http://localhost:8002/anonymize_selected",
                        files=files,
                        data=data,
                        timeout=30
                    )
                
                if response.status_code == 200:
                    result = response.json()
                    print(f"✅ Анонимизация завершена: {result.get('status')}")
                    print(f"🔢 Замен применено: {result.get('replacements_applied')}")
                    
                    # Сохраняем результат для анализа
                    if 'anonymized_document_base64' in result:
                        import base64
                        doc_data = base64.b64decode(result['anonymized_document_base64'])
                        output_path = "debug_uuid_test_output.docx"
                        
                        with open(output_path, 'wb') as f:
                            f.write(doc_data)
                        
                        print(f"💾 Результат сохранен в: {output_path}")
                        
                        # ЭТАП 3: Анализируем результат
                        print(f"\n🔍 ЭТАП 3: АНАЛИЗ РЕЗУЛЬТАТА")
                        print("-" * 30)
                        
                        doc = Document(output_path)
                        uuid_pattern = re.compile(r'[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}')
                        
                        # Ищем UUID в table_2
                        if len(doc.tables) > 2:
                            table_2 = doc.tables[2]
                            found_uuids = []
                            
                            for row_idx, row in enumerate(table_2.rows):
                                for cell_idx, cell in enumerate(row.cells):
                                    text = cell.text.strip()
                                    uuids = uuid_pattern.findall(text.lower())
                                    if uuids:
                                        for uuid_found in uuids:
                                            found_uuids.append({
                                                'uuid': uuid_found,
                                                'row': row_idx,
                                                'cell': cell_idx,
                                                'text': text[:100] + ('...' if len(text) > 100 else '')
                                            })
                            
                            print(f"🎯 UUID найдено в table_2: {len(found_uuids)}")
                            
                            # Группируем по UUID
                            uuid_groups = {}
                            for item in found_uuids:
                                uuid_val = item['uuid']
                                if uuid_val not in uuid_groups:
                                    uuid_groups[uuid_val] = []
                                uuid_groups[uuid_val].append(item)
                            
                            for uuid_val, items in uuid_groups.items():
                                print(f"\n🔄 UUID: {uuid_val}")
                                print(f"   Использован: {len(items)} раз(а)")
                                for item in items:
                                    print(f"   • Строка {item['row']}, ячейка {item['cell']}: {item['text']}")
                                
                                # КЛЮЧЕВОЙ ВОПРОС: совпадает ли этот UUID с отправленными?
                                matching_sent = [s for s in selected_items if s['uuid'] == uuid_val]
                                if matching_sent:
                                    print(f"   ✅ Соответствует отправленному UUID")
                                else:
                                    print(f"   ❌ НЕ СООТВЕТСТВУЕТ отправленным UUID!")
                                    print(f"       Отправленные: {[s['uuid'] for s in selected_items]}")
                        
                        print(f"\n💡 ДИАГНОЗ:")
                        print(f"   • Отправлено элементов: {len(selected_items)}")
                        print(f"   • Все с РАЗНЫМИ UUID: {len(set(s['uuid'] for s in selected_items)) == len(selected_items)}")
                        print(f"   • Замен применено: {result.get('replacements_applied')}")
                        print(f"   • UUID в результате: {len(uuid_groups)} уникальных")
                        
                        if len(uuid_groups) == 1 and len(selected_items) > 1:
                            print(f"   🚨 ПРОБЛЕМА: Отправлены разные UUID, получен один!")
                            print(f"   💡 ПРИЧИНА: Где-то в процессе замены UUID перезаписывается")
                else:
                    print(f"❌ Ошибка анонимизации: {response.status_code} - {response.text}")
            
            finally:
                import os
                if os.path.exists(temp_input_path):
                    os.unlink(temp_input_path)
        else:
            print(f"❌ Ошибка анализа: {response.status_code}")
            
    except Exception as e:
        print(f"❌ Ошибка: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    find_real_uuid_duplication_cause()