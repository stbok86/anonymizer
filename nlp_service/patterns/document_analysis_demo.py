#!/usr/bin/env python3
"""
Анализ реального документа: демонстрация различий spaCy NER vs Phrase Matcher
Упрощенная версия без spaCy для демонстрации концепций
"""

from docx import Document
import os
import re
import time
from typing import List, Dict, Any, Set

class DocumentAnalysisDemo:
    """Демонстрация концепций NER vs Phrase Matcher на реальном документе"""
    
    def __init__(self):
        self.doc_path = r"C:\Projects\Anonymizer\unified_document_service\test_docs\test_01_1_4_SD.docx"
        
        # Словарь для Phrase Matcher (точные совпадения)
        self.government_phrases = {
            # Федеральные органы
            "министерство внутренних дел российской федерации",
            "мвд россии", "мвд рф",
            "федеральная налоговая служба", "фнс россии",
            "роскомнадзор",
            "министерство здравоохранения российской федерации",
            "минздрав россии", "минздрав рф",
            "министерство образования и науки российской федерации",
            "минобрнауки россии", "минобрнауки рф",
            "федеральная служба безопасности", "фсб россии",
            "министерство чрезвычайных ситуаций", "мчс россии",
            "федеральная антимонопольная служба", "фас россии",
            "правительство российской федерации", "правительство рф",
            "администрация президента российской федерации",
            
            # Региональные органы
            "министерство информационного развития и связи пермского края",
            "правительство пермского края",
            "администрация губернатора пермского края",
            "департамент образования и науки кировской области",
            "управление внутренних дел по свердловской области",
            "департамент здравоохранения города москвы",
            "комитет по образованию санкт-петербурга",
            
            # Муниципальные органы
            "администрация города перми",
            "городская дума города перми",
            "мэрия города екатеринбурга",
            "администрация ленинского района",
            
            # Судебные и силовые
            "прокуратура пермского края",
            "следственный комитет российской федерации"
        }
        
        # Паттерны для симуляции NER (regex-подход)
        self.ner_patterns = [
            # Министерства
            r'\bминистерство\s+[\w\s]+(?:российской\s+федерации|рф|россии|края|области|республики)\b',
            # Департаменты и управления
            r'\b(?:департамент|управление|комитет)\s+[\w\s]+(?:края|области|города|района|рф|россии)?\b',
            # Федеральные службы
            r'\bфедеральная\s+(?:служба|антимонопольная\s+служба|налоговая\s+служба)\s*[\w\s]*\b',
            # Администрации
            r'\bадминистрация\s+(?:президента|губернатора|города|района)\s*[\w\s]*\b',
            # Правительство
            r'\bправительство\s+(?:российской\s+федерации|рф|россии|[\w\s]+края|[\w\s]+области)\b',
            # Аббревиатуры
            r'\b(?:мвд|фнс|фсб|мчс|фас)\s+(?:россии|рф)\b',
            r'\b(?:роскомнадзор|росстат|ростуризм|росреестр)\b',
            # Суды и прокуратура
            r'\b(?:прокуратура|суд)\s+[\w\s]+(?:края|области|района|рф)\b',
            # Думы и советы
            r'\b(?:дума|совет)\s+[\w\s]*(?:города|района|края|области)\b'
        ]
    
    def extract_document_text(self) -> str:
        """Извлекает текст из Word документа"""
        
        if not os.path.exists(self.doc_path):
            raise FileNotFoundError(f"Документ не найден: {self.doc_path}")
        
        print(f"📄 Загружаем документ: {os.path.basename(self.doc_path)}")
        
        try:
            doc = Document(self.doc_path)
            full_text = []
            
            # Извлекаем текст из параграфов
            paragraph_count = 0
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text.strip())
                    paragraph_count += 1
            
            # Извлекаем текст из таблиц
            table_count = 0
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(" | ".join(row_text))
                table_count += 1
            
            document_text = "\n".join(full_text)
            
            print(f"✅ Документ загружен:")
            print(f"   Общая длина: {len(document_text)} символов")
            print(f"   Параграфов с текстом: {paragraph_count}")
            print(f"   Таблиц: {table_count}")
            
            # Показываем первые 300 символов для контекста
            print(f"\n📖 Начало документа:")
            print(f"   {document_text[:300]}...")
            
            return document_text
            
        except Exception as e:
            raise Exception(f"Ошибка чтения документа: {e}")
    
    def simulate_phrase_matcher(self, text: str) -> List[Dict[str, Any]]:
        """Симулирует работу Phrase Matcher - точный поиск известных фраз"""
        
        print(f"\n📚 PHRASE MATCHER ПОДХОД")
        print(f"{'='*60}")
        print(f"🔍 Ищем точные совпадения с {len(self.government_phrases)} известными организациями...")
        
        start_time = time.time()
        
        text_lower = text.lower()
        results = []
        
        # Поиск каждой известной фразы
        for phrase in self.government_phrases:
            # Ищем точные совпадения (с учетом границ слов)
            pattern = r'\b' + re.escape(phrase) + r'\b'
            matches = re.finditer(pattern, text_lower, re.IGNORECASE)
            
            for match in matches:
                # Получаем оригинальный текст (с правильным регистром)
                original_text = text[match.start():match.end()]
                
                results.append({
                    'text': original_text,
                    'start': match.start(),
                    'end': match.end(),
                    'confidence': 0.95,  # Высокая уверенность для точных совпадений
                    'method': 'phrase_matcher',
                    'matched_phrase': phrase
                })
        
        processing_time = time.time() - start_time
        
        # Удаляем дубликаты (если одна организация найдена несколько раз)
        unique_results = []
        seen_positions = set()
        
        for result in results:
            pos_key = (result['start'], result['end'])
            if pos_key not in seen_positions:
                unique_results.append(result)
                seen_positions.add(pos_key)
        
        print(f"⏱️ Время обработки: {processing_time*1000:.1f} мс")
        print(f"📊 Найдено совпадений: {len(unique_results)}")
        
        print(f"\n🏛️ НАЙДЕННЫЕ ГОСУДАРСТВЕННЫЕ ОРГАНИЗАЦИИ:")
        for i, result in enumerate(unique_results, 1):
            print(f"   {i:2d}. 🏛️ '{result['text']}' (confidence: {result['confidence']})")
            print(f"       Позиция: {result['start']}-{result['end']}")
            
            # Показываем контекст
            context_start = max(0, result['start'] - 40)
            context_end = min(len(text), result['end'] + 40)
            context = text[context_start:context_end].replace('\n', ' ')
            print(f"       Контекст: ...{context}...")
        
        print(f"\n🎯 ОСОБЕННОСТИ PHRASE MATCHER:")
        print(f"   ✅ Очень высокая точность (нет ложных срабатываний)")
        print(f"   ✅ Очень быстрая работа")
        print(f"   ✅ Находит точные совпадения даже в сложном тексте")
        print(f"   ❌ НЕ найдет неизвестные организации")
        print(f"   ❌ НЕ найдет вариации названий")
        
        return unique_results
    
    def simulate_spacy_ner(self, text: str) -> List[Dict[str, Any]]:
        """Симулирует работу spaCy NER - интеллектуальный анализ"""
        
        print(f"\n🤖 SPACY NER ПОДХОД (СИМУЛЯЦИЯ)")
        print(f"{'='*60}")
        print(f"🔍 Анализируем текст с помощью паттернов и эвристик...")
        
        start_time = time.time()
        
        all_organizations = []
        
        # Применяем каждый NER паттерн
        for i, pattern in enumerate(self.ner_patterns, 1):
            matches = re.finditer(pattern, text, re.IGNORECASE)
            
            for match in matches:
                org_text = match.group().strip()
                confidence = self._calculate_ner_confidence(org_text, text, match)
                
                all_organizations.append({
                    'text': org_text,
                    'start': match.start(),
                    'end': match.end(),
                    'confidence': confidence,
                    'method': 'spacy_ner_simulation',
                    'pattern_id': i,
                    'is_government': True  # Все наши паттерны для госорганов
                })
        
        # Дополнительный поиск потенциальных организаций
        potential_orgs = self._find_potential_organizations(text)
        all_organizations.extend(potential_orgs)
        
        processing_time = time.time() - start_time
        
        # Удаляем пересекающиеся результаты, оставляем лучшие по confidence
        filtered_results = self._remove_overlapping_detections(all_organizations)
        
        print(f"⏱️ Время обработки: {processing_time*1000:.1f} мс")
        print(f"📊 Всего найдено: {len(all_organizations)} совпадений")
        print(f"📊 После фильтрации: {len(filtered_results)} организаций")
        
        # Сортируем по уверенности
        filtered_results.sort(key=lambda x: x['confidence'], reverse=True)
        
        print(f"\n🏛️ НАЙДЕННЫЕ ОРГАНИЗАЦИИ:")
        government_count = 0
        for i, result in enumerate(filtered_results, 1):
            is_gov = result.get('is_government', False)
            gov_marker = "🏛️" if is_gov else "🏢"
            
            print(f"   {i:2d}. {gov_marker} '{result['text']}' (confidence: {result['confidence']:.2f})")
            print(f"       Позиция: {result['start']}-{result['end']}")
            
            if is_gov:
                government_count += 1
            
            # Показываем контекст
            context_start = max(0, result['start'] - 40)
            context_end = min(len(text), result['end'] + 40)
            context = text[context_start:context_end].replace('\n', ' ')
            print(f"       Контекст: ...{context}...")
        
        print(f"\n📊 Из них государственных: {government_count}")
        
        print(f"\n🎯 ОСОБЕННОСТИ SPACY NER:")
        print(f"   ✅ Может найти НЕИЗВЕСТНЫЕ организации")
        print(f"   ✅ Понимает контекст и вариации")
        print(f"   ✅ Обобщает паттерны на новые случаи")
        print(f"   ❌ Может давать ложные срабатывания")
        print(f"   ❌ Медленнее чем Phrase Matcher")
        
        return filtered_results
    
    def _calculate_ner_confidence(self, org_text: str, full_text: str, match) -> float:
        """Рассчитывает уверенность для NER детекции"""
        
        base_confidence = 0.70
        
        # Бонус за длину названия
        length_bonus = min(0.20, len(org_text.split()) * 0.04)
        
        # Бонус за заглавные буквы
        caps_bonus = 0.05 if org_text[0].isupper() else 0
        
        # Бонус за ключевые слова
        keyword_bonus = 0
        gov_keywords = ['министерство', 'департамент', 'управление', 'федеральная', 'администрация']
        for keyword in gov_keywords:
            if keyword in org_text.lower():
                keyword_bonus = 0.15
                break
        
        # Бонус за контекст
        context_start = max(0, match.start() - 100)
        context_end = min(len(full_text), match.end() + 100)
        context = full_text[context_start:context_end].lower()
        
        context_bonus = 0
        context_markers = ['сообщил', 'объявил', 'утвердил', 'провел', 'постановил']
        for marker in context_markers:
            if marker in context:
                context_bonus = 0.05
                break
        
        return min(0.95, base_confidence + length_bonus + caps_bonus + keyword_bonus + context_bonus)
    
    def _find_potential_organizations(self, text: str) -> List[Dict[str, Any]]:
        """Ищет потенциальные организации по общим признакам"""
        
        potential_orgs = []
        
        # Ищем слова с заглавной буквы, которые могут быть названиями организаций
        org_indicators = [
            r'\b[А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+(?:\s+[А-ЯЁ][а-яё]+)*\b',  # Несколько слов с заглавной
        ]
        
        for pattern in org_indicators:
            matches = re.finditer(pattern, text)
            
            for match in matches:
                candidate = match.group().strip()
                
                # Фильтруем только потенциальные организации
                if self._is_likely_organization(candidate):
                    confidence = self._calculate_ner_confidence(candidate, text, match)
                    
                    potential_orgs.append({
                        'text': candidate,
                        'start': match.start(),
                        'end': match.end(),
                        'confidence': confidence * 0.7,  # Снижаем уверенность для потенциальных
                        'method': 'spacy_ner_potential',
                        'is_government': self._is_likely_government(candidate)
                    })
        
        return potential_orgs
    
    def _is_likely_organization(self, text: str) -> bool:
        """Определяет, может ли текст быть названием организации"""
        
        # Простая эвристика
        words = text.split()
        
        # Слишком короткие или длинные названия
        if len(words) < 2 or len(words) > 8:
            return False
        
        # Исключаем обычные фразы
        common_phrases = ['в том числе', 'в связи', 'по вопросам', 'в соответствии']
        if any(phrase in text.lower() for phrase in common_phrases):
            return False
        
        return True
    
    def _is_likely_government(self, org_text: str) -> bool:
        """Определяет, является ли организация государственной"""
        
        org_lower = org_text.lower()
        gov_keywords = [
            'министерство', 'департамент', 'управление', 'служба', 'комитет',
            'администрация', 'правительство', 'дума', 'прокуратура', 'суд',
            'федеральная', 'государственная', 'муниципальная', 'региональная'
        ]
        
        return any(keyword in org_lower for keyword in gov_keywords)
    
    def _remove_overlapping_detections(self, detections: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Удаляет пересекающиеся детекции"""
        
        if not detections:
            return []
        
        # Сортируем по уверенности
        detections.sort(key=lambda x: x['confidence'], reverse=True)
        
        filtered = []
        
        for detection in detections:
            is_overlapping = False
            
            for existing in filtered:
                if self._positions_overlap(detection, existing):
                    is_overlapping = True
                    break
            
            if not is_overlapping:
                filtered.append(detection)
        
        return filtered
    
    def _positions_overlap(self, det1: Dict, det2: Dict) -> bool:
        """Проверяет пересечение позиций"""
        return not (det1['end'] <= det2['start'] or det2['end'] <= det1['start'])
    
    def compare_results(self, phrase_results: List[Dict], ner_results: List[Dict], text: str):
        """Сравнивает результаты двух подходов"""
        
        print(f"\n🆚 СРАВНИТЕЛЬНЫЙ АНАЛИЗ НА РЕАЛЬНОМ ДОКУМЕНТЕ")
        print(f"{'='*80}")
        
        # Нормализуем тексты для сравнения
        phrase_texts = {result['text'].lower().strip() for result in phrase_results}
        ner_texts = {result['text'].lower().strip() for result in ner_results}
        
        # Анализируем пересечения
        both_found = phrase_texts & ner_texts
        phrase_only = phrase_texts - ner_texts
        ner_only = ner_texts - phrase_texts
        
        print(f"📊 СТАТИСТИКА ДЛЯ ДАННОГО ДОКУМЕНТА:")
        print(f"   📚 Phrase Matcher нашел: {len(phrase_results)} организаций")
        print(f"   🤖 NER симуляция нашла: {len(ner_results)} организаций")
        print(f"   🎯 Найдены обоими: {len(both_found)} организаций")
        print(f"   📚 Только Phrase Matcher: {len(phrase_only)} организаций") 
        print(f"   🤖 Только NER: {len(ner_only)} организаций")
        
        if both_found:
            print(f"\n✅ НАЙДЕНЫ ОБОИМИ МЕТОДАМИ:")
            for org in sorted(both_found):
                print(f"   • {org}")
        
        if phrase_only:
            print(f"\n📚 ТОЛЬКО PHRASE MATCHER:")
            for org in sorted(phrase_only):
                print(f"   • {org}")
                print(f"     → Точное совпадение с известной организацией")
        
        if ner_only:
            print(f"\n🤖 ТОЛЬКО NER ПОДХОД:")
            for org in sorted(ner_only):
                # Находим полную информацию об организации
                full_info = next(r for r in ner_results if r['text'].lower().strip() == org)
                gov_status = "🏛️ госорган" if full_info.get('is_government') else "🏢 коммерческая"
                print(f"   • {org} {gov_status} (conf: {full_info['confidence']:.2f})")
        
        # Анализ производительности
        phrase_time = sum(0.1 for _ in phrase_results)  # Примерное время
        ner_time = sum(0.5 for _ in ner_results)  # NER медленнее
        
        print(f"\n⚡ ВЫВОДЫ ПО ПРОИЗВОДИТЕЛЬНОСТИ:")
        print(f"   📚 Phrase Matcher: быстрый, точный для известных названий")
        print(f"   🤖 NER подход: медленнее, но находит больше вариантов")
        
        print(f"\n💡 РЕКОМЕНДАЦИИ ДЛЯ ВАШЕГО ДОКУМЕНТА:")
        
        if len(phrase_results) > 0 and len(ner_results) > len(phrase_results):
            print(f"   1. Phrase Matcher нашел {len(phrase_results)} точных совпадений")
            print(f"   2. NER нашел дополнительно {len(ner_results) - len(phrase_results)} потенциальных организаций")
            print(f"   3. Рекомендуется использовать ОБА подхода:")
            print(f"      • Phrase Matcher для высокоточного поиска известных госорганов")
            print(f"      • NER для обнаружения неизвестных или вариативных названий")
        
        elif len(phrase_results) > 0:
            print(f"   • В документе преобладают известные госорганы")
            print(f"   • Phrase Matcher показывает отличные результаты")
            print(f"   • NER можно использовать как дополнение")
        
        else:
            print(f"   • В документе мало известных госорганов")
            print(f"   • NER подход более эффективен для этого типа документа")
            print(f"   • Стоит расширить словарь Phrase Matcher")

def run_real_document_demo():
    """Запускает демонстрацию на реальном документе"""
    
    print(f"🔬 АНАЛИЗ РЕАЛЬНОГО ДОКУМЕНТА: SPACY NER vs PHRASE MATCHER")
    print(f"{'='*80}")
    print(f"📄 Документ: test_01_1_4_SD.docx")
    print()
    
    try:
        analyzer = DocumentAnalysisDemo()
        
        # Загружаем документ
        document_text = analyzer.extract_document_text()
        
        # Анализируем двумя подходами
        phrase_results = analyzer.simulate_phrase_matcher(document_text)
        ner_results = analyzer.simulate_spacy_ner(document_text)
        
        # Сравниваем результаты
        analyzer.compare_results(phrase_results, ner_results, document_text)
        
        print(f"\n🎓 ИТОГОВЫЕ ВЫВОДЫ:")
        print(f"{'='*50}")
        print(f"""
🔑 КЛЮЧЕВЫЕ РАЗЛИЧИЯ НА ПРАКТИКЕ:

📚 PHRASE MATCHER:
   • Работает как точный словарь
   • Ищет ТОЛЬКО известные названия организаций
   • Очень высокая скорость и точность
   • НЕ может найти неизвестные организации
   • Идеален для стандартных документов с типовыми названиями

🤖 SPACY NER:
   • Работает как интеллектуальный анализатор
   • Может найти НОВЫЕ и неизвестные организации
   • Понимает контекст и морфологию
   • Может давать ложные срабатывания
   • Идеален для разнообразных документов и исследований

💡 ПРАКТИЧЕСКИЕ РЕКОМЕНДАЦИИ:

1. Для МАКСИМАЛЬНОГО ПОКРЫТИЯ → используйте ОБА подхода
2. Для ВЫСОКОЙ СКОРОСТИ → используйте только Phrase Matcher  
3. Для ПОИСКА НОВЫХ ОРГАНИЗАЦИЙ → используйте NER
4. Для ГОСУДАРСТВЕННЫХ ДОКУМЕНТОВ → комбинация дает лучший результат

🎯 ДЛЯ ВАШЕЙ СИСТЕМЫ АНОНИМИЗАЦИИ:
   Оптимальная стратегия - начать с Phrase Matcher для известных 
   госорганов, затем добавить NER для обнаружения пропущенных случаев.
        """)
        
    except Exception as e:
        print(f"❌ Ошибка анализа: {e}")
        print(f"\n💡 Возможные причины:")
        print(f"   • Документ не найден по указанному пути")
        print(f"   • Нет прав на чтение файла")
        print(f"   • Документ поврежден или не в формате .docx")

if __name__ == "__main__":
    run_real_document_demo()