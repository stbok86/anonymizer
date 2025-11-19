#!/usr/bin/env python3
"""
Анализ реального документа для демонстрации spaCy NER vs Phrase Matcher
"""

import spacy
from spacy.matcher import PhraseMatcher
from docx import Document
import os
import time
from typing import List, Dict, Any

class RealDocumentAnalysis:
    """Анализ реального документа для сравнения методов"""
    
    def __init__(self):
        print("🔄 Загружаем spaCy модель...")
        self.nlp = spacy.load("ru_core_news_lg")
        
        # Настраиваем Phrase Matcher с государственными организациями
        self.phrase_matcher = PhraseMatcher(self.nlp.vocab, attr="LOWER")
        self._setup_government_phrases()
        
        self.doc_path = r"C:\Projects\Anonymizer\unified_document_service\test_docs\test_01_1_4_SD.docx"
    
    def _setup_government_phrases(self):
        """Настраиваем Phrase Matcher с известными госорганизациями"""
        
        government_orgs = [
            # Федеральные органы
            "Министерство внутренних дел Российской Федерации",
            "МВД России", "МВД РФ",
            "Федеральная налоговая служба", "ФНС России",
            "Роскомнадзор", "Федеральная служба по надзору в сфере связи",
            "Министерство здравоохранения Российской Федерации",
            "Минздрав России", "Минздрав РФ",
            "Министерство образования и науки Российской Федерации",
            "Минобрнауки России", "Минобрнауки РФ",
            "Федеральная служба безопасности", "ФСБ России",
            "Министерство чрезвычайных ситуаций", "МЧС России",
            "Федеральная антимонопольная служба", "ФАС России",
            "Правительство Российской Федерации", "Правительство РФ",
            "Администрация Президента Российской Федерации",
            
            # Региональные органы
            "Министерство информационного развития и связи Пермского края",
            "Правительство Пермского края",
            "Администрация губернатора Пермского края",
            "Департамент образования и науки Кировской области",
            "Управление внутренних дел по Свердловской области",
            "Департамент здравоохранения города Москвы",
            "Комитет по образованию Санкт-Петербурга",
            
            # Муниципальные органы
            "Администрация города Перми",
            "Городская дума города Перми",
            "Мэрия города Екатеринбурга",
            "Администрация Ленинского района",
            
            # Судебные органы
            "Верховный суд Российской Федерации",
            "Конституционный суд Российской Федерации",
            "Арбитражный суд Пермского края",
            "Пермский районный суд",
            
            # Силовые структуры
            "Прокуратура Пермского края",
            "Следственный комитет Российской Федерации",
            "Управление Федеральной службы исполнения наказаний"
        ]
        
        # Создаем паттерны
        patterns = [self.nlp(org) for org in government_orgs]
        self.phrase_matcher.add("GOVERNMENT_ORG", patterns)
        
        print(f"✅ Phrase Matcher настроен с {len(government_orgs)} государственными организациями")
    
    def extract_document_text(self) -> str:
        """Извлекает текст из Word документа"""
        
        if not os.path.exists(self.doc_path):
            raise FileNotFoundError(f"Документ не найден: {self.doc_path}")
        
        print(f"📄 Загружаем документ: {os.path.basename(self.doc_path)}")
        
        try:
            doc = Document(self.doc_path)
            full_text = []
            
            # Извлекаем текст из параграфов
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text.strip())
            
            # Извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            document_text = "\n".join(full_text)
            
            print(f"✅ Документ загружен:")
            print(f"   Общая длина: {len(document_text)} символов")
            print(f"   Параграфов: {len(doc.paragraphs)}")
            print(f"   Таблиц: {len(doc.tables)}")
            
            # Показываем первые 500 символов для контекста
            print(f"\n📖 Начало документа:")
            print(f"   {document_text[:500]}...")
            
            return document_text
            
        except Exception as e:
            raise Exception(f"Ошибка чтения документа: {e}")
    
    def analyze_with_spacy_ner(self, text: str) -> Dict[str, Any]:
        """Анализирует документ с помощью spaCy NER"""
        
        print(f"\n🤖 АНАЛИЗ С ПОМОЩЬЮ SPACY NER")
        print(f"{'='*60}")
        
        start_time = time.time()
        doc = self.nlp(text)
        processing_time = time.time() - start_time
        
        # Находим все сущности
        all_entities = list(doc.ents)
        organizations = [ent for ent in doc.ents if ent.label_ == "ORG"]
        persons = [ent for ent in doc.ents if ent.label_ in ["PER", "PERSON"]]
        locations = [ent for ent in doc.ents if ent.label_ in ["LOC", "GPE"]]
        
        print(f"⏱️ Время обработки: {processing_time*1000:.1f} мс")
        print(f"📊 Найдено сущностей:")
        print(f"   Всего: {len(all_entities)}")
        print(f"   Организации (ORG): {len(organizations)}")
        print(f"   Персоны (PER): {len(persons)}")
        print(f"   Локации (LOC): {len(locations)}")
        
        print(f"\n🏛️ НАЙДЕННЫЕ ОРГАНИЗАЦИИ:")
        org_results = []
        for i, org in enumerate(organizations, 1):
            confidence = self._estimate_ner_confidence(org, text)
            
            # Определяем, является ли организация государственной
            is_government = self._is_likely_government(org.text)
            gov_marker = "🏛️" if is_government else "🏢"
            
            print(f"   {i:2d}. {gov_marker} '{org.text}' (confidence: {confidence:.2f})")
            print(f"       Позиция: {org.start_char}-{org.end_char}")
            print(f"       Контекст: ...{text[max(0, org.start_char-30):org.start_char]}[{org.text}]{text[org.end_char:org.end_char+30]}...")
            
            org_results.append({
                'text': org.text,
                'start': org.start_char,
                'end': org.end_char,
                'confidence': confidence,
                'is_government': is_government,
                'method': 'spacy_ner'
            })
        
        return {
            'results': org_results,
            'processing_time': processing_time,
            'total_entities': len(all_entities),
            'organizations': len(organizations),
            'government_orgs': len([r for r in org_results if r['is_government']])
        }
    
    def analyze_with_phrase_matcher(self, text: str) -> Dict[str, Any]:
        """Анализирует документ с помощью Phrase Matcher"""
        
        print(f"\n📚 АНАЛИЗ С ПОМОЩЬЮ PHRASE MATCHER")
        print(f"{'='*60}")
        
        start_time = time.time()
        doc = self.nlp(text)
        matches = self.phrase_matcher(doc)
        processing_time = time.time() - start_time
        
        print(f"⏱️ Время обработки: {processing_time*1000:.1f} мс")
        print(f"📊 Найдено совпадений: {len(matches)}")
        
        phrase_results = []
        
        print(f"\n🏛️ НАЙДЕННЫЕ ГОСУДАРСТВЕННЫЕ ОРГАНИЗАЦИИ:")
        for i, (match_id, start, end) in enumerate(matches, 1):
            span = doc[start:end]
            
            print(f"   {i:2d}. 🏛️ '{span.text}' (confidence: 0.95)")
            print(f"       Позиция: {span.start_char}-{span.end_char}")
            print(f"       Контекст: ...{text[max(0, span.start_char-30):span.start_char]}[{span.text}]{text[span.end_char:span.end_char+30]}...")
            
            phrase_results.append({
                'text': span.text,
                'start': span.start_char,
                'end': span.end_char,
                'confidence': 0.95,  # Высокая уверенность для точных совпадений
                'is_government': True,  # Все найденные - госорганы
                'method': 'phrase_matcher'
            })
        
        return {
            'results': phrase_results,
            'processing_time': processing_time,
            'matches_found': len(matches)
        }
    
    def compare_results(self, ner_results: Dict, phrase_results: Dict, original_text: str):
        """Сравнивает результаты двух подходов"""
        
        print(f"\n🆚 СРАВНИТЕЛЬНЫЙ АНАЛИЗ РЕЗУЛЬТАТОВ")
        print(f"{'='*80}")
        
        ner_orgs = {r['text'].lower(): r for r in ner_results['results']}
        phrase_orgs = {r['text'].lower(): r for r in phrase_results['results']}
        
        # Находим пересечения и различия
        ner_only = set(ner_orgs.keys()) - set(phrase_orgs.keys())
        phrase_only = set(phrase_orgs.keys()) - set(ner_orgs.keys())
        both_found = set(ner_orgs.keys()) & set(phrase_orgs.keys())
        
        print(f"📊 СТАТИСТИКА:")
        print(f"   spaCy NER нашел: {len(ner_orgs)} организаций")
        print(f"   Phrase Matcher нашел: {len(phrase_orgs)} организаций")
        print(f"   Найдены обоими: {len(both_found)} организаций")
        print(f"   Только NER: {len(ner_only)} организаций")
        print(f"   Только Phrase Matcher: {len(phrase_only)} организаций")
        
        if both_found:
            print(f"\n🎯 НАЙДЕНЫ ОБОИМИ МЕТОДАМИ:")
            for org_name in sorted(both_found):
                print(f"   ✅ '{ner_orgs[org_name]['text']}'")
        
        if ner_only:
            print(f"\n🤖 ТОЛЬКО SPACY NER НАШЕЛ:")
            for org_name in sorted(ner_only):
                org = ner_orgs[org_name]
                gov_status = "🏛️ (госорган)" if org['is_government'] else "🏢 (коммерческая)"
                print(f"   • '{org['text']}' {gov_status} (conf: {org['confidence']:.2f})")
        
        if phrase_only:
            print(f"\n📚 ТОЛЬКО PHRASE MATCHER НАШЕЛ:")
            for org_name in sorted(phrase_only):
                org = phrase_orgs[org_name]
                print(f"   • '{org['text']}' 🏛️ (conf: {org['confidence']:.2f})")
        
        # Производительность
        print(f"\n⚡ ПРОИЗВОДИТЕЛЬНОСТЬ:")
        ner_time = ner_results['processing_time'] * 1000
        phrase_time = phrase_results['processing_time'] * 1000
        speedup = ner_time / phrase_time if phrase_time > 0 else 0
        
        print(f"   spaCy NER: {ner_time:.1f} мс")
        print(f"   Phrase Matcher: {phrase_time:.1f} мс")
        print(f"   Phrase Matcher быстрее в {speedup:.1f} раз")
        
        # Анализ качества для госорганов
        ner_gov_orgs = len([r for r in ner_results['results'] if r['is_government']])
        phrase_gov_orgs = len(phrase_results['results'])  # Все найденные - госорганы
        
        print(f"\n🏛️ ГОСУДАРСТВЕННЫЕ ОРГАНИЗАЦИИ:")
        print(f"   spaCy NER (отфильтрованные): {ner_gov_orgs}")
        print(f"   Phrase Matcher (все): {phrase_gov_orgs}")
        
        return {
            'ner_total': len(ner_orgs),
            'phrase_total': len(phrase_orgs),
            'overlap': len(both_found),
            'ner_only': len(ner_only),
            'phrase_only': len(phrase_only),
            'ner_time': ner_time,
            'phrase_time': phrase_time
        }
    
    def _estimate_ner_confidence(self, ent, text: str) -> float:
        """Оценивает уверенность для NER сущности"""
        base_confidence = 0.75
        
        # Бонус за длину названия
        length_bonus = min(0.15, len(ent.text.split()) * 0.03)
        
        # Бонус за заглавные буквы
        caps_bonus = 0.05 if ent.text[0].isupper() else 0
        
        # Бонус за контекст (если рядом слова-маркеры)
        context_start = max(0, ent.start_char - 50)
        context_end = min(len(text), ent.end_char + 50)
        context = text[context_start:context_end].lower()
        
        context_bonus = 0
        gov_markers = ['министерство', 'департамент', 'управление', 'служба', 'комитет', 'администрация']
        for marker in gov_markers:
            if marker in context:
                context_bonus = 0.08
                break
        
        return min(0.90, base_confidence + length_bonus + caps_bonus + context_bonus)
    
    def _is_likely_government(self, org_text: str) -> bool:
        """Определяет, является ли организация государственной"""
        org_lower = org_text.lower()
        
        gov_keywords = [
            'министерство', 'департамент', 'управление', 'служба', 'комитет',
            'администрация', 'правительство', 'дума', 'прокуратура', 'суд',
            'федеральная', 'государственная', 'муниципальная', 'региональная',
            'мвд', 'фнс', 'фсб', 'мчс', 'роскомнадзор', 'ростуризм', 'росстат'
        ]
        
        return any(keyword in org_lower for keyword in gov_keywords)

def run_real_document_analysis():
    """Запускает анализ реального документа"""
    
    print(f"🔬 АНАЛИЗ РЕАЛЬНОГО ДОКУМЕНТА: SPACY NER vs PHRASE MATCHER")
    print(f"{'='*80}")
    
    try:
        analyzer = RealDocumentAnalysis()
        
        # Загружаем документ
        document_text = analyzer.extract_document_text()
        
        # Анализируем двумя методами
        ner_results = analyzer.analyze_with_spacy_ner(document_text)
        phrase_results = analyzer.analyze_with_phrase_matcher(document_text)
        
        # Сравниваем результаты
        comparison = analyzer.compare_results(ner_results, phrase_results, document_text)
        
        # Выводы
        print(f"\n🎯 ВЫВОДЫ ДЛЯ ДАННОГО ДОКУМЕНТА:")
        print(f"{'='*50}")
        
        if comparison['overlap'] > 0:
            print(f"✅ Оба метода нашли {comparison['overlap']} общих организаций")
        
        if comparison['ner_only'] > 0:
            print(f"🤖 spaCy NER нашел {comparison['ner_only']} дополнительных организаций")
            print(f"   (может включать коммерческие и неизвестные госорганы)")
        
        if comparison['phrase_only'] > 0:
            print(f"📚 Phrase Matcher нашел {comparison['phrase_only']} организаций, пропущенных NER")
            print(f"   (точные совпадения с известными госорганами)")
        
        speedup = comparison['ner_time'] / comparison['phrase_time']
        print(f"⚡ Phrase Matcher быстрее в {speedup:.1f} раз")
        
        print(f"\n💡 РЕКОМЕНДАЦИИ:")
        print(f"   1. Используйте Phrase Matcher для быстрого поиска известных госорганов")
        print(f"   2. Используйте spaCy NER для обнаружения новых и неизвестных организаций") 
        print(f"   3. Комбинируйте оба подхода для максимального покрытия")
        print(f"   4. Фильтруйте результаты NER по государственным ключевым словам")
        
    except Exception as e:
        print(f"❌ Ошибка анализа: {e}")

if __name__ == "__main__":
    run_real_document_analysis()