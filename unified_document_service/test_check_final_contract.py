#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document

anon_path = 'test_contract_diagnostic_anon.docx'

print("=" * 80)
print("ПРОВЕРКА ФИНАЛЬНОГО АНОНИМИЗИРОВАННОГО ДОКУМЕНТА")
print("=" * 80)

doc = Document(anon_path)

print("\n📝 Проверка первых 5 параграфов:")
for i, para in enumerate(doc.paragraphs[:5]):
    print(f"\nПараграф {i}:")
    print(f"  {para.text[:200]}")

print("\n\n📋 Проверка таблицы #2, строки 6 (Подрядчик):")
if len(doc.tables) > 2:
    table = doc.tables[2]
    if len(table.rows) > 6:
        cell = table.rows[6].cells[1]
        print(f"  {cell.text[:200]}")

print("\n\n🔍 Поиск любых упоминаний '13' или 'ОК-2023':")
found = False
for i, para in enumerate(doc.paragraphs):
    if '13/ОК-2023' in para.text or 'ОК-2023' in para.text:
        print(f"❌ Найдено в параграфе {i}: {para.text[:100]}")
        found = True

for t_idx, table in enumerate(doc.tables):
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            if '13/ОК-2023' in cell.text or 'ОК-2023' in cell.text:
                print(f"❌ Найдено в таблице {t_idx}, ряд {r_idx}: {cell.text[:100]}")
                found = True

if not found:
    print("✅ Оригинальные номера контрактов НЕ найдены (успешно заменены)")

print("\n" + "=" * 80)
