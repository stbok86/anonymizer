#!/usr/bin/env python3
"""
Анализ всех таблиц в документе
"""

import sys
import os
sys.path.append(os.path.join(os.path.dirname(__file__), 'app'))

from docx import Document

def analyze_all_tables():
    """
    Анализируем все таблицы в документе
    """
    doc_path = r"C:\Projects\Anonymizer\unified_document_service\test_docs\test_01_1_4.docx"
    
    print("🔍 АНАЛИЗ ВСЕХ ТАБЛИЦ В ДОКУМЕНТЕ")
    print("=" * 50)
    
    doc = Document(doc_path)
    print(f"Документ: {doc_path}")
    print(f"Всего таблиц: {len(doc.tables)}")
    
    search_text = "14 августа 2023"
    
    for table_idx, table in enumerate(doc.tables):
        print(f"\n📋 ТАБЛИЦА {table_idx} (table_{table_idx})")
        print(f"Размер: {len(table.rows)} строк × {len(table.rows[0].cells) if table.rows else 0} столбцов")
        
        found_positions = []
        
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text
                if search_text in cell_text:
                    found_positions.append((row_idx, cell_idx))
                    print(f"  ✅ Найдено в ячейке [{row_idx}][{cell_idx}]")
                    print(f"     Содержимое: '{cell_text.strip()}'")
        
        if not found_positions:
            print(f"  ❌ Дата '{search_text}' не найдена")
        
        print(f"  Всего найдено: {len(found_positions)} вхождений")
    
    # Также проверим параграфы
    print(f"\n📝 ПАРАГРАФЫ:")
    para_count = 0
    for para_idx, para in enumerate(doc.paragraphs):
        if search_text in para.text:
            para_count += 1
            print(f"  ✅ Параграф {para_idx}: '{para.text.strip()[:100]}...'")
    
    print(f"  Всего в параграфах: {para_count}")

if __name__ == "__main__":
    analyze_all_tables()