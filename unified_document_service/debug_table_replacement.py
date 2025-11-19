#!/usr/bin/env python3
"""
Тест для понимания проблемы с UUID в таблицах
"""

import sys
import os
sys.path.append(os.path.join(os.path.dirname(__file__), 'app'))

from docx import Document
from app.formatter_applier import FormatterApplier
import uuid
from docx.table import Table


def test_table_replacement_issue():
    """
    Тест проблемы с заменами в таблице
    """
    print("🔍 ТЕСТ ПРОБЛЕМЫ С UUID В ТАБЛИЦАХ")
    print("=" * 50)
    
    # Загружаем документ
    doc_path = r"C:\Projects\Anonymizer\unified_document_service\test_docs\test_01_1_4.docx"
    
    if not os.path.exists(doc_path):
        print(f"❌ Документ не найден: {doc_path}")
        return
    
    doc = Document(doc_path)
    print(f"✅ Документ загружен: {doc_path}")
    
    # Находим table_2
    tables = doc.tables
    if len(tables) < 3:
        print("❌ Недостаточно таблиц в документе")
        return
    
    table_2 = tables[2]  # table_2 - это третья таблица (индекс 2)
    print(f"✅ Найдена table_2 с {len(table_2.rows)} строками")
    
    # Создаем тестовые замены с разными UUID для одного текста
    test_replacements = []
    
    # Генерируем 3 разных UUID для текста "14 августа 2023"
    for i in range(3):
        test_uuid = str(uuid.uuid4())
        replacement = {
            'uuid': test_uuid,
            'element': table_2,
            'original_value': '14 августа 2023',
            'category': 'date',
            'block_id': f'table_2_test_{i}',
            'position': {
                'start': 500 + i * 100,
                'end': 515 + i * 100
            }
        }
        test_replacements.append(replacement)
        print(f"🔧 Создана замена #{i+1}: UUID={test_uuid[:8]}...")
    
    print("\n" + "=" * 50)
    print("🚀 ТЕСТИРУЕМ ПРИМЕНЕНИЕ ЗАМЕН")
    print("=" * 50)
    
    # Создаем FormatterApplier
    formatter = FormatterApplier()
    
    # Применяем замены одну за другой
    for i, replacement in enumerate(test_replacements):
        print(f"\n--- ЗАМЕНА #{i+1} ---")
        print(f"UUID: {replacement['uuid'][:8]}...")
        print(f"Текст: '{replacement['original_value']}'")
        
        success = formatter._apply_single_replacement(replacement)
        print(f"Результат: {'✅ Успех' if success else '❌ Ошибка'}")
    
    print("\n" + "=" * 50)
    print("🔍 АНАЛИЗИРУЕМ РЕЗУЛЬТАТ")
    print("=" * 50)
    
    # Анализируем что получилось в таблице
    date_count = 0
    uuid_found = {}
    
    for row_idx, row in enumerate(table_2.rows):
        for cell_idx, cell in enumerate(row.cells):
            cell_text = cell.text
            if '14 августа 2023' in cell_text:
                date_count += 1
                print(f"Найдена дата в ячейке [{row_idx}][{cell_idx}]: '{cell_text.strip()}'")
            
            # Ищем UUID в тексте
            for para in cell.paragraphs:
                para_text = para.text
                if 'uuid-' in para_text or any(c in para_text for c in '0123456789abcdef-' if len(para_text) > 30):
                    # Возможно содержит UUID
                    words = para_text.split()
                    for word in words:
                        if len(word) == 36 and '-' in word:  # Формат UUID
                            if word in uuid_found:
                                uuid_found[word] += 1
                            else:
                                uuid_found[word] = 1
                            print(f"UUID в ячейке [{row_idx}][{cell_idx}]: {word}")
    
    print(f"\nИтог:")
    print(f"Всего вхождений текста '14 августа 2023': {date_count}")
    print(f"Найденные UUID:")
    for uuid_str, count in uuid_found.items():
        print(f"  {uuid_str}: {count} раз(а)")
    
    if len(uuid_found) == 1 and list(uuid_found.values())[0] > 1:
        print("❌ ПРОБЛЕМА: Один UUID используется несколько раз!")
    elif len(uuid_found) == 3:
        print("✅ ОК: Каждый UUID используется один раз")
    else:
        print("❓ Неожиданный результат")


def analyze_table_content_before_replacement():
    """
    Анализируем содержимое таблицы до замены
    """
    print("\n" + "=" * 50)
    print("🔍 АНАЛИЗ СОДЕРЖИМОГО table_2 ДО ЗАМЕНЫ")
    print("=" * 50)
    
    doc_path = r"C:\Projects\Anonymizer\unified_document_service\test_docs\test_01_1_4.docx"
    doc = Document(doc_path)
    table_2 = doc.tables[2]
    
    date_positions = []
    
    for row_idx, row in enumerate(table_2.rows):
        for cell_idx, cell in enumerate(row.cells):
            cell_text = cell.text
            if '14 августа 2023' in cell_text:
                start_pos = cell_text.find('14 августа 2023')
                print(f"Ячейка [{row_idx}][{cell_idx}]: позиция {start_pos}")
                print(f"  Контекст: '{cell_text[max(0, start_pos-10):start_pos+25]}'")
                date_positions.append((row_idx, cell_idx, start_pos))
    
    print(f"\nВсего найдено {len(date_positions)} вхождений в таблице")
    return date_positions


if __name__ == "__main__":
    analyze_table_content_before_replacement()
    test_table_replacement_issue()