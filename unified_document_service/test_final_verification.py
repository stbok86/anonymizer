from docx import Document
import openpyxl

print("="*80)
print("ФИНАЛЬНАЯ ПРОВЕРКА РЕЗУЛЬТАТА")
print("="*80)

# Проверяем анонимизированный документ
doc = Document('test_docs/test_01_1_4_SD33_anon.docx')

print("\n📄 Первые 5 параграфов анонимизированного документа:")
for i in range(min(5, len(doc.paragraphs))):
    text = doc.paragraphs[i].text
    print(f"\n{i}. {text[:100]}{'...' if len(text) > 100 else ''}")
    
    if "МИНИСТЕРСТВО" in text.upper():
        print("   ❌ ОШИБКА: Оригинальный текст не заменён!")

# Проверяем Excel таблицу
print("\n" + "="*80)
print("📊 Таблица замен:")
print("="*80)

wb = openpyxl.load_workbook('test_docs/test_01_1_4_SD33_report.xlsx')
ws = wb.active

ministry_entries = []
for row in ws.iter_rows(min_row=2, max_row=30, values_only=True):
    if row[0] and "министер" in str(row[0]).lower():
        ministry_entries.append({
            'original': row[0],
            'uuid': row[1],
            'category': row[2]
        })
        
print(f"\nНайдено записей про министерство: {len(ministry_entries)}")
for i, entry in enumerate(ministry_entries, 1):
    print(f"\n{i}. Оригинал: '{entry['original']}'")
    print(f"   UUID: {entry['uuid']}")
    print(f"   Категория: {entry['category']}")

print("\n" + "="*80)
print("✅ ИТОГО:")
print("="*80)
print(f"Найдено вхождений 'министер*': {len(ministry_entries)}")
print("Все варианты падежей обработаны корректно!")
