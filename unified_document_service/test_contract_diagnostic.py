#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import requests
from docx import Document
import pandas as pd
import base64

doc_path = 'test_docs/test_01_1_4_SD33.docx'

print("=" * 80)
print("ДИАГНОСТИКА: Контракту от 14 августа 2023 г. № 13/ОК-2023")
print("=" * 80)

# Шаг 1: Проверяем оригинальный документ
print("\n📄 Шаг 1: Поиск в оригинальном документе...")
doc = Document(doc_path)

target_text = "Контракту от 14 августа 2023 г. № 13/ОК-2023"
target_short = "13/ОК-2023"

found_original = False
for i, para in enumerate(doc.paragraphs):
    if target_short in para.text:
        print(f"✅ Найдено в параграфе #{i}:")
        print(f"   '{para.text}'")
        found_original = True

# Проверяем таблицы
for t_idx, table in enumerate(doc.tables):
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            if target_short in cell.text:
                print(f"✅ Найдено в таблице #{t_idx}, ряд {r_idx}, ячейка {c_idx}:")
                print(f"   '{cell.text}'")
                found_original = True

# Шаг 2: Анонимизация
print("\n🔄 Шаг 2: Анонимизация документа...")

with open(doc_path, 'rb') as f:
    files = {'file': ('test.docx', f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
    data = {
        'patterns_file': 'patterns/sensitive_patterns.xlsx',
        'generate_excel_report': 'true'
    }
    
    response = requests.post('http://localhost:8009/anonymize_full', files=files, data=data, timeout=120)

if response.status_code == 200:
    result = response.json()
    print(f"✅ Анонимизация выполнена")
    
    # Сохраняем анонимизированный документ
    doc_data = base64.b64decode(result['files_base64']['anonymized_document_base64'])
    anon_path = 'test_contract_diagnostic_anon.docx'
    with open(anon_path, 'wb') as f:
        f.write(doc_data)
    
    # Сохраняем Excel отчет
    excel_data = base64.b64decode(result['files_base64']['excel_report_base64'])
    excel_path = 'test_contract_diagnostic_report.xlsx'
    with open(excel_path, 'wb') as f:
        f.write(excel_data)
    
    # Шаг 3: Проверяем анонимизированный документ
    print("\n🔍 Шаг 3: Проверка анонимизированного документа...")
    anon_doc = Document(anon_path)
    
    for i, para in enumerate(anon_doc.paragraphs):
        if "ОК-2023" in para.text:
            print(f"\n📝 Параграф #{i}:")
            print(f"   '{para.text}'")
    
    for t_idx, table in enumerate(anon_doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                if "ОК-2023" in cell.text:
                    print(f"\n📋 Таблица #{t_idx}, ряд {r_idx}, ячейка {c_idx}:")
                    print(f"   '{cell.text}'")
    
    # Шаг 4: Проверяем таблицу замен
    print("\n📊 Шаг 4: Проверка таблицы замен...")
    df = pd.read_excel(excel_path)
    
    # Ищем все записи связанные с "ОК-2023"
    contract_rows = df[df['Исходные данные'].str.contains('ОК-2023|августа 2023', case=False, na=False)]
    
    if len(contract_rows) > 0:
        print(f"\n✅ Найдено {len(contract_rows)} записей:")
        for idx, row in contract_rows.iterrows():
            print(f"\n{idx + 1}. Оригинал: '{row['Исходные данные']}'")
            print(f"   Замена: '{row['Замена (идентификатор)']}'")
    
    # Ищем конкретное значение
    specific = df[df['Исходные данные'].str.contains('Контракту от 14 августа', case=False, na=False)]
    if len(specific) > 0:
        print(f"\n🎯 Конкретное значение 'Контракту от 14 августа...':")
        for idx, row in specific.iterrows():
            print(f"   Оригинал: '{row['Исходные данные']}'")
            print(f"   Замена: '{row['Замена (идентификатор)']}'")
    else:
        print(f"\n⚠️ Точное значение 'Контракту от 14 августа...' НЕ найдено в таблице замен")

else:
    print(f"❌ Ошибка: {response.status_code}")
    print(response.text)

print("\n" + "=" * 80)
