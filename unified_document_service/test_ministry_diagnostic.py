"""
Диагностика: почему "МИНИСТЕРСТВО ИНФОРМАЦИОННОГО РАЗВИТИЯ И СВЯЗИ" не заменяется
"""
import os
import sys
import requests
import base64
from docx import Document
import json

sys.stdout.reconfigure(encoding='utf-8')
os.environ['PYTHONIOENCODING'] = 'utf-8'

doc_path = r'C:\Projects\Anonymizer\unified_document_service\test_docs\test_01_1_4_SD33.docx'

print("=" * 80)
print("🔍 ДИАГНОСТИКА: МИНИСТЕРСТВО ИНФОРМАЦИОННОГО РАЗВИТИЯ И СВЯЗИ")
print("=" * 80)
print()

# Шаг 1: Читаем документ
doc = Document(doc_path)

# Ищем это значение в документе
target_text = "МИНИСТЕРСТВО ИНФОРМАЦИОННОГО РАЗВИТИЯ И СВЯЗИ"
found_in_doc = False

print("📄 Шаг 1: Поиск в оригинальном документе...")
print()

for i, para in enumerate(doc.paragraphs[:20]):
    if target_text in para.text.upper():
        print(f"✅ Найдено в параграфе #{i}: '{para.text}'")
        found_in_doc = True
        break

if not found_in_doc:
    print(f"❌ '{target_text}' не найдено в первых 20 параграфах")
    print("\n🔍 Проверяем заголовки и другие части...")
    for section in doc.sections:
        header = section.header
        for para in header.paragraphs:
            if target_text in para.text.upper():
                print(f"✅ Найдено в ЗАГОЛОВКЕ: '{para.text}'")
                found_in_doc = True
                break

print()

# Шаг 2: Анализ через NLP Service
print("🤖 Шаг 2: Проверка обнаружения через NLP Service...")
print()

# Читаем весь текст документа
full_text = '\n'.join([p.text for p in doc.paragraphs])

# Вызываем NLP Service напрямую
nlp_url = "http://localhost:8006/analyze"

blocks = [{"block_id": "test_block", "content": full_text[:5000]}]  # Первые 5000 символов

response = requests.post(nlp_url, json={"blocks": blocks})

if response.status_code == 200:
    result = response.json()
    detections = result.get('detections', [])
    
    # Ищем наше значение
    ministry_detections = [d for d in detections if 'МИНИСТЕРСТВО' in d.get('text', '').upper() or 'МИНИСТЕРСТВО' in d.get('original_value', '').upper()]
    
    print(f"NLP Service нашел {len(detections)} детекций")
    print(f"Из них связанных с МИНИСТЕРСТВОМ: {len(ministry_detections)}")
    print()
    
    if ministry_detections:
        print("🎯 Детекции МИНИСТЕРСТВО:")
        for d in ministry_detections:
            print(f"   - Текст: '{d.get('text', d.get('original_value', 'N/A'))}'")
            print(f"     Категория: {d.get('category', 'N/A')}")
            print(f"     Метод: {d.get('method', 'N/A')}")
            print(f"     Confidence: {d.get('confidence', 'N/A')}")
            print()
else:
    print(f"❌ Ошибка NLP Service: {response.status_code}")

print()

# Шаг 3: Полная анонимизация
print("🔄 Шаг 3: Полная анонимизация документа...")
print()

with open(doc_path, 'rb') as f:
    files = {'file': ('test.docx', f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
    data = {
        'patterns_file': 'patterns/sensitive_patterns.xlsx',
        'generate_excel_report': 'true'
    }
    
    response = requests.post('http://localhost:8009/anonymize_full', files=files, data=data, timeout=120)

if response.status_code == 200:
    result = response.json()
    
    print(f"✅ Анонимизация выполнена")
    print(f"   Всего замен: {result.get('statistics', {}).get('total_replacements', 0)}")
    print()
    
    # Сохраняем анонимизированный документ
    doc_data = base64.b64decode(result['files_base64']['anonymized_document_base64'])
    anon_path = 'test_anonymized_diagnostic.docx'
    with open(anon_path, 'wb') as f:
        f.write(doc_data)
    
    # Читаем анонимизированный документ
    anon_doc = Document(anon_path)
    
    # Проверяем, заменился ли текст
    print("🔍 Шаг 4: Проверка замены в анонимизированном документе...")
    print()
    
    found_in_anon = False
    for i, para in enumerate(anon_doc.paragraphs[:20]):
        if target_text in para.text.upper():
            print(f"❌ ПРОБЛЕМА: Текст НЕ ЗАМЕНЕН в параграфе #{i}")
            print(f"   Оригинал: '{para.text}'")
            found_in_anon = True
            break
    
    # Проверяем заголовки
    for section in anon_doc.sections:
        header = section.header
        for para in header.paragraphs:
            if target_text in para.text.upper():
                print(f"❌ ПРОБЛЕМА: Текст НЕ ЗАМЕНЕН в ЗАГОЛОВКЕ")
                print(f"   Текст: '{para.text}'")
                found_in_anon = True
                break
    
    if not found_in_anon:
        print("✅ Текст был заменен (не найден в анонимизированном документе)")
    
    print()
    
    # Проверяем таблицу замен
    excel_data = base64.b64decode(result['files_base64']['excel_report_base64'])
    excel_path = 'test_replacements_diagnostic.xlsx'
    with open(excel_path, 'wb') as f:
        f.write(excel_data)
    
    import pandas as pd
    df = pd.read_excel(excel_path)
    
    print("📊 Шаг 5: Проверка таблицы замен...")
    print()
    
    ministry_in_table = df[df['Исходные данные'].str.contains('МИНИСТЕРСТВО', case=False, na=False)]
    
    if len(ministry_in_table) > 0:
        print(f"✅ Найдено в таблице замен: {len(ministry_in_table)} записей")
        for idx, row in ministry_in_table.iterrows():
            print(f"   - '{row['Исходные данные']}'")
            print(f"     UUID: {row['Замена (идентификатор)']}")
        print()
    else:
        print("❌ НЕ найдено в таблице замен!")
        print()
        print("🔍 Все записи в таблице:")
        for idx, row in df.head(10).iterrows():
            print(f"   {idx+1}. '{row['Исходные данные']}'")
        print()

else:
    print(f"❌ Ошибка анонимизации: {response.status_code}")

print()
print("=" * 80)
print("🎯 ДИАГНОЗ")
print("=" * 80)
