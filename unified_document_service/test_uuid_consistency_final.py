#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import requests
from docx import Document
import pandas as pd
import base64

doc_path = 'test_docs/test_01_1_4_SD33.docx'

print("=" * 80)
print("ТЕСТ КОНСИСТЕНТНОСТИ UUID: Документ vs Таблица замен")
print("=" * 80)

# Анонимизация
print("\n📤 Отправка документа на анонимизацию...")

with open(doc_path, 'rb') as f:
    files = {'file': ('test.docx', f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
    data = {
        'patterns_file': 'patterns/sensitive_patterns.xlsx',
        'generate_excel_report': 'true'
    }
    
    response = requests.post('http://localhost:8009/anonymize_full', files=files, data=data, timeout=120)

if response.status_code == 200:
    result = response.json()
    print(f"✅ Анонимизация выполнена")
    print(f"   Всего замен: {result.get('statistics', {}).get('total_replacements', 0)}")
    
    # Сохраняем файлы
    doc_data = base64.b64decode(result['files_base64']['anonymized_document_base64'])
    anon_path = 'test_uuid_consistency_anon.docx'
    with open(anon_path, 'wb') as f:
        f.write(doc_data)
    
    excel_data = base64.b64decode(result['files_base64']['excel_report_base64'])
    excel_path = 'test_uuid_consistency_report.xlsx'
    with open(excel_path, 'wb') as f:
        f.write(excel_data)
    
    # Читаем таблицу замен
    df = pd.read_excel(excel_path)
    
    print(f"\n📊 Таблица замен: {len(df)} записей")
    print("\nПервые 5 записей:")
    for idx, row in df.head(5).iterrows():
        print(f"{idx+1}. '{row['Исходные данные'][:50]}...' → '{row['Замена (идентификатор)']}'")
    
    # Читаем анонимизированный документ
    anon_doc = Document(anon_path)
    
    # Собираем все UUID из документа
    doc_uuids = set()
    for para in anon_doc.paragraphs:
        text = para.text
        # Ищем UUID (формат: xxxxxxxx-xxxx-xxxx-xxxx-xxxxxxxxxxxx)
        import re
        uuids_in_para = re.findall(r'[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}', text)
        doc_uuids.update(uuids_in_para)
    
    for table in anon_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                uuids_in_cell = re.findall(r'[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}', cell.text)
                doc_uuids.update(uuids_in_cell)
    
    print(f"\n📄 Документ: найдено {len(doc_uuids)} уникальных UUID")
    
    # Собираем UUID из таблицы
    table_uuids = set(df['Замена (идентификатор)'].astype(str).tolist())
    print(f"📊 Таблица: {len(table_uuids)} UUID")
    
    # Сравнение
    print(f"\n🔍 АНАЛИЗ КОНСИСТЕНТНОСТИ:")
    
    # UUID которые есть в документе, но НЕТ в таблице
    doc_only = doc_uuids - table_uuids
    if doc_only:
        print(f"\n❌ UUID только в документе (НЕТ в таблице): {len(doc_only)}")
        for uuid in list(doc_only)[:3]:
            print(f"   - {uuid}")
        if len(doc_only) > 3:
            print(f"   ... и еще {len(doc_only) - 3}")
    else:
        print(f"\n✅ Все UUID из документа есть в таблице")
    
    # UUID которые есть в таблице, но НЕТ в документе
    table_only = table_uuids - doc_uuids
    if table_only:
        print(f"\n❌ UUID только в таблице (НЕТ в документе): {len(table_only)}")
        print(f"\nПОЛНЫЙ СПИСОК:")
        for uuid in table_only:
            # Найдем оригинал для этого UUID
            orig = df[df['Замена (идентификатор)'] == uuid]['Исходные данные'].values
            if len(orig) > 0:
                print(f"   - {uuid}")
                print(f"     Оригинал: '{orig[0]}'")
    else:
        print(f"\n✅ Все UUID из таблицы есть в документе")
    
    # Общие UUID
    common = doc_uuids & table_uuids
    print(f"\n✅ Совпадающих UUID: {len(common)}")
    
    if len(doc_only) == 0 and len(table_only) == 0:
        print(f"\n🎉 ОТЛИЧНО! UUID полностью консистентны между документом и таблицей!")
    else:
        print(f"\n⚠️ ПРОБЛЕМА! UUID не совпадают между документом и таблицей")

else:
    print(f"❌ Ошибка: {response.status_code}")
    print(response.text)

print("\n" + "=" * 80)
