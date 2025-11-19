#!/usr/bin/env python3
"""
Тест позиционной логики для параграфов
"""

import sys
import os
sys.path.append(os.path.join(os.path.dirname(__file__), 'app'))

from docx import Document
from app.formatter_applier import FormatterApplier
import uuid

def test_paragraph_position_logic():
    """
    Тест позиционной логики для параграфов
    """
    print("🔍 ТЕСТ ПОЗИЦИОННОЙ ЛОГИКИ ДЛЯ ПАРАГРАФОВ")
    print("=" * 50)
    
    doc_path = r"C:\Projects\Anonymizer\unified_document_service\test_docs\test_01_1_4.docx"
    doc = Document(doc_path)
    
    # Найдем все параграфы содержащие "14 августа 2023"
    target_paragraphs = []
    
    for para_idx, para in enumerate(doc.paragraphs):
        para_text = para.text
        if "14 августа 2023" in para_text:
            target_paragraphs.append({
                'index': para_idx,
                'paragraph': para,
                'text': para_text,
                'context': para_text[:100] + "..." if len(para_text) > 100 else para_text
            })
            print(f"Параграф {para_idx}: '{para_text[:80]}{'...' if len(para_text) > 80 else ''}'")
    
    print(f"\nВсего найдено {len(target_paragraphs)} параграфов с датой")
    
    if len(target_paragraphs) == 0:
        print("❌ Не найдено параграфов с датой")
        return
    
    # Создаем тестовые замены с разными UUID для каждого параграфа
    test_replacements = []
    
    for i, para_info in enumerate(target_paragraphs):
        test_uuid = str(uuid.uuid4())
        replacement = {
            'uuid': test_uuid,
            'element': para_info['paragraph'],
            'original_value': '14 августа 2023',
            'category': 'date',
            'block_id': f'para_{para_info["index"]}_test_{i}',
            'position': {
                'start': 100 * i,  # Разные позиции для тестирования
                'end': 100 * i + 15
            }
        }
        test_replacements.append(replacement)
        print(f"🔧 Создана замена #{i+1} для параграфа {para_info['index']}: UUID={test_uuid[:8]}...")
    
    print("\n" + "=" * 50)
    print("🚀 ТЕСТИРУЕМ ПРИМЕНЕНИЕ ЗАМЕН В ПАРАГРАФАХ")
    print("=" * 50)
    
    formatter = FormatterApplier()
    successful_replacements = 0
    
    for i, replacement in enumerate(test_replacements):
        print(f"\n--- ЗАМЕНА #{i+1} ---")
        print(f"UUID: {replacement['uuid'][:8]}...")
        print(f"Элемент: параграф")
        print(f"Позиция: {replacement['position']['start']}")
        
        success = formatter._apply_single_replacement(replacement)
        if success:
            successful_replacements += 1
            print(f"Результат: ✅ Успех")
        else:
            print(f"Результат: ❌ Ошибка")
    
    print(f"\n📊 ИТОГ: {successful_replacements}/{len(test_replacements)} замен в параграфах успешно")
    
    # Анализируем результат
    print("\n" + "=" * 50)
    print("🔍 ФИНАЛЬНЫЙ АНАЛИЗ ПАРАГРАФОВ")
    print("=" * 50)
    
    uuid_found = {}
    date_count = 0
    
    for para_idx, para in enumerate(doc.paragraphs):
        para_text = para.text
        if '14 августа 2023' in para_text:
            date_count += 1
            print(f"Остался оригинальный текст в параграфе {para_idx}")
        
        # Поиск UUID в параграфах
        words = para_text.split()
        for word in words:
            if len(word) == 36 and word.count('-') == 4:  # Формат UUID
                if word in uuid_found:
                    uuid_found[word] += 1
                else:
                    uuid_found[word] = 1
                print(f"UUID в параграфе {para_idx}: {word}")
    
    print(f"\nИтог:")
    print(f"Остается оригинального текста: {date_count}")
    print(f"Найдено уникальных UUID: {len(uuid_found)}")
    for uuid_str, count in uuid_found.items():
        print(f"  {uuid_str}: {count} раз(а)")
    
    if len(uuid_found) == successful_replacements and all(count == 1 for count in uuid_found.values()):
        print("🎉 УСПЕХ: Каждый UUID в параграфах используется ровно один раз!")
    elif len(uuid_found) > 0:
        print("✅ ЧАСТИЧНЫЙ УСПЕХ: Замены в параграфах работают, но может потребоваться доработка позиций")
    else:
        print("❌ Проблема: замены в параграфах не работают")

def test_position_shift_handling():
    """
    Тест влияния замены на позиции последующих элементов
    """
    print("\n" + "=" * 70)
    print("🔧 ТЕСТ ВЛИЯНИЯ ЗАМЕНЫ НА ПОЗИЦИИ")
    print("=" * 70)
    
    # Демонстрируем как меняются позиции при замене
    original_text = "текст от 14 августа 2023"
    uuid_replacement = "2a3b4c5d-6e7f-8901-2345-67890abcdef1"
    
    print(f"Оригинальный текст: '{original_text}' (длина: {len(original_text)})")
    print(f"UUID для замены: '{uuid_replacement}' (длина: {len(uuid_replacement)})")
    print(f"Разница в длине: {len(uuid_replacement) - len('14 августа 2023')} символов")
    print("\nЭто означает, что все позиции ПОСЛЕ замены сдвинутся!")
    print("Поэтому важно обрабатывать замены в обратном порядке по позициям.")

if __name__ == "__main__":
    test_paragraph_position_logic()
    test_position_shift_handling()