#!/usr/bin/env python3
"""
Тест с реальными позициями UUID дубликации
"""

import sys
import os
sys.path.append(os.path.join(os.path.dirname(__file__), 'app'))

from docx import Document
from app.formatter_applier import FormatterApplier
import uuid

def test_real_positions():
    """
    Тест с реальными позициями из документа
    """
    print("🔍 ТЕСТ С РЕАЛЬНЫМИ ПОЗИЦИЯМИ")
    print("=" * 50)
    
    doc_path = r"C:\Projects\Anonymizer\unified_document_service\test_docs\test_01_1_4.docx"
    doc = Document(doc_path)
    table_2 = doc.tables[2]
    
    # Найдем реальные позиции текста "14 августа 2023" в table_2
    real_positions = []
    current_pos = 0
    
    for row_idx, row in enumerate(table_2.rows):
        for cell_idx, cell in enumerate(row.cells):
            cell_text = cell.text
            if "14 августа 2023" in cell_text:
                text_start_in_cell = cell_text.find("14 августа 2023")
                absolute_position = current_pos + text_start_in_cell
                real_positions.append({
                    'row': row_idx,
                    'cell': cell_idx, 
                    'position': absolute_position,
                    'cell_context': cell_text[max(0, text_start_in_cell-10):text_start_in_cell+25]
                })
                print(f"Найдено в ячейке [{row_idx}][{cell_idx}] на позиции {absolute_position}")
                print(f"  Контекст: '{cell_text[max(0, text_start_in_cell-10):text_start_in_cell+25]}'")
            current_pos += len(cell_text)
    
    print(f"\nВсего найдено {len(real_positions)} вхождений")
    
    # Создаем тестовые замены с реальными позициями
    test_replacements = []
    
    for i, pos_info in enumerate(real_positions):
        test_uuid = str(uuid.uuid4())
        replacement = {
            'uuid': test_uuid,
            'element': table_2,
            'original_value': '14 августа 2023',
            'category': 'date',
            'block_id': f'table_2_real_{i}',
            'position': {
                'start': pos_info['position'],
                'end': pos_info['position'] + 15  # длина "14 августа 2023"
            }
        }
        test_replacements.append(replacement)
        print(f"🔧 Создана замена #{i+1}: UUID={test_uuid[:8]}..., позиция={pos_info['position']}")
    
    # Тестируем применение
    print("\n" + "=" * 50)
    print("🚀 ТЕСТИРУЕМ С РЕАЛЬНЫМИ ПОЗИЦИЯМИ")
    print("=" * 50)
    
    formatter = FormatterApplier()
    successful_replacements = 0
    
    for i, replacement in enumerate(test_replacements):
        print(f"\n--- ЗАМЕНА #{i+1} ---")
        print(f"UUID: {replacement['uuid'][:8]}...")
        print(f"Позиция: {replacement['position']['start']}")
        
        success = formatter._apply_single_replacement(replacement)
        if success:
            successful_replacements += 1
            print(f"Результат: ✅ Успех")
        else:
            print(f"Результат: ❌ Ошибка")
    
    print(f"\n📊 ИТОГ: {successful_replacements}/{len(test_replacements)} замен успешно")
    
    # Анализируем результат
    print("\n" + "=" * 50)
    print("🔍 ФИНАЛЬНЫЙ АНАЛИЗ")
    print("=" * 50)
    
    uuid_found = {}
    date_count = 0
    
    for row_idx, row in enumerate(table_2.rows):
        for cell_idx, cell in enumerate(row.cells):
            cell_text = cell.text
            if '14 августа 2023' in cell_text:
                date_count += 1
                print(f"Остался оригинальный текст в ячейке [{row_idx}][{cell_idx}]")
            
            # Поиск UUID
            for para in cell.paragraphs:
                para_text = para.text
                words = para_text.split()
                for word in words:
                    if len(word) == 36 and word.count('-') == 4:  # Формат UUID
                        if word in uuid_found:
                            uuid_found[word] += 1
                        else:
                            uuid_found[word] = 1
                        print(f"UUID в ячейке [{row_idx}][{cell_idx}]: {word}")
    
    print(f"\nИтог:")
    print(f"Остается оригинального текста: {date_count}")
    print(f"Найдено уникальных UUID: {len(uuid_found)}")
    for uuid_str, count in uuid_found.items():
        print(f"  {uuid_str}: {count} раз(а)")
    
    if len(uuid_found) == successful_replacements and all(count == 1 for count in uuid_found.values()):
        print("🎉 УСПЕХ: Каждый UUID используется ровно один раз!")
    else:
        print("❌ Проблема все еще есть")

if __name__ == "__main__":
    test_real_positions()