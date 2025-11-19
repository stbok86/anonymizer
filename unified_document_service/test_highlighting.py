#!/usr/bin/env python3
"""
Тест выделения UUID желтым цветом
"""

import sys
import os
sys.path.append(os.path.join(os.path.dirname(__file__), 'app'))

from docx import Document
from app.formatter_applier import FormatterApplier
import uuid

def test_highlighting_functionality():
    """
    Тест функции выделения UUID желтым цветом
    """
    print("🎨 ТЕСТ ВЫДЕЛЕНИЯ UUID ЖЕЛТЫМ ЦВЕТОМ")
    print("=" * 50)
    
    # Загружаем документ
    doc_path = r"C:\Projects\Anonymizer\unified_document_service\test_docs\test_01_1_4.docx"
    doc = Document(doc_path)
    
    print(f"✅ Документ загружен: {doc_path}")
    
    # Найдем один параграф с датой для теста
    target_para = None
    for para_idx, para in enumerate(doc.paragraphs):
        if "14 августа 2023" in para.text:
            target_para = para
            print(f"📄 Найден параграф {para_idx}: '{para.text[:80]}...'")
            break
    
    if not target_para:
        print("❌ Параграф с датой не найден")
        return
    
    # Создаем тест с выделением
    print("\n🎨 ТЕСТ #1: Выделение включено (по умолчанию)")
    
    formatter_with_highlight = FormatterApplier(highlight_replacements=True)
    test_uuid = str(uuid.uuid4())
    
    replacement = {
        'uuid': test_uuid,
        'element': target_para,
        'original_value': '14 августа 2023',
        'category': 'date',
        'block_id': 'test_highlighting_on',
        'position': {'start': 100, 'end': 115}
    }
    
    print(f"🔧 Замена с выделением: '{replacement['original_value']}' → '{test_uuid[:8]}...'")
    success = formatter_with_highlight._apply_single_replacement(replacement)
    print(f"Результат: {'✅ Успех' if success else '❌ Ошибка'}")
    
    # Проверяем что UUID добавлен
    if test_uuid in target_para.text:
        print(f"✅ UUID найден в тексте параграфа")
        
        # Проверяем выделение
        highlight_found = False
        for run in target_para.runs:
            if test_uuid in run.text:
                try:
                    if hasattr(run.font, 'highlight_color') and run.font.highlight_color:
                        print(f"🎨 ✅ UUID выделен цветом: {run.font.highlight_color}")
                        highlight_found = True
                    else:
                        print(f"⚠️ UUID найден в run, но выделение не обнаружено")
                except Exception as e:
                    print(f"⚠️ Ошибка проверки выделения: {e}")
        
        if highlight_found:
            print("🎉 ВЫДЕЛЕНИЕ РАБОТАЕТ!")
        else:
            print("❌ Выделение НЕ применено")
    else:
        print(f"❌ UUID не найден в тексте")
    
    print("\n🎨 ТЕСТ #2: Выделение выключено")
    
    # Найдем другой параграф для второго теста
    target_para2 = None
    for para_idx, para in enumerate(doc.paragraphs):
        if "14 августа 2023" in para.text and para != target_para:
            target_para2 = para
            print(f"📄 Найден второй параграф {para_idx}: '{para.text[:80]}...'")
            break
    
    if target_para2:
        formatter_no_highlight = FormatterApplier(highlight_replacements=False)
        test_uuid2 = str(uuid.uuid4())
        
        replacement2 = {
            'uuid': test_uuid2,
            'element': target_para2,
            'original_value': '14 августа 2023',
            'category': 'date',
            'block_id': 'test_highlighting_off',
            'position': {'start': 200, 'end': 215}
        }
        
        print(f"🔧 Замена без выделения: '{replacement2['original_value']}' → '{test_uuid2[:8]}...'")
        success2 = formatter_no_highlight._apply_single_replacement(replacement2)
        print(f"Результат: {'✅ Успех' if success2 else '❌ Ошибка'}")
        
        # Проверяем отсутствие выделения
        if test_uuid2 in target_para2.text:
            print(f"✅ UUID найден в тексте параграфа")
            
            highlight_found2 = False
            for run in target_para2.runs:
                if test_uuid2 in run.text:
                    try:
                        if hasattr(run.font, 'highlight_color') and run.font.highlight_color:
                            print(f"⚠️ UUID выделен (не ожидалось): {run.font.highlight_color}")
                            highlight_found2 = True
                        else:
                            print(f"✅ UUID без выделения (как ожидалось)")
                    except Exception as e:
                        print(f"ℹ️ Проверка выделения: {e}")
            
            if not highlight_found2:
                print("✅ Выделение корректно выключено")
            else:
                print("❌ Выделение применено, хотя должно быть выключено")
    
    print("\n📋 ИТОГ ТЕСТА ВЫДЕЛЕНИЯ:")
    print("- По умолчанию FormatterApplier включает выделение UUID желтым цветом")
    print("- Выделение можно отключить параметром highlight_replacements=False") 
    print("- Все новые UUID в документе будут выделены желтым фоном")

def save_test_document():
    """
    Сохраняем тестовый документ с выделениями
    """
    print("\n💾 СОХРАНЕНИЕ ТЕСТОВОГО ДОКУМЕНТА")
    print("=" * 50)
    
    doc_path = r"C:\Projects\Anonymizer\unified_document_service\test_docs\test_01_1_4.docx"
    doc = Document(doc_path)
    
    # Применяем несколько замен с выделением
    formatter = FormatterApplier(highlight_replacements=True)
    
    replacements_applied = 0
    for para_idx, para in enumerate(doc.paragraphs):
        if "14 августа 2023" in para.text and replacements_applied < 2:
            test_uuid = str(uuid.uuid4())
            replacement = {
                'uuid': test_uuid,
                'element': para,
                'original_value': '14 августа 2023',
                'category': 'date',
                'block_id': f'save_test_{para_idx}',
                'position': {'start': 100 + replacements_applied * 100, 'end': 115 + replacements_applied * 100}
            }
            
            success = formatter._apply_single_replacement(replacement)
            if success:
                replacements_applied += 1
                print(f"✅ Замена {replacements_applied}: параграф {para_idx}")
    
    # Сохраняем результат
    output_path = r"C:\Projects\Anonymizer\unified_document_service\test_highlighting_result.docx"
    doc.save(output_path)
    print(f"💾 Документ с выделениями сохранен: {output_path}")
    print("🎨 Откройте документ в Word чтобы увидеть желтые выделения UUID!")

if __name__ == "__main__":
    test_highlighting_functionality()
    save_test_document()