#!/usr/bin/env python3
"""
Диагностический скрипт для анализа структуры DOCX документа
и выявления проблем с извлечением колонтитулов
"""

from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
import sys
import os

def analyze_docx_structure(file_path):
    """Детальный анализ структуры DOCX файла"""
    print(f"=== Анализ файла: {file_path} ===")
    
    try:
        doc = Document(file_path)
        print(f"Количество секций: {len(doc.sections)}")
        
        # Анализ основного содержимого
        print(f"\n--- ОСНОВНОЕ СОДЕРЖИМОЕ ---")
        print(f"Количество параграфов: {len(doc.paragraphs)}")
        print(f"Количество таблиц: {len(doc.tables)}")
        
        # Анализ headers для каждой секции
        print(f"\n--- АНАЛИЗ HEADERS ---")
        for s_idx, section in enumerate(doc.sections):
            print(f"\nСекция {s_idx}:")
            header = section.header
            print(f"  Header существует: {header is not None}")
            
            if header:
                # Проверяем paragraphs через header.paragraphs
                print(f"  Параграфов в header (через .paragraphs): {len(header.paragraphs)}")
                for p_idx, para in enumerate(header.paragraphs):
                    text = para.text.strip()
                    print(f"    Параграф {p_idx}: '{text}' (len={len(text)})")
                
                # Проверяем элементы через _element
                print(f"  Элементов в header._element: {len(list(header._element))}")
                for el_idx, el in enumerate(header._element):
                    tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
                    print(f"    Элемент {el_idx}: {tag}")
                    
                    if el.tag.endswith('}p'):
                        try:
                            paragraph = Paragraph(el, header)
                            text = paragraph.text.strip()
                            print(f"      -> Текст параграфа: '{text}'")
                        except Exception as e:
                            print(f"      -> Ошибка создания Paragraph: {e}")
        
        # Анализ footers для каждой секции  
        print(f"\n--- АНАЛИЗ FOOTERS ---")
        for s_idx, section in enumerate(doc.sections):
            print(f"\nСекция {s_idx}:")
            footer = section.footer
            print(f"  Footer существует: {footer is not None}")
            
            if footer:
                print(f"  Параграфов в footer: {len(footer.paragraphs)}")
                for p_idx, para in enumerate(footer.paragraphs):
                    text = para.text.strip()
                    print(f"    Параграф {p_idx}: '{text}' (len={len(text)})")
        
        # Поиск специфического текста "ЕИСУФХД"
        print(f"\n--- ПОИСК СПЕЦИФИЧЕСКОГО ТЕКСТА ---")
        target_text = "ЕИСУФХД"
        
        # В основном содержимом
        found_in_main = False
        for para in doc.paragraphs:
            if target_text in para.text:
                print(f"  Найден в основном содержимом: '{para.text.strip()}'")
                found_in_main = True
        
        # В headers
        found_in_headers = False
        for s_idx, section in enumerate(doc.sections):
            header = section.header
            if header:
                for para in header.paragraphs:
                    if target_text in para.text:
                        print(f"  Найден в header секции {s_idx}: '{para.text.strip()}'")
                        found_in_headers = True
        
        # В footers
        found_in_footers = False
        for s_idx, section in enumerate(doc.sections):
            footer = section.footer
            if footer:
                for para in footer.paragraphs:
                    if target_text in para.text:
                        print(f"  Найден в footer секции {s_idx}: '{para.text.strip()}'")
                        found_in_footers = True
        
        if not (found_in_main or found_in_headers or found_in_footers):
            print(f"  Текст '{target_text}' не найден в стандартных местах!")
            print(f"  Проверяем XML структуру...")
            
            # Дополнительный поиск в XML
            xml_content = str(doc._element.xml)
            if target_text in xml_content:
                print(f"  Текст найден в XML структуре документа!")
            else:
                print(f"  Текст НЕ найден даже в XML!")
                
    except Exception as e:
        print(f"Ошибка анализа: {e}")

if __name__ == "__main__":
    # Тестируем на известном файле
    test_file = r"C:\Projects\Anonymizer\unified_document_service\test_docs\test_01_1_4.docx"
    
    if os.path.exists(test_file):
        analyze_docx_structure(test_file)
    else:
        print(f"Файл не найден: {test_file}")
        print("Доступные файлы в test_docs:")
        test_dir = r"C:\Projects\Anonymizer\unified_document_service\test_docs"
        if os.path.exists(test_dir):
            for f in os.listdir(test_dir):
                if f.endswith('.docx'):
                    print(f"  - {f}")