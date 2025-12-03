#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Тест детекции организации в реальном документе"""

import sys
import os
sys.path.insert(0, os.path.dirname(__file__))

from docx import Document
import requests

# Путь к документу
doc_path = r"C:\Projects\Anonymizer\unified_document_service\test_docs\test_01_1_4_SD22.docx"

print("="*100)
print("ТЕСТ ДЕТЕКЦИИ ОРГАНИЗАЦИИ В РЕАЛЬНОМ ДОКУМЕНТЕ")
print("="*100)

# Читаем документ
doc = Document(doc_path)

# Ищем организацию в тексте
target_org = "Общество с ограниченной ответственностью «КАМА Технологии»"
print(f"\nИщем организацию: '{target_org}'")

# Собираем весь текст
all_paragraphs = []
for para in doc.paragraphs:
    if para.text.strip():
        all_paragraphs.append(para.text)

# Также проверяем таблицы
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if cell.text.strip():
                all_paragraphs.append(cell.text)

# Ищем организацию
found_in = []
for i, text in enumerate(all_paragraphs):
    if "КАМА Технологии" in text or "ограниченной ответственностью" in text:
        found_in.append((i, text))
        
if found_in:
    print(f"\nОрганизация найдена в {len(found_in)} местах:")
    for idx, (i, text) in enumerate(found_in[:3], 1):
        print(f"\n  {idx}. Параграф/ячейка #{i}:")
        print(f"     Текст: '{text[:250]}...'")
        
        # Тестируем детекцию через NLP Service
        print(f"\n     Тестируем детекцию через NLP Service...")
        
        response = requests.post(
            "http://localhost:8006/analyze",
            json={"blocks": [{"block_id": f"test_{i}", "content": text}]},
            timeout=10
        )
        
        if response.status_code == 200:
            result = response.json()
            detections = result.get('detections', [])
            
            org_detections = [d for d in detections if d.get('category') == 'organization']
            
            print(f"     Статус: {response.status_code}")
            print(f"     Всего детекций: {len(detections)}")
            print(f"     Детекций organization: {len(org_detections)}")
            
            if org_detections:
                for j, det in enumerate(org_detections, 1):
                    print(f"\n       Детекция #{j}:")
                    print(f"         - Текст: '{det.get('original_value', '')}'")
                    print(f"         - Уверенность: {det.get('confidence', 0):.2f}")
                    print(f"         - Метод: {det.get('method', 'N/A')}")
                    print(f"         - Позиция: {det.get('position', {}).get('start', 0)}-{det.get('position', {}).get('end', 0)}")
            else:
                print(f"       ОРГАНИЗАЦИЯ НЕ НАЙДЕНА!")
                
                # Показываем что найдено
                if detections:
                    print(f"\n       Найдены другие категории:")
                    for det in detections[:5]:
                        print(f"         - {det.get('category')}: '{det.get('original_value', '')[:80]}'")
        else:
            print(f"     Ошибка: {response.status_code}")
else:
    print("\nОрганизация НЕ НАЙДЕНА в документе!")

# Дополнительно - тестируем чистое название
print("\n" + "="*100)
print("ПРЯМОЙ ТЕСТ ОРГАНИЗАЦИИ")
print("="*100)

test_orgs = [
    "Общество с ограниченной ответственностью «КАМА Технологии»",
    "ООО «КАМА Технологии»",
    "«КАМА Технологии»",
    "Общество с ограниченной ответственностью"
]

for org in test_orgs:
    print(f"\nТест организации: '{org}'")
    
    response = requests.post(
        "http://localhost:8006/analyze",
        json={"blocks": [{"block_id": "test", "content": org}]},
        timeout=10
    )
    
    if response.status_code == 200:
        result = response.json()
        detections = result.get('detections', [])
        org_detections = [d for d in detections if d.get('category') == 'organization']
        
        print(f"  Детекций organization: {len(org_detections)}")
        
        if org_detections:
            for det in org_detections:
                print(f"    - '{det.get('original_value', '')}' (conf: {det.get('confidence', 0):.2f}, метод: {det.get('method')})")
        else:
            print(f"    НЕ НАЙДЕНО! Всего детекций: {len(detections)}")
            if detections:
                for det in detections[:2]:
                    print(f"    - {det.get('category')}: '{det.get('original_value', '')}'")
