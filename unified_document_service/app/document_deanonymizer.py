#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Document Deanonymizer - Модуль для деанонимизации документов
Выполняет обратную замену UUID на оригинальные чувствительные данные
"""

import pandas as pd
from docx import Document
import re
import tempfile
import os
from datetime import datetime
from typing import Dict, List, Tuple, Any, Optional
import uuid
import logging

# Настройка логирования
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentDeanonymizer:
    """
    Класс для деанонимизации документов
    Заменяет UUID обратно на оригинальные чувствительные данные
    """
    
    def __init__(self):
        self.replacement_mapping = {}
        self.statistics = {
            'total_uuids_found': 0,
            'total_replacements': 0,
            'successful_replacements': 0,
            'failed_replacements': 0,
            'replacement_details': []
        }
    
    def deanonymize_document(self, document_path: str, table_path: str) -> Dict[str, Any]:
        """
        Основная функция деанонимизации документа
        
        Args:
            document_path: Путь к анонимизированному DOCX файлу
            table_path: Путь к таблице соответствий (Excel/CSV)
            
        Returns:
            Словарь с результатами деанонимизации
        """
        
        try:
            logger.info(f"Начало деанонимизации документа: {document_path}")
            logger.info(f"[INFO] Таблица замен: {table_path}")
            
            # Загружаем таблицу соответствий
            if not self._load_replacement_table(table_path):
                return {
                    'success': False,
                    'error': 'Не удалось загрузить таблицу соответствий',
                    'statistics': self.statistics
                }
            
            # Загружаем документ
            try:
                document = Document(document_path)
            except Exception as e:
                return {
                    'success': False,
                    'error': f'Не удалось загрузить документ: {str(e)}',
                    'statistics': self.statistics
                }
            
            # Выполняем деанонимизацию
            self._process_document(document)
            
            # Сохраняем деанонимизированный документ
            output_path = tempfile.mktemp(suffix='_deanonymized.docx')
            document.save(output_path)
            
            # Создаем отчет
            report_path = self._create_deanonymization_report()
            
            result = {
                'success': True,
                'output_path': output_path,
                'report_path': report_path,
                'statistics': self.statistics,
                'message': f'Успешно деанонимизировано: {self.statistics["successful_replacements"]} из {self.statistics["total_uuids_found"]} UUID'
            }
            
            logger.info(f"Деанонимизация завершена успешно")
            logger.info(f"[STATS] Статистика: {self.statistics}")
            
            return result
            
        except Exception as e:
            logger.error(f"[ERROR] Ошибка при деанонимизации: {str(e)}")
            return {
                'success': False,
                'error': f'Внутренняя ошибка: {str(e)}',
                'statistics': self.statistics
            }
    
    def _load_replacement_table(self, table_path: str) -> bool:
        """Загружает таблицу соответствий UUID ↔ оригинальные данные"""
        
        try:
            # Определяем формат файла и загружаем
            if table_path.endswith('.xlsx'):
                df = pd.read_excel(table_path)
            elif table_path.endswith('.csv'):
                df = pd.read_csv(table_path)
            else:
                logger.error(f"[ERROR] Неподдерживаемый формат таблицы: {table_path}")
                return False
            
            logger.info(f"Загружена таблица размером: {len(df)} строк")
            logger.info(f"Колонки: {list(df.columns)}")
            
            # Ищем нужные колонки (различные варианты названий)
            uuid_column = self._find_column(df, [
                'uuid', 'UUID', 'uuid_value', 'replacement_uuid',
                'замена', 'Замена', 'замена (идентификатор)', 'Замена (идентификатор)',
                'идентификатор', 'Идентификатор', 'replacement', 'Replacement'
            ])
            original_column = self._find_column(df, [
                'original_value', 'original', 'value', 'sensitive_data',
                'исходные данные', 'Исходные данные', 'оригинал', 'Оригинал',
                'данные', 'Данные', 'текст', 'Текст'
            ])
            
            if not uuid_column:
                logger.error("[ERROR] Не найдена колонка с UUID")
                return False
            
            if not original_column:
                logger.error("[ERROR] Не найдена колонка с оригинальными данными")
                return False
            
            logger.info(f"[SUCCESS] Найдены колонки: UUID='{uuid_column}', Original='{original_column}'")
            
            # Строим маппинг UUID -> оригинальное значение
            for _, row in df.iterrows():
                uuid_val = str(row[uuid_column]).strip()
                original_val = str(row[original_column]).strip()
                
                if uuid_val and original_val and uuid_val != 'nan' and original_val != 'nan':
                    self.replacement_mapping[uuid_val] = original_val
            
            logger.info(f"[INFO] Загружено {len(self.replacement_mapping)} соответствий для деанонимизации")
            
            # Показываем примеры для проверки
            if len(self.replacement_mapping) > 0:
                sample_items = list(self.replacement_mapping.items())[:3]
                for uuid_val, original_val in sample_items:
                    logger.info(f"   [SAMPLE] {uuid_val[:8]}... -> '{original_val}'")
            
            return len(self.replacement_mapping) > 0
            
        except Exception as e:
            logger.error(f"[ERROR] Ошибка при загрузке таблицы: {str(e)}")
            return False
    
    def _find_column(self, df: pd.DataFrame, possible_names: List[str]) -> Optional[str]:
        """Находит колонку по списку возможных названий"""
        
        for col_name in df.columns:
            if col_name.lower() in [name.lower() for name in possible_names]:
                return col_name
        return None
    
    def _process_document(self, document: Document):
        """Обрабатывает документ и заменяет UUID на оригинальные данные"""
        
        logger.info("Начинаем обработку параграфов документа")
        
        # Обрабатываем основной текст документа
        for paragraph in document.paragraphs:
            self._process_paragraph(paragraph)
        
        # Обрабатываем таблицы
        for table in document.tables:
            self._process_table(table)
        
        # Обрабатываем колонтитулы
        for section in document.sections:
            # Header
            if section.header:
                for paragraph in section.header.paragraphs:
                    self._process_paragraph(paragraph)
            
            # Footer
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    self._process_paragraph(paragraph)
        
        logger.info(f"[SUCCESS] Обработка документа завершена")
    
    def _process_paragraph(self, paragraph):
        """Обрабатывает параграф и заменяет UUID в тексте"""
        
        if not paragraph.text.strip():
            return
        
        # Ищем UUID в тексте параграфа
        uuids_in_text = self._find_uuids_in_text(paragraph.text)
        
        if not uuids_in_text:
            return
        
        logger.debug(f"[DEBUG] Найдено UUID в параграфе: {len(uuids_in_text)}")
        
        # Заменяем UUID на оригинальные данные, сохраняя форматирование
        for run in paragraph.runs:
            if run.text.strip():
                original_text = run.text
                modified_text = self._replace_uuids_in_text(original_text)
                
                if modified_text != original_text:
                    run.text = modified_text
                    logger.debug(f"[DEBUG] Замена в run: '{original_text[:50]}...' -> '{modified_text[:50]}...'")
    
    def _process_table(self, table):
        """Обрабатывает таблицу и заменяет UUID в ячейках"""
        
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self._process_paragraph(paragraph)
    
    def _find_uuids_in_text(self, text: str) -> List[str]:
        """Находит все UUID в тексте"""
        
        # Паттерн для UUID версии 4 (8-4-4-4-12 символов)
        uuid_pattern = r'\b[0-9a-fA-F]{8}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-[0-9a-fA-F]{12}\b'
        
        found_uuids = re.findall(uuid_pattern, text)
        self.statistics['total_uuids_found'] += len(found_uuids)
        
        return found_uuids
    
    def _replace_uuids_in_text(self, text: str) -> str:
        """Заменяет UUID в тексте на оригинальные данные"""
        
        # Паттерн для UUID
        uuid_pattern = r'\b[0-9a-fA-F]{8}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-[0-9a-fA-F]{12}\b'
        
        def replace_uuid(match):
            uuid_str = match.group(0)
            self.statistics['total_replacements'] += 1
            
            # Ищем соответствие в таблице замен
            if uuid_str in self.replacement_mapping:
                original_value = self.replacement_mapping[uuid_str]
                self.statistics['successful_replacements'] += 1
                
                # Записываем детали замены
                self.statistics['replacement_details'].append({
                    'uuid': uuid_str,
                    'original_value': original_value,
                    'status': 'success'
                })
                
                logger.debug(f"[SUCCESS] Замена: {uuid_str} -> '{original_value}'")
                return original_value
            else:
                self.statistics['failed_replacements'] += 1
                
                # Записываем неудачную замену
                self.statistics['replacement_details'].append({
                    'uuid': uuid_str,
                    'original_value': None,
                    'status': 'failed',
                    'reason': 'UUID не найден в таблице замен'
                })
                
                logger.warning(f"[WARNING] UUID не найден в таблице: {uuid_str}")
                return uuid_str  # Оставляем без изменений
        
        # Выполняем замену
        result = re.sub(uuid_pattern, replace_uuid, text)
        return result
    
    def _create_deanonymization_report(self) -> str:
        """Создает отчет о процессе деанонимизации"""
        
        try:
            # Создаем DataFrame с деталями
            details_df = pd.DataFrame(self.statistics['replacement_details'])
            
            # Создаем общую статистику
            summary_data = {
                'Метрика': [
                    'Всего UUID найдено в документе',
                    'Всего попыток замен', 
                    'Успешных замен',
                    'Неудачных замен',
                    'Процент успешности'
                ],
                'Значение': [
                    self.statistics['total_uuids_found'],
                    self.statistics['total_replacements'],
                    self.statistics['successful_replacements'],
                    self.statistics['failed_replacements'],
                    f"{round((self.statistics['successful_replacements'] / max(1, self.statistics['total_replacements'])) * 100, 1)}%"
                ]
            }
            summary_df = pd.DataFrame(summary_data)
            
            # Сохраняем в Excel
            report_path = tempfile.mktemp(suffix='_deanonymization_report.xlsx')
            
            with pd.ExcelWriter(report_path, engine='openpyxl') as writer:
                summary_df.to_excel(writer, sheet_name='Статистика', index=False)
                if not details_df.empty:
                    details_df.to_excel(writer, sheet_name='Детали замен', index=False)
            
            logger.info(f"[INFO] Создан отчет о деанонимизации: {report_path}")
            return report_path
            
        except Exception as e:
            logger.error(f"[ERROR] Ошибка при создании отчета: {str(e)}")
            return None
    
    def validate_uuid_format(self, uuid_str: str) -> bool:
        """Проверяет, является ли строка валидным UUID"""
        
        try:
            uuid.UUID(uuid_str)
            return True
        except ValueError:
            return False


def main():
    """Функция для тестирования модуля"""
    
    # Пример использования
    deanonymizer = DocumentDeanonymizer()
    
    # Тестовые пути (заменить на реальные)
    doc_path = "test_anonymized.docx"
    table_path = "test_replacements.xlsx"
    
    if os.path.exists(doc_path) and os.path.exists(table_path):
        result = deanonymizer.deanonymize_document(doc_path, table_path)
        print(f"Результат: {result}")
    else:
        print("Тестовые файлы не найдены")


if __name__ == "__main__":
    main()