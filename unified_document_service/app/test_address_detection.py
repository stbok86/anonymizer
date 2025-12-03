#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Тест детекции адреса в реальном документе"""

import sys
import os
sys.path.insert(0, os.path.dirname(__file__))

from docx import Document
import requests

# Путь к документу
doc_path = r"C:\Projects\Anonymizer\unified_document_service\test_docs\test_01_1_4_SD22.docx"

print("="*100)
print("ТЕСТ ДЕТЕКЦИИ АДРЕСА В РЕАЛЬНОМ ДОКУМЕНТЕ")
print("="*100)

# Читаем документ
doc = Document(doc_path)

# Ищем адрес в тексте
target_address = "614000, Пермский край, г. Пермь, ул. Ленина, д. 66"
print(f"\nИщем адрес: '{target_address}'")

# Собираем весь текст
all_paragraphs = []
for para in doc.paragraphs:
    if para.text.strip():
        all_paragraphs.append(para.text)

# Также проверяем таблицы
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if cell.text.strip():
                all_paragraphs.append(cell.text)

# Ищем адрес
found_in = []
for i, text in enumerate(all_paragraphs):
    if target_address in text or "Пермский край" in text or "ул. Ленина" in text:
        found_in.append((i, text))
        
if found_in:
    print(f"\nАдрес найден в {len(found_in)} местах:")
    for idx, (i, text) in enumerate(found_in[:3], 1):
        print(f"\n  {idx}. Параграф/ячейка #{i}:")
        print(f"     Текст: '{text[:200]}...'")
        
        # Тестируем детекцию через NLP Service
        print(f"\n     Тестируем детекцию через NLP Service...")
        
        response = requests.post(
            "http://localhost:8006/analyze",
            json={"blocks": [{"block_id": f"test_{i}", "content": text}]},
            timeout=10
        )
        
        if response.status_code == 200:
            result = response.json()
            detections = result.get('detections', [])
            
            address_detections = [d for d in detections if d.get('category') == 'address']
            
            print(f"     Статус: {response.status_code}")
            print(f"     Всего детекций: {len(detections)}")
            print(f"     Детекций address: {len(address_detections)}")
            
            if address_detections:
                for j, det in enumerate(address_detections, 1):
                    print(f"\n       Детекция #{j}:")
                    print(f"         - Текст: '{det.get('original_value', '')}'")
                    print(f"         - Уверенность: {det.get('confidence', 0):.2f}")
                    print(f"         - Метод: {det.get('method', 'N/A')}")
                    print(f"         - Позиция: {det.get('position', {}).get('start', 0)}-{det.get('position', {}).get('end', 0)}")
            else:
                print(f"       АДРЕС НЕ НАЙДЕН!")
                
                # Показываем что найдено
                if detections:
                    print(f"\n       Найдены другие категории:")
                    for det in detections[:3]:
                        print(f"         - {det.get('category')}: '{det.get('original_value', '')[:50]}'")
        else:
            print(f"     Ошибка: {response.status_code}")
else:
    print("\nАдрес НЕ НАЙДЕН в документе!")

# Дополнительно - тестируем чистый адрес
print("\n" + "="*100)
print("ПРЯМОЙ ТЕСТ АДРЕСА")
print("="*100)

test_addresses = [
    "614000, Пермский край, г. Пермь, ул. Ленина, д. 66",
    "г. Пермь, ул. Ленина, д. 66",
    "ул. Ленина, д. 66",
    "Пермский край, г. Пермь"
]

for addr in test_addresses:
    print(f"\nТест адреса: '{addr}'")
    
    response = requests.post(
        "http://localhost:8006/analyze",
        json={"blocks": [{"block_id": "test", "content": addr}]},
        timeout=10
    )
    
    if response.status_code == 200:
        result = response.json()
        detections = result.get('detections', [])
        address_detections = [d for d in detections if d.get('category') == 'address']
        
        print(f"  Детекций address: {len(address_detections)}")
        
        if address_detections:
            for det in address_detections:
                print(f"    - '{det.get('original_value', '')}' (conf: {det.get('confidence', 0):.2f})")
        else:
            print(f"    НЕ НАЙДЕНО! Всего детекций: {len(detections)}")
            if detections:
                for det in detections[:2]:
                    print(f"    - {det.get('category')}: '{det.get('original_value', '')}'")
