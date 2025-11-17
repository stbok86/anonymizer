
from fastapi import FastAPI, UploadFile, File, Form
from fastapi.responses import JSONResponse, FileResponse
import tempfile
import shutil
import os
import base64
import requests

try:
    from docx import Document
    from block_builder import BlockBuilder
    from rule_adapter import RuleEngineAdapter
    from formatter_applier import FormatterApplier
    from full_anonymizer import FullAnonymizer
except ImportError as e:
    print(f"Warning: Could not import required modules: {e}")
    # В production эти модули должны быть доступны

app = FastAPI()

# URL для NLP Service
NLP_SERVICE_URL = "http://localhost:8006"

# Middleware для логирования всех запросов
@app.middleware("http")
async def log_requests(request, call_next):
    print(f"🌐 [REQUEST] {request.method} {request.url.path} от {request.client.host}")
    if request.method == "POST":
        print(f"🌐 [REQUEST] Content-Type: {request.headers.get('content-type', 'NOT_SET')}")
    response = await call_next(request)
    print(f"🌐 [RESPONSE] {response.status_code}")
    return response

@app.get("/health")
def health():
    return {"status": "ok", "service": "unified_document_service"}

@app.get("/healthz")
def healthz():
    return {"status": "ok"}

@app.get("/readyz")
def readyz():
    return {"status": "ready"}

@app.post("/parse_docx_blocks")
async def parse_docx_blocks(file: UploadFile = File(...)):
    # Сохраняем загруженный файл во временный файл
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        shutil.copyfileobj(file.file, tmp)
        tmp_path = tmp.name
    try:
        doc = Document(tmp_path)
        builder = BlockBuilder()
        blocks = builder.build_blocks(doc)
        # Возвращаем только метаинформацию (без element)
        blocks_out = [
            {k: v for k, v in b.items() if k != "element"}
            for b in blocks
        ]
        return JSONResponse(content={"blocks": blocks_out})
    finally:
        os.remove(tmp_path)

@app.post("/analyze_document")
async def analyze_document(file: UploadFile = File(...), patterns_file: str = "patterns/sensitive_patterns.xlsx"):
    """
    Анализ документа: парсинг → поиск чувствительных данных (без анонимизации)
    """
    print(f"🔍 [DEBUG] analyze_document вызван с patterns_file: {patterns_file}")
    
    # Сохраняем загруженный файл
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        shutil.copyfileobj(file.file, tmp)
        tmp_path = tmp.name
    
    try:
        print(f"🔍 [DEBUG] Загружен файл: {file.filename}")
        
        # ЭТАП 1: Загрузка документа
        doc = Document(tmp_path)
        
        # ЭТАП 2: Извлечение блоков
        builder = BlockBuilder()
        blocks = builder.build_blocks(doc)
        print(f"🔍 [DEBUG] Извлечено блоков: {len(blocks)}")
        
        # ЭТАП 3: Поиск чувствительных данных
        print(f"🔍 [DEBUG] Начинаем поиск чувствительных данных с файлом паттернов: {patterns_file}")
        rule_engine = RuleEngineAdapter(patterns_file)
        processed_blocks = rule_engine.apply_rules_to_blocks(blocks)
        
        # Собираем найденные элементы Rule Engine
        rule_engine_items = []
        for block in processed_blocks:
            if 'sensitive_patterns' in block:
                for pattern in block['sensitive_patterns']:
                    found_item = {
                        'block_id': block['block_id'],
                        'category': pattern['category'],
                        'original_value': pattern['original_value'],
                        'uuid': pattern['uuid'],
                        'position': pattern['position'],
                        'confidence': pattern.get('confidence', 1.0),
                        'method': 'regex',  # Rule Engine всегда использует regex
                        'source': 'Rule Engine',
                        'block_text': block.get('text', block.get('content', 'Контекст недоступен'))
                    }
                    rule_engine_items.append(found_item)
        
        print(f"🔍 [DEBUG] Rule Engine нашел: {len(rule_engine_items)} элементов")
        
        # ЭТАП 4: NLP анализ тех же блоков
        nlp_items = []
        try:
            print(f"🧠 [DEBUG] Начинаем NLP анализ {len(blocks)} блоков...")
            
            # Подготавливаем блоки для NLP Service
            text_blocks = [
                {
                    "content": block.get('text', block.get('content', '')),
                    "block_id": block.get('block_id', f'block_{i}'),
                    "block_type": block.get('type', 'text')
                }
                for i, block in enumerate(blocks)
                if block.get('text') or block.get('content')
            ]
            
            if text_blocks:
                nlp_data = {
                    "blocks": text_blocks,
                    "options": {"confidence_threshold": 0.6}
                }
                
                print(f"🧠 [DEBUG] Отправляем в NLP Service {len(text_blocks)} блоков")
                
                nlp_response = requests.post(
                    f"{NLP_SERVICE_URL}/analyze",
                    json=nlp_data,
                    timeout=60
                )
                
                print(f"🧠 [DEBUG] NLP Service ответил: {nlp_response.status_code}")
                
                if nlp_response.status_code == 200:
                    nlp_result = nlp_response.json()
                    
                    print(f"🧠 [DEBUG] NLP Response структура:")
                    print(f"  success: {nlp_result.get('success', 'НЕ НАЙДЕН')}")
                    print(f"  detections: {nlp_result.get('detections', 'НЕ НАЙДЕНО')}")
                    print(f"  Полный ответ: {nlp_result}")
                    
                    if nlp_result.get('success', False) and 'detections' in nlp_result:
                        nlp_counter = len(rule_engine_items)
                        for detection in nlp_result['detections']:
                            nlp_item = {
                                'block_id': detection.get('block_id', f'nlp_block_{nlp_counter}'),
                                'category': detection.get('category', 'NLP Detection'),
                                'original_value': detection.get('original_value', ''),
                                'uuid': detection.get('uuid', f"NLP_{nlp_counter}_{detection.get('category', 'unknown')}"),
                                'position': detection.get('position', {}),
                                'confidence': detection.get('confidence', 0.8),
                                'method': detection.get('method', 'nlp_unknown'),  # Добавляем поле method
                                'spacy_label': detection.get('spacy_label', ''),   # Добавляем spacy_label
                                'source': 'NLP Service',
                                'block_text': detection.get('context', detection.get('original_value', ''))
                            }
                            nlp_items.append(nlp_item)
                            nlp_counter += 1
                        
                        print(f"🧠 [DEBUG] NLP Service нашел: {len(nlp_result['detections'])} элементов")
                    else:
                        print(f"🧠 [DEBUG] NLP Service не нашел чувствительных данных или вернул success=false")
                else:
                    print(f"🧠 [DEBUG] NLP Service error: {nlp_response.status_code} - {nlp_response.text}")
            else:
                print(f"🧠 [DEBUG] Нет текстовых блоков для NLP анализа")
                
        except Exception as e:
            print(f"🧠 [DEBUG] Ошибка NLP анализа: {str(e)}")
        
        # Объединяем результаты БЕЗ дедупликации - каждый сервис отвечает за свои данные
        all_found_items = rule_engine_items + nlp_items
        
        print(f"📊 [ANALYZE] Объединение результатов: Rule Engine={len(rule_engine_items)}, NLP Service={len(nlp_items)}, Итого={len(all_found_items)}")
        print(f"💡 [INFO] Дедупликация отключена - каждый сервис обрабатывает свои типы данных")
        
        return JSONResponse(content={
            "status": "success",
            "found_items": all_found_items,
            "total_items": len(all_found_items),
            "rule_engine_items": len(rule_engine_items),
            "nlp_items": len(nlp_items),
            "duplicates_removed": 0,  # Дедупликация отключена
            "blocks_processed": len(blocks),
            "filename": file.filename
        })
        
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"status": "error", "message": str(e)}
        )
    finally:
        os.remove(tmp_path)

@app.post("/anonymize_document")
async def anonymize_document(file: UploadFile = File(...), patterns_file: str = "patterns/sensitive_patterns.xlsx"):
    """
    Полная анонимизация документа: парсинг → поиск → замена с сохранением форматирования
    """
    # Сохраняем загруженный файл
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        shutil.copyfileobj(file.file, tmp)
        tmp_path = tmp.name
    
    # Создаем путь для результата
    output_path = tmp_path.replace(".docx", "_anonymized.docx")
    
    try:
        # ЭТАП 1: Загрузка документа
        doc = Document(tmp_path)
        
        # ЭТАП 2: Извлечение блоков
        builder = BlockBuilder()
        blocks = builder.build_blocks(doc)
        
        # ЭТАП 3: Поиск чувствительных данных
        rule_engine = RuleEngineAdapter(patterns_file)
        processed_blocks = rule_engine.apply_rules_to_blocks(blocks)
        
        # Собираем совпадения для замены
        all_matches = []
        for block in processed_blocks:
            if 'sensitive_patterns' in block:
                for pattern in block['sensitive_patterns']:
                    match = {
                        'block_id': block['block_id'],
                        'original_value': pattern['original_value'],
                        'uuid': pattern['uuid'],
                        'position': pattern['position'],
                        'element': block.get('element'),
                        'category': pattern['category']
                    }
                    all_matches.append(match)
        
        # ЭТАП 4: Применение замен с сохранением форматирования
        formatter = FormatterApplier(highlight_replacements=True)
        stats = formatter.apply_replacements_to_document(doc, all_matches)
        
        # ЭТАП 5: Сохранение результата
        doc.save(output_path)
        
        # Генерируем отчет
        report = rule_engine.generate_report(processed_blocks)
        
        response_data = {
            "message": "Документ успешно анонимизирован",
            "statistics": stats,
            "patterns_found": report['pattern_statistics'],
            "total_blocks": len(blocks),
            "matches_count": len(all_matches),
            "download_url": f"/download_anonymized/{os.path.basename(output_path)}"
        }
        
        return JSONResponse(content=response_data)
        
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"error": f"Ошибка анонимизации: {str(e)}"}
        )
    finally:
        # Очищаем временный входной файл
        if os.path.exists(tmp_path):
            os.remove(tmp_path)

@app.get("/download_anonymized/{filename}")
async def download_anonymized(filename: str):
    """Скачивание анонимизированного документа"""
    file_path = os.path.join(tempfile.gettempdir(), filename)
    
    if os.path.exists(file_path):
        return FileResponse(
            path=file_path,
            filename=filename,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    else:
        return JSONResponse(
            status_code=404,
            content={"error": "Файл не найден"}
        )

@app.post("/anonymize_full")
async def anonymize_full(file: UploadFile = File(...), 
                        patterns_file: str = "patterns/sensitive_patterns.xlsx",
                        generate_excel_report: bool = True,
                        generate_json_ledger: bool = True):
    """
    ПОЛНЫЙ ЦИКЛ АНОНИМИЗАЦИИ с Excel отчетом и JSON журналом
    Этап 2.4: Replacement Ledger и базовый Anonymize
    """
    # Создаем временные пути
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_input:
        shutil.copyfileobj(file.file, tmp_input)
        input_path = tmp_input.name
    
    # Создаем пути для выходных файлов  
    temp_dir = tempfile.mkdtemp()
    filename_base = os.path.splitext(file.filename)[0]
    
    output_path = os.path.join(temp_dir, f"{filename_base}_anonymized.docx")
    excel_path = os.path.join(temp_dir, f"{filename_base}_report.xlsx") if generate_excel_report else None
    json_path = os.path.join(temp_dir, f"{filename_base}_ledger.json") if generate_json_ledger else None
    
    try:
        # Инициализируем полный анонимизатор
        anonymizer = FullAnonymizer(patterns_path=patterns_file)
        
        # Выполняем полный цикл анонимизации
        result = anonymizer.anonymize_document(
            input_path=input_path,
            output_path=output_path,
            excel_report_path=excel_path,
            json_ledger_path=json_path
        )
        
        if result['status'] == 'success':
            import base64
            
            # Подготавливаем файлы в base64 для прямого скачивания
            files_base64 = {}
            
            # Анонимизированный документ
            if os.path.exists(output_path):
                with open(output_path, 'rb') as f:
                    files_base64['anonymized_document_base64'] = base64.b64encode(f.read()).decode('utf-8')
            
            # Excel отчет
            if excel_path and os.path.exists(excel_path):
                with open(excel_path, 'rb') as f:
                    files_base64['excel_report_base64'] = base64.b64encode(f.read()).decode('utf-8')
            
            # JSON журнал (если нужен)
            if json_path and os.path.exists(json_path):
                with open(json_path, 'rb') as f:
                    files_base64['json_ledger_base64'] = base64.b64encode(f.read()).decode('utf-8')
            
            # Подготавливаем ссылки для скачивания (для обратной совместимости)
            download_links = {
                'anonymized_document': f"/download_file/{os.path.basename(output_path)}",
            }
            
            if excel_path and os.path.exists(excel_path):
                download_links['excel_report'] = f"/download_file/{os.path.basename(excel_path)}"
            
            if json_path and os.path.exists(json_path):
                download_links['json_ledger'] = f"/download_file/{os.path.basename(json_path)}"
            
            # Обновляем результат и файлами в base64, и ссылками
            result['files_base64'] = files_base64
            result['download_links'] = download_links
            result['temp_directory'] = temp_dir  # Для очистки позже
            
            return JSONResponse(content=result)
        else:
            return JSONResponse(
                status_code=500,
                content=result
            )
    
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={
                "status": "error",
                "error_message": f"Ошибка полного цикла анонимизации: {str(e)}"
            }
        )
    finally:
        # Очищаем входной временный файл
        if os.path.exists(input_path):
            os.remove(input_path)

@app.post("/anonymize_selected")
async def anonymize_selected(file: UploadFile = File(...), selected_items: str = Form(None)):
    """
    Селективная анонимизация документа на основе выбранных пользователем элементов
    """
    print(f"🔧 [ANONYMIZE] Запрос анонимизации файла: {file.filename}")
    print(f"🔧 [ANONYMIZE] Получено selected_items: {selected_items}")
    
    if not selected_items:
        print(f"❌ [ANONYMIZE] Нет selected_items")
        return JSONResponse(
            status_code=400,
            content={"status": "error", "message": "Не указаны элементы для анонимизации"}
        )
    
    # Сохраняем загруженный файл
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        shutil.copyfileobj(file.file, tmp)
        tmp_path = tmp.name
    
    # Создаем путь для результата
    output_path = tmp_path.replace(".docx", "_anonymized.docx")
    
    try:
        # Парсим выбранные элементы из JSON
        import json
        selected_items_list = json.loads(selected_items)
        
        print(f"🔧 [ANONYMIZE] Парсинг JSON успешен: {len(selected_items_list)} элементов")
        print(f"🔧 [ANONYMIZE] Первые 3 элемента: {selected_items_list[:3] if len(selected_items_list) > 3 else selected_items_list}")
        
        # Инициализируем полный анонимизатор
        anonymizer = FullAnonymizer()
        
        print(f"🔧 [ANONYMIZE] Вызываем anonymize_selected_items...")
        
        # Выполняем селективную анонимизацию
        result = anonymizer.anonymize_selected_items(
            input_path=tmp_path,
            output_path=output_path,
            selected_items=selected_items_list
        )
        
        print(f"🔧 [ANONYMIZE] Результат анонимизации: {result.get('status', 'NO_STATUS')}")
        
        if result['status'] == 'success':
            # Кодируем файл в base64 для прямого скачивания
            with open(output_path, 'rb') as f:
                file_base64 = base64.b64encode(f.read()).decode('utf-8')
            
            result['anonymized_document_base64'] = file_base64
            result['download_url'] = f"/download_file/{os.path.basename(output_path)}"
            result['filename'] = file.filename.replace('.docx', '_anonymized.docx')
            
            return JSONResponse(content=result)
        else:
            return JSONResponse(
                status_code=500,
                content=result
            )
            
    except json.JSONDecodeError as e:
        print(f"❌ [ANONYMIZE] Ошибка парсинга JSON: {str(e)}")
        print(f"❌ [ANONYMIZE] selected_items: {selected_items}")
        return JSONResponse(
            status_code=400,
            content={"status": "error", "message": "Некорректный формат данных selected_items"}
        )
    except Exception as e:
        print(f"❌ [ANONYMIZE] Ошибка селективной анонимизации: {str(e)}")
        print(f"❌ [ANONYMIZE] Тип ошибки: {type(e).__name__}")
        import traceback
        print(f"❌ [ANONYMIZE] Traceback: {traceback.format_exc()}")
        return JSONResponse(
            status_code=500,
            content={"status": "error", "message": f"Ошибка селективной анонимизации: {str(e)}"}
        )
    finally:
        # Очищаем временные файлы
        if os.path.exists(tmp_path):
            os.remove(tmp_path)
        if os.path.exists(output_path):
            # Файл удалится автоматически через некоторое время или при перезапуске
            pass

@app.get("/download_file/{filename}")
async def download_file(filename: str):
    """Универсальное скачивание файлов (DOCX, Excel, JSON)"""
    # Ищем файл в временных директориях
    temp_dir = tempfile.gettempdir()
    file_path = None
    
    # Поиск в основной temp директории
    candidate_path = os.path.join(temp_dir, filename)
    if os.path.exists(candidate_path):
        file_path = candidate_path
    else:
        # Поиск в поддиректориях temp
        for root, dirs, files in os.walk(temp_dir):
            if filename in files:
                file_path = os.path.join(root, filename)
                break
    
    if file_path and os.path.exists(file_path):
        # Определяем MIME тип по расширению
        if filename.endswith('.docx'):
            media_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif filename.endswith('.xlsx'):
            media_type = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        elif filename.endswith('.json'):
            media_type = 'application/json'
        else:
            media_type = 'application/octet-stream'
        
        return FileResponse(
            path=file_path,
            filename=filename,
            media_type=media_type
        )
    else:
        return JSONResponse(
            status_code=404,
            content={"error": f"Файл {filename} не найден"}
        )

# Запуск сервера при прямом выполнении
if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8003, reload=True)
