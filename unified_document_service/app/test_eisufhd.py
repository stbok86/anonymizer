#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Анализ алгоритма детекции аббревиатур ИС и тест ЕИСУФХД"""

import requests
from docx import Document

print("="*100)
print("АЛГОРИТМ ДЕТЕКЦИИ АББРЕВИАТУР ИНФОРМАЦИОННЫХ СИСТЕМ")
print("="*100)

print("""
АЛГОРИТМ состоит из 4 этапов (в порядке выполнения):

1. _search_complex_abbreviations() - СЛОЖНЫЕ АББРЕВИАТУРЫ БЕЗ ПРОБЕЛОВ
   - Паттерны из nlp_patterns.json с priority=1 и type='regex'
   - Примеры паттернов:
     * \\b(ЕИС)[А-ЯA-Z]{2,}\\b  -> ЕИС + минимум 2 заглавные буквы
     * \\b(АИС)[А-ЯA-Z]{2,}\\b  -> АИС + минимум 2 заглавные буквы
     * \\b(ГИС)[А-ЯA-Z]{2,}\\b  -> ГИС + минимум 2 заглавные буквы
   
   - Группы захвата:
     * group(1) = анонимная часть (ЕИС, АИС, ГИС)
     * group(2) = приватная часть (УФХД, ПК и т.д.)
     * group(0) = полное совпадение (ЕИСУФХД)
   
   - Фильтрация ложных срабатываний:
     * Проверка символов ДО и ПОСЛЕ совпадения
     * Если до/после стоит строчная буква -> пропустить (это часть слова)
   
   - Результат: список детекций с confidence=0.9, method='complex_abbreviation'

2. _simple_pattern_search() - ПРОСТЫЕ ПАТТЕРНЫ
   - Дополнительные regex паттерны
   - Не применяется к аббревиатурам (используется для других случаев)

3. _search_simple_abbreviations() - ПРОСТЫЕ АББРЕВИАТУРЫ
   - Фразовые паттерны из JSON (type='phrase')
   - Пример: "ГИС ЖКХ"
   - Использует PhraseMatcher от spaCy

4. _search_spaced_abbreviations_filtered() - АББРЕВИАТУРЫ С ПРОБЕЛАМИ
   - Паттерны с priority=2 и type='regex'
   - Примеры: "ЕИС УФХД", "ФГИС ПК"
   - Проверка пересечений с уже найденными детекциями
   - Только если НЕ конфликтует с результатами этапа 1-3

5. _remove_duplicates() - УДАЛЕНИЕ ДУБЛИКАТОВ
   - Пороговое значение threshold=0.7
   - Удаляет перекрывающиеся детекции

КЛЮЧЕВОЙ МОМЕНТ:
Для "ЕИСУФХД" используется этап 1 (_search_complex_abbreviations)
Паттерн: \\b(ЕИС)[А-ЯA-Z]{2,}\\b
- \\b = граница слова
- (ЕИС) = группа 1, анонимная часть
- [А-ЯA-Z]{2,} = группа 2, минимум 2 заглавные буквы кириллицы/латиницы
- \\b = граница слова

ПРОБЛЕМА может быть:
1. Граница слова \\b не распознается корректно
2. Символы до/после не соответствуют ожиданиям
3. Паттерн не компилируется правильно
4. "УФХД" содержит всего 4 буквы, паттерн требует минимум 2 - должно работать
""")

print("\n" + "="*100)
print("ТЕСТ ДЕТЕКЦИИ 'ЕИСУФХД' В РЕАЛЬНОМ ДОКУМЕНТЕ")
print("="*100)

doc_path = r"C:\Projects\Anonymizer\unified_document_service\test_docs\test_01_1_4_SD22.docx"
doc = Document(doc_path)

# Ищем ЕИСУФХД в документе
all_text = []
for para in doc.paragraphs:
    all_text.append(para.text)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            all_text.append(cell.text)

found_blocks = []
for i, text in enumerate(all_text):
    if "ЕИСУФХД" in text or "ЕИС УФХД" in text:
        found_blocks.append((i, text))

if found_blocks:
    print(f"\nНайдено {len(found_blocks)} блоков с 'ЕИСУФХД' или 'ЕИС УФХД':")
    
    for idx, (i, text) in enumerate(found_blocks[:3], 1):
        print(f"\n{idx}. Блок #{i}:")
        # Показываем контекст
        eisufhd_pos = text.find("ЕИСУФХД")
        eis_ufhd_pos = text.find("ЕИС УФХД")
        
        if eisufhd_pos >= 0:
            context_start = max(0, eisufhd_pos - 30)
            context_end = min(len(text), eisufhd_pos + 38)
            context = text[context_start:context_end]
            print(f"   Контекст (ЕИСУФХД): '...{context}...'")
            print(f"   Позиция: {eisufhd_pos}")
            
            # Проверяем символы до и после
            before_char = text[eisufhd_pos - 1] if eisufhd_pos > 0 else ''
            after_char = text[eisufhd_pos + 7] if eisufhd_pos + 7 < len(text) else ''
            print(f"   Символ ДО: '{before_char}' (код: {ord(before_char) if before_char else 'N/A'})")
            print(f"   Символ ПОСЛЕ: '{after_char}' (код: {ord(after_char) if after_char else 'N/A'})")
            print(f"   ДО - строчная?: {before_char.islower() if before_char else False}")
            print(f"   ПОСЛЕ - строчная?: {after_char.islower() if after_char else False}")
        
        if eis_ufhd_pos >= 0:
            context_start = max(0, eis_ufhd_pos - 30)
            context_end = min(len(text), eis_ufhd_pos + 38)
            context = text[context_start:context_end]
            print(f"   Контекст (ЕИС УФХД): '...{context}...'")
            print(f"   Позиция: {eis_ufhd_pos}")
        
        # Тест через NLP Service
        print(f"\n   Тестируем через NLP Service...")
        response = requests.post(
            "http://localhost:8006/analyze",
            json={"blocks": [{"block_id": f"test_{i}", "content": text}]},
            timeout=10
        )
        
        if response.status_code == 200:
            result = response.json()
            detections = result.get('detections', [])
            is_detections = [d for d in detections if d.get('category') == 'information_system']
            
            print(f"   Статус: {response.status_code}")
            print(f"   Детекций ИС: {len(is_detections)}")
            
            if is_detections:
                for det in is_detections:
                    print(f"     ✅ Найдено: '{det.get('original_value')}' (метод: {det.get('method')}, conf: {det.get('confidence'):.2f})")
            else:
                print(f"     ❌ ЕИСУФХД НЕ ДЕТЕКТИРОВАН!")

# Прямой тест паттерна
print("\n" + "="*100)
print("ПРЯМОЙ ТЕСТ АББРЕВИАТУР")
print("="*100)

test_cases = [
    "ЕИСУФХД",
    "ЕИС УФХД",
    "ЕИСУФХД ПК",
    "В системе ЕИСУФХД реализовано",
    "использование ЕИСУФХД для",
]

for text in test_cases:
    print(f"\nТест: '{text}'")
    response = requests.post(
        "http://localhost:8006/analyze",
        json={"blocks": [{"block_id": "test", "content": text}]},
        timeout=10
    )
    
    if response.status_code == 200:
        result = response.json()
        detections = result.get('detections', [])
        is_detections = [d for d in detections if d.get('category') == 'information_system']
        
        if is_detections:
            for det in is_detections:
                print(f"  ✅ '{det.get('original_value')}' (метод: {det.get('method')}, conf: {det.get('confidence'):.2f})")
        else:
            print(f"  ❌ НЕ НАЙДЕНО")
