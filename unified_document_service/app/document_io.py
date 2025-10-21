from docx import Document
from typing import Dict, List, Any

class DocxWalker:
    """
    Класс для обхода структуры DOCX-документа: paragraphs, tables, headers, footers.
    Возвращает ссылки на объекты для дальнейшей работы.
    """
    def __init__(self, docx_path: str):
        self.doc = Document(docx_path)

    def walk(self) -> Dict[str, List[Any]]:
        result = {
            "paragraphs": list(self.doc.paragraphs),
            "tables": list(self.doc.tables),
            "headers": [section.header for section in self.doc.sections],
            "footers": [section.footer for section in self.doc.sections],
        }
        return result
