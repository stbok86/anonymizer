"""
Модуль для применения замен в документах с сохранением форматирования
"""

import re
import uuid
from typing import List, Dict, Any, Optional, Tuple
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX

try:
    from uuid_mapper import UUIDMapper
except ImportError:
    import sys
    import os
    sys.path.append(os.path.dirname(__file__))
    from uuid_mapper import UUIDMapper

try:
    from docx_metadata_handler import DocxMetadataHandler
except ImportError:
    import sys
    import os
    sys.path.append(os.path.dirname(__file__))
    from docx_metadata_handler import DocxMetadataHandler


class FormatterApplier:
    def __init__(self, highlight_replacements: bool = True):
        """
        Инициализация применителя форматирования
        
        Args:
            highlight_replacements: Выделять ли замененный текст жёлтым цветом (по умолчанию True)
        """
        self.highlight_replacements = highlight_replacements
        self.replacement_color = WD_COLOR_INDEX.YELLOW  # Жёлтый цвет для выделения UUID
        
        # 🎯 ЦЕНТРАЛИЗОВАННЫЙ UUID MAPPER
        self.uuid_mapper = UUIDMapper(namespace="document-anonymization")
        
    def apply_replacements(self, doc, replacements_table: List[Dict]) -> int:
        """
        Метод-алиас для совместимости с предыдущими версиями
        
        Args:
            doc: Документ DOCX
            replacements_table: Таблица замен
            
        Returns:
            Количество применененных замен
        """
        result = self.apply_replacements_to_document(doc, replacements_table)
        return result.get('total_replacements', 0)
    
    def apply_replacements_to_document(self, doc, replacements: List[Dict]) -> Dict[str, Any]:
        """
        Применение замен к документу с сохранением форматирования
        
        Args:
            doc: Документ DOCX 
            replacements: Список замен с информацией о позициях
            
        Returns:
            Статистика применения замен
        """
        print(f"📝 [FORMATTER_APPLIER] Получено замен для обработки: {len(replacements)}")
        
        # 🎯 НОРМАЛИЗАЦИЯ ЗАМЕН: обеспечиваем консистентные UUID
        normalized_replacements = self._normalize_replacements_with_centralized_uuids(replacements)
        
        for i, match in enumerate(normalized_replacements[:5]):  # Показываем первые 5
            print(f"📝 [FORMATTER_APPLIER] Замена {i+1}: '{match.get('original_value', 'N/A')}' → '{match.get('uuid', 'N/A')}'")
        if len(normalized_replacements) > 5:
            print(f"📝 [FORMATTER_APPLIER] ... и еще {len(replacements) - 5} замен")
            
        if not replacements:
            return {
                'total_replacements': 0,
                'categories': {},
                'blocks_processed': 0
            }
        
        stats = {
            'total_replacements': 0,
            'categories': {},
            'blocks_processed': 0,
            'replacement_details': []
        }
        
        # Группируем замены по блокам для эффективной обработки
        replacements_by_block = {}
        for replacement in normalized_replacements:  # Используем нормализованные с UUID
            block_id = replacement.get('block_id')
            if block_id not in replacements_by_block:
                replacements_by_block[block_id] = []
            replacements_by_block[block_id].append(replacement)
        
        # Обрабатываем каждый блок
        for block_id, block_replacements in replacements_by_block.items():
            try:
                # Сортируем замены по позиции (в обратном порядке для корректной замены)
                block_replacements.sort(key=lambda x: x.get('position', {}).get('start', 0), reverse=True)
                
                block_stats = self._apply_replacements_to_block(block_replacements)
                
                # Агрегируем статистику
                stats['total_replacements'] += block_stats['replacements_made']
                stats['blocks_processed'] += 1
                
                # Подсчет по категориям
                for replacement in block_replacements:
                    category = replacement.get('category', 'unknown')
                    if category not in stats['categories']:
                        stats['categories'][category] = 0
                    stats['categories'][category] += 1
                
                stats['replacement_details'].extend(block_stats['details'])
                
            except Exception as e:
                print(f"Ошибка при обработке блока {block_id}: {str(e)}")
                continue
        
        # 🎯 ДОПОЛНИТЕЛЬНАЯ ОБРАБОТКА: Headers & Footers ПОСЛЕ основного анонимизирования
        print(f"📝 [FORMATTER_APPLIER] Начинаем дополнительную обработку заголовков и колонтитулов...")
        header_footer_stats = self._apply_replacements_to_headers_footers(doc, normalized_replacements)
        
        # Агрегируем статистику headers/footers
        stats['total_replacements'] += header_footer_stats['total_replacements']
        stats['headers_footers_processed'] = header_footer_stats['headers_footers_processed']
        
        # Объединяем статистику по категориям
        for category, count in header_footer_stats['categories'].items():
            if category not in stats['categories']:
                stats['categories'][category] = 0
            stats['categories'][category] += count
        
        stats['replacement_details'].extend(header_footer_stats['replacement_details'])
        
        print(f"📝 [FORMATTER_APPLIER] Дополнительная обработка завершена. Замен в headers/footers: {header_footer_stats['total_replacements']}")
        
        # 🎯 ДОПОЛНИТЕЛЬНАЯ ОБРАБОТКА: Headers & Footers ПОСЛЕ основного анонимизирования
        print(f"📝 [FORMATTER_APPLIER] Начинаем дополнительную обработку заголовков и колонтитулов...")
        header_footer_stats = self._apply_replacements_to_headers_footers(doc, normalized_replacements)
        
        # Агрегируем статистику headers/footers
        stats['total_replacements'] += header_footer_stats['total_replacements']
        stats['headers_footers_processed'] = header_footer_stats['headers_footers_processed']
        
        # Объединяем статистику по категориям
        for category, count in header_footer_stats['categories'].items():
            if category not in stats['categories']:
                stats['categories'][category] = 0
            stats['categories'][category] += count
        
        stats['replacement_details'].extend(header_footer_stats['replacement_details'])
        
        print(f"📝 [FORMATTER_APPLIER] Дополнительная обработка завершена. Замен в headers/footers: {header_footer_stats['total_replacements']}")
        
        # 🎯 ВАЖНО: Сохраняем нормализованные замены для использования в отчетах
        # Убираем element для возможности JSON сериализации
        serializable_normalized = []
        for r in normalized_replacements:
            r_copy = r.copy()
            r_copy.pop('element', None)  # Удаляем XML элемент
            serializable_normalized.append(r_copy)
        
        stats['normalized_replacements'] = serializable_normalized
        
        return stats
    
    def _apply_replacements_to_block(self, block_replacements: List[Dict]) -> Dict[str, Any]:
        """
        Применение замен к конкретному блоку
        
        Args:
            block_replacements: Список замен для данного блока
            
        Returns:
            Статистика по блоку
        """
        block_stats = {
            'replacements_made': 0,
            'details': []
        }
        
        for replacement in block_replacements:
            try:
                success = self._apply_single_replacement(replacement)
                if success:
                    block_stats['replacements_made'] += 1
                    block_stats['details'].append({
                        'uuid': replacement.get('uuid'),
                        'category': replacement.get('category'),
                        'original_value': replacement.get('original_value'),
                        'success': True
                    })
                    
            except Exception as e:
                print(f"Ошибка при применении замены {replacement.get('uuid', 'unknown')}: {str(e)}")
                block_stats['details'].append({
                    'uuid': replacement.get('uuid'),
                    'category': replacement.get('category'),
                    'original_value': replacement.get('original_value'),
                    'success': False,
                    'error': str(e)
                })
        
        return block_stats
    
    def _apply_single_replacement(self, replacement: Dict) -> bool:
        """
        Применение одной конкретной замены
        
        Args:
            replacement: Информация о замене
            
        Returns:
            True если замена применена успешно
        """
        try:
            element = replacement.get('element')
            original_value = replacement.get('original_value', '')
            position = replacement.get('position', {})
            
            print(f"\n🔧 [SINGLE_REPLACEMENT] Начинаем замену:")
            print(f"🔧 [SINGLE_REPLACEMENT] Оригинал: '{original_value}'")
            print(f"🔧 [SINGLE_REPLACEMENT] Element type: {type(element) if element else 'None'}")
            print(f"🔧 [SINGLE_REPLACEMENT] Position: {position}")
            print(f"🔧 [SINGLE_REPLACEMENT] Block ID: {replacement.get('block_id', 'N/A')}")
            print(f"🔧 [SINGLE_REPLACEMENT] Category: {replacement.get('category', 'N/A')}")
            
            if element is None:
                print(f"🔧 [SINGLE_REPLACEMENT] ❌ Element is None")
                return False
                
            if not original_value:
                print(f"🔧 [SINGLE_REPLACEMENT] ❌ original_value is empty")
                return False
                
            # Дополнительная проверка для None значений
            if original_value is None:
                print(f"🔧 [SINGLE_REPLACEMENT] ❌ original_value is None")
                return False
            
            # 🎯 ЦЕНТРАЛИЗОВАННАЯ генерация UUID (игнорируем placeholder)
            existing_uuid = replacement.get('uuid', '')
            if not existing_uuid or existing_uuid == 'placeholder':
                # Генерируем централизованный детерминистический UUID
                replacement_value = self._generate_replacement_value(
                    original_value, 
                    replacement.get('category', 'unknown')
                )
            else:
                # Используем уже готовый UUID (если он реальный)
                replacement_value = existing_uuid
            
            print(f"🔧 [SINGLE_REPLACEMENT] UUID замены: '{replacement_value}'")
            
            # Проверяем содержание элемента перед заменой
            if hasattr(element, 'text'):
                current_text = getattr(element, 'text', '') or ''
                print(f"🔧 [SINGLE_REPLACEMENT] Текущий text: '{current_text}'")
            
            if hasattr(element, 'rows'):
                print(f"🔧 [SINGLE_REPLACEMENT] Таблица с {len(element.rows)} строками")
            

            # --- SDT (lxml.etree._Element) ---
            try:
                import lxml.etree
            except ImportError:
                lxml = None
            if 'lxml' in str(type(element)) or (hasattr(element, 'tag') and hasattr(element, 'xpath')):
                # SDT-элемент: ищем w:t и заменяем текст
                try:
                    # Используем xpath без параметра namespaces для BaseOxmlElement
                    # Пространство имен указывается прямо в XPath строке
                    try:
                        # Пытаемся использовать xpath с namespaces (для lxml.etree._Element)
                        text_elements = element.xpath('.//w:t', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                    except TypeError:
                        # Если ошибка - используем xpath без namespaces (для BaseOxmlElement)
                        # BaseOxmlElement уже знает о пространствах имен из документа
                        text_elements = element.xpath('.//w:t')
                    
                    replaced = False
                    for text_element in text_elements:
                        current_text = text_element.text or ''
                        if original_value and original_value in current_text:
                            new_text = current_text.replace(original_value, replacement_value, 1)
                            text_element.text = new_text
                            print(f"🔧 [SINGLE_REPLACEMENT][SDT] ✅ Замена в SDT: '{current_text}' → '{new_text}'")
                            replaced = True
                            break  # Только первое вхождение
                    if replaced:
                        return True
                    else:
                        print(f"🔧 [SINGLE_REPLACEMENT][SDT] ❌ Значение '{original_value}' не найдено в SDT")
                        return False
                except Exception as e:
                    print(f"🔧 [SINGLE_REPLACEMENT][SDT] ❌ Ошибка при замене в SDT: {str(e)}")
                    return False

            # --- Таблица ---
            if hasattr(element, 'rows'):
                print(f"🔧 [SINGLE_REPLACEMENT] Обрабатываем таблицу")
                result = self._replace_in_table(element, original_value, replacement_value, position)
                print(f"🔧 [SINGLE_REPLACEMENT] Результат замены в таблице: {result}")
                return result
            # --- Параграф ---
            elif hasattr(element, 'text'):
                print(f"🔧 [SINGLE_REPLACEMENT] Обрабатываем параграф")
                result = self._replace_in_paragraph(element, original_value, replacement_value, position)
                print(f"🔧 [SINGLE_REPLACEMENT] Результат замены в параграфе: {result}")
                return result
            # --- Общий случай ---
            else:
                print(f"🔧 [SINGLE_REPLACEMENT] Общий случай замены")
                current_text = getattr(element, 'text', '')
                if current_text is None:
                    current_text = ''
                print(f"🔧 [SINGLE_REPLACEMENT] Текущий текст элемента: '{current_text}'")
                if original_value and original_value in current_text:
                    new_text = current_text.replace(original_value, replacement_value)
                    element.text = new_text
                    print(f"🔧 [SINGLE_REPLACEMENT] ✅ Общая замена: '{current_text}' → '{new_text}'")
                    return True
                else:
                    print(f"🔧 [SINGLE_REPLACEMENT] ❌ Значение '{original_value}' не найдено в тексте '{current_text}'")
                print(f"🔧 [SINGLE_REPLACEMENT] ❌ Замена не выполнена")
                return False
            
        except Exception as e:
            print(f"🔧 [SINGLE_REPLACEMENT] ❌ Ошибка при применении замены: {str(e)}")
            import traceback
            print(f"🔧 [SINGLE_REPLACEMENT] Traceback: {traceback.format_exc()}")
            return False
            
        except Exception as e:
            print(f"Ошибка при применении замены: {str(e)}")
            return False
    
    def _normalize_text(self, text: str) -> str:
        """
        Нормализует текст, заменяя различные типы пробелов на обычные пробелы
        и удаляя лишние пробелы
        """
        if not text:
            return ''
        
        # Заменяем неразрывные пробелы (160) и другие виды пробелов на обычный пробел (32)
        text = text.replace('\u00A0', ' ')  # неразрывный пробел
        text = text.replace('\u2009', ' ')  # тонкий пробел
        text = text.replace('\u2007', ' ')  # цифровой пробел
        text = text.replace('\u2008', ' ')  # пунктуационный пробел
        text = text.replace('\u202F', ' ')  # узкий неразрывный пробел
        text = text.replace('\u3000', ' ')  # идеографический пробел
        
        # Нормализуем пробелы
        return ' '.join(text.split())

    def _replace_in_paragraph(self, paragraph, original_value: str, replacement_value: str, position: Dict) -> bool:
        """
        Замена текста в параграфе с сохранением форматирования и учетом позиции
        
        Args:
            paragraph: Параграф документа
            original_value: Исходное значение для замены
            replacement_value: Замещающее значение
            position: Позиция замены для точного попадания
            
        Returns:
            True если замена применена
        """
        try:
            # print(f"🔧 [PARAGRAPH] Попытка замены: '{original_value}' → '{replacement_value}'")
            # print(f"🔧 [PARAGRAPH] Информация о позиции: {position}")
            # Получаем полный текст параграфа для проверки
            paragraph_text = getattr(paragraph, 'text', '') or ''
            # print(f"🔧 [PARAGRAPH] Полный текст параграфа: '{paragraph_text}'")
            # print(f"🔧 [PARAGRAPH] Количество runs: {len(paragraph.runs)}")
            # Нормализуем текст для поиска
            original_value_normalized = self._normalize_text(original_value)
            paragraph_text_normalized = self._normalize_text(paragraph_text)
            # print(f"🔧 [PARAGRAPH] Нормализованный искомый текст: '{original_value_normalized}'")
            # print(f"🔧 [PARAGRAPH] Нормализованный текст параграфа: '{paragraph_text_normalized}'")
            if not original_value_normalized or original_value_normalized not in paragraph_text_normalized:
                # print(f"🔧 [PARAGRAPH] ❌ Текст не найден после нормализации")
                return False
            # Получаем целевую позицию для проверки
            target_position = position.get('start') if position else None
            # print(f"🔧 [PARAGRAPH] Целевая позиция: {target_position}")
            # Если позиция указана, проверяем соответствие
            if target_position is not None:
                # Ищем позицию текста в параграфе
                text_position_in_paragraph = paragraph_text_normalized.find(original_value_normalized)
                if text_position_in_paragraph == -1:
                    # print(f"🔧 [PARAGRAPH] ❌ Текст не найден в параграфе")
                    return False
                # print(f"🔧 [PARAGRAPH] Позиция текста в параграфе: {text_position_in_paragraph}")
                # print(f"🔧 [PARAGRAPH] Целевая позиция в документе: {target_position}")
                # Для параграфов используем менее строгую проверку позиции
                # так как позиция может отличаться из-за разной структуры документа
                position_match = True  # Для параграфов пока принимаем любую позицию
                if not position_match:
                    # print(f"🔧 [PARAGRAPH] ❌ Позиция не совпадает, пропускаем")
                    return False
                else:
                    # print(f"🔧 [PARAGRAPH] ✅ Позиция подходит для замены")
                    pass
            replacement_made = False
            # Стратегия 1: Прямой поиск с нормализацией в runs
            for i, run in enumerate(paragraph.runs):
                run_text = run.text or ''
                run_text_normalized = self._normalize_text(run_text)
                # print(f"🔧 [PARAGRAPH] Run {i}: '{run_text}' (нормализован: '{run_text_normalized}')")
                # Пробуем прямое совпадение
                if original_value in run_text or original_value_normalized in run_text_normalized:
                    # print(f"🔧 [PARAGRAPH] ✅ Найден в run {i}, заменяем")
                    # Заменяем в исходном тексте run'а
                    old_run_text = run.text
                    if original_value in run_text:
                        run.text = run_text.replace(original_value, replacement_value, 1)  # Заменяем только первое вхождение
                    else:
                        # Если точное совпадение не найдено, но нормализованное есть
                        run.text = self._replace_with_normalization(run_text, original_value, replacement_value)
                    replacement_made = True
                    # Применяем выделение к UUID
                    if self.highlight_replacements:
                        try:
                            run.font.highlight_color = self.replacement_color
                            # print(f"🔧 [PARAGRAPH] ✅ Жёлтое выделение UUID применено к run {i}")
                        except Exception as e:
                            # print(f"🔧 [PARAGRAPH] ⚠️ Не удалось применить выделение: {e}")
                            pass
                    # print(f"🔧 [PARAGRAPH] ✅ Замена в run {i} завершена: '{old_run_text}' → '{run.text}'")
                    # print(f"🔧 [PARAGRAPH] 🎯 Замена выполнена, выходим из поиска")
                    break  # Выходим после первой успешной замены
            # Стратегия 2: Если не нашли в отдельных runs, ищем в пересечении runs с нормализацией
            if not replacement_made and len(paragraph.runs) > 1:
                # print(f"🔧 [PARAGRAPH] Стратегия 2: поиск в пересечении runs с нормализацией")
                # Собираем текст из всех runs для точного поиска
                full_text = ''.join(run.text for run in paragraph.runs)
                full_text_normalized = self._normalize_text(full_text)
                # print(f"🔧 [PARAGRAPH] Полный текст из runs: '{full_text}'")
                # print(f"🔧 [PARAGRAPH] Нормализованный полный текст: '{full_text_normalized}'")
                if original_value in full_text or original_value_normalized in full_text_normalized:
                    # print(f"🔧 [PARAGRAPH] ✅ Найден в пересечении runs")
                    # Найдем позицию в нормализованном тексте
                    search_text = original_value if original_value in full_text else original_value_normalized
                    search_target = full_text if original_value in full_text else full_text_normalized
                    start_pos = search_target.find(search_text)
                    end_pos = start_pos + len(search_text)
                    # print(f"🔧 [PARAGRAPH] Позиция в полном тексте: {start_pos}-{end_pos}")
                    # print(f"🔧 [PARAGRAPH] Search text: '{search_text}'")
                    # print(f"🔧 [PARAGRAPH] Search target: '{search_target}'")
                    # print(f"🔧 [PARAGRAPH] Using normalized? {search_target == full_text_normalized}")
                    # Если мы работаем с нормализованным текстом, нужно найти соответствие в исходном
                    if search_target == full_text_normalized:
                        # print(f"🔧 [PARAGRAPH] Вызываем _replace_in_normalized_runs")
                        replacement_made = self._replace_in_normalized_runs(paragraph, original_value, replacement_value)
                    else:
                        # print(f"🔧 [PARAGRAPH] Вызываем _replace_across_runs")
                        replacement_made = self._replace_across_runs(paragraph, original_value, replacement_value, start_pos, end_pos)
                else:
                    # print(f"🔧 [PARAGRAPH] ❌ Текст не найден даже в полном тексте runs")
                    pass
            if replacement_made:
                # print(f"🔧 [PARAGRAPH] ✅ Замена выполнена успешно")
                pass
            else:
                # print(f"🔧 [PARAGRAPH] ❌ Замена не выполнена")
                pass
            return replacement_made
        except Exception as e:
            # print(f"🔧 [PARAGRAPH] ❌ Ошибка при замене в параграфе: {str(e)}")
            return False
    
    def _replace_with_normalization(self, run_text: str, original_value: str, replacement_value: str) -> str:
        """
        Замена текста с учетом нормализации пробелов (только первое вхождение)
        """
        # Создаем карту позиций символов для нормализованного текста к исходному
        normalized = self._normalize_text(run_text)
        original_normalized = self._normalize_text(original_value)
        
        if original_normalized in normalized:
            # Простая замена, если текст точно совпадает после нормализации
            # Заменяем все виды пробелов в исходном тексте на обычные пробелы
            result = run_text
            result = result.replace('\u00A0', ' ')  # неразрывный пробел
            result = result.replace('\u2009', ' ')  # тонкий пробел
            result = result.replace('\u2007', ' ')  # цифровой пробел
            result = result.replace('\u2008', ' ')  # пунктуационный пробел
            result = result.replace('\u202F', ' ')  # узкий неразрывный пробел
            result = result.replace('\u3000', ' ')  # идеографический пробел
            
            # Теперь заменяем нормализованный текст (только первое вхождение)
            if original_normalized in self._normalize_text(result):
                return result.replace(original_value, replacement_value, 1)
        
        return run_text
    
    def _replace_in_normalized_runs(self, paragraph, original_value: str, replacement_value: str) -> bool:
        """
        Замена текста, который найден только после нормализации пробелов
        """
        print(f"🔧 [NORMALIZED_RUNS] Начинаем замену в нормализованном тексте")
        print(f"🔧 [NORMALIZED_RUNS] Искомый текст: '{original_value}'")
        
        # Собираем полный текст и заменяем в нем
        full_text = ''.join(run.text for run in paragraph.runs)
        
        # Нормализуем полный текст 
        full_text_normalized = self._normalize_text(full_text)
        original_normalized = self._normalize_text(original_value)
        
        print(f"🔧 [NORMALIZED_RUNS] Полный текст: '{full_text}'")
        print(f"🔧 [NORMALIZED_RUNS] Нормализованный полный: '{full_text_normalized}'")
        print(f"🔧 [NORMALIZED_RUNS] Нормализованный искомый: '{original_normalized}'")
        
        if original_normalized not in full_text_normalized:
            print(f"🔧 [NORMALIZED_RUNS] ❌ Текст не найден даже после нормализации")
            return False
        
        # Найдем позицию в нормализованном тексте
        normalized_start = full_text_normalized.find(original_normalized)
        normalized_end = normalized_start + len(original_normalized)
        
        print(f"🔧 [NORMALIZED_RUNS] Позиция в нормализованном тексте: {normalized_start}-{normalized_end}")
        
        # Теперь нужно найти соответствующую позицию в исходном тексте
        # Создадим карту соответствий между исходным и нормализованным текстом
        original_to_normalized = []
        normalized_pos = 0
        
        for original_pos, char in enumerate(full_text):
            if char in ['\u00A0', '\u2009', '\u2007', '\u2008', '\u202F', '\u3000']:
                # Неразрывные пробелы - заменяются на обычный пробел
                original_to_normalized.append(normalized_pos)
                normalized_pos += 1
            elif char.isspace():
                # Обычные пробелы остаются
                original_to_normalized.append(normalized_pos)
                normalized_pos += 1
            else:
                # Обычные символы остаются
                original_to_normalized.append(normalized_pos)
                normalized_pos += 1
        
        # Найдем границы в исходном тексте
        original_start = None
        original_end = None
        
        for i, norm_pos in enumerate(original_to_normalized):
            if norm_pos == normalized_start and original_start is None:
                original_start = i
            if norm_pos == normalized_end - 1:
                original_end = i + 1
                break
        
        if original_start is None or original_end is None:
            print(f"🔧 [NORMALIZED_RUNS] ❌ Не удалось найти границы в исходном тексте")
            return False
        
        print(f"🔧 [NORMALIZED_RUNS] Позиция в исходном тексте: {original_start}-{original_end}")
        
        # Теперь заменяем по runs используя позицию в исходном тексте
        return self._replace_across_runs(paragraph, original_value, replacement_value, original_start, original_end)
    
    def _replace_across_runs(self, paragraph, original_value: str, replacement_value: str, start_pos: int, end_pos: int) -> bool:
        """
        Замена текста, распределенного по нескольким runs
        """
        try:
            # Определяем какие runs затронуты
            current_pos = 0
            affected_runs = []
            for i, run in enumerate(paragraph.runs):
                run_start = current_pos
                run_end = current_pos + len(run.text)
                # Проверяем пересечение с искомым текстом
                if not (run_end <= start_pos or run_start >= end_pos):
                    affected_runs.append({
                        'index': i,
                        'run': run,
                        'run_start': run_start,
                        'run_end': run_end,
                        'text_start': max(0, start_pos - run_start),
                        'text_end': min(len(run.text), end_pos - run_start)
                    })
                current_pos = run_end
            # print(f"🔧 [PARAGRAPH] Затронутые runs: {[r['index'] for r in affected_runs]}")
            if affected_runs:
                replacement_made = False
                # Заменяем текст в затронутых runs
                for i, run_info in enumerate(affected_runs):
                    run = run_info['run']
                    text_start = run_info['text_start']
                    text_end = run_info['text_end']
                    if i == 0:
                        # Первый run - добавляем replacement_value
                        if text_start == 0 and text_end == len(run.text):
                            # Весь run заменяется
                            run.text = replacement_value
                        elif text_start == 0:
                            # Заменяем начало
                            run.text = replacement_value + run.text[text_end:]
                        elif text_end == len(run.text):
                            # Заменяем конец
                            run.text = run.text[:text_start] + replacement_value
                        else:
                            # Заменяем середину
                            run.text = run.text[:text_start] + replacement_value + run.text[text_end:]
                        replacement_made = True
                        # print(f"🔧 [PARAGRAPH] ✅ Run {run_info['index']} заменен")
                        # Применяем выделение к UUID
                        if self.highlight_replacements:
                            try:
                                run.font.highlight_color = self.replacement_color
                                # print(f"🔧 [PARAGRAPH] ✅ Жёлтое выделение UUID применено к run {run_info['index']}")
                            except Exception as e:
                                # print(f"🔧 [PARAGRAPH] ⚠️ Не удалось применить выделение: {e}")
                                pass
                    else:
                        # Остальные runs - убираем затронутый текст
                        if text_start == 0 and text_end == len(run.text):
                            # Весь run удаляется
                            run.text = ""
                        elif text_start == 0:
                            # Удаляем начало
                            run.text = run.text[text_end:]
                        elif text_end == len(run.text):
                            # Удаляем конец
                            run.text = run.text[:text_start]
                        else:
                            # Удаляем середину (редкий случай)
                            run.text = run.text[:text_start] + run.text[text_end:]
                        # print(f"🔧 [PARAGRAPH] ✅ Run {run_info['index']} обрезан")
                        # Применяем выделение, если в run есть текст
                        if self.highlight_replacements and run.text.strip():
                            try:
                                run.font.highlight_color = self.replacement_color
                                # print(f"🔧 [PARAGRAPH] ✅ Выделение применено к обрезанному run {run_info['index']}")
                            except Exception as e:
                                # print(f"🔧 [PARAGRAPH] ⚠️ Не удалось применить выделение: {e}")
                                pass
                return replacement_made
            return False
        except Exception as e:
            # print(f"🔧 [PARAGRAPH] ❌ Ошибка при замене через runs: {str(e)}")
            return False

    def _replace_in_table(self, table, original_value: str, replacement_value: str, position_info: dict = None) -> bool:
        """
        Замена текста в таблице
        
        Args:
            table: Таблица документа
            original_value: Исходное значение
            replacement_value: Замещающее значение
            position_info: Информация о позиции для точной замены (опционально)
            
        Returns:
            True если замена применена
        """
        try:
            print(f"🔧 [TABLE] Начало замены в таблице: '{original_value}' → '{replacement_value}'")
            print(f"🔧 [TABLE] Информация о позиции: {position_info}")
            replacement_made = False
            
            # КРИТИЧНО: Позиция уже относительная к блоку (таблице), не нужно пересчитывать!
            target_position = position_info.get('start') if position_info else None
            
            # Извлекаем весь текст таблицы для поиска позиции
            table_text = ""
            cell_positions = []  # Мапим позиции на ячейки
            
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = getattr(cell, 'text', '') or ''
                    cell_start = len(table_text)
                    table_text += cell_text
                    cell_end = len(table_text)
                    
                    cell_positions.append({
                        'row': row_idx,
                        'col': cell_idx,
                        'start': cell_start,
                        'end': cell_end,
                        'cell': cell,
                        'text': cell_text
                    })
                    
                    # Добавляем разделитель как в BlockBuilder
                    if cell_idx < len(row.cells) - 1:
                        table_text += " | "
                    
                # Новая строка после каждой строки таблицы
                table_text += "\n"
            
            print(f"🔧 [TABLE] Текст таблицы ({len(table_text)} символов): '{table_text[:100]}...'")
            print(f"🔧 [TABLE] Ищем '{original_value}' на позиции {target_position}")
            
            # Находим ячейку по позиции
            # КРИТИЧНО: Нормализуем текст для поиска (заменяем \xa0 на обычный пробел)
            normalized_original = self._normalize_text(original_value)
            
            target_cell = None
            for cell_info in cell_positions:
                normalized_cell_text = self._normalize_text(cell_info['text'])
                if normalized_original in normalized_cell_text:
                    # Если позиция не указана, берём первую найденную
                    if target_position is None:
                        target_cell = cell_info
                        print(f"🔧 [TABLE] ✅ Найден текст в ячейке [{cell_info['row']}][{cell_info['col']}] (без проверки позиции)")
                        break
                    else:
                        # Проверяем что позиция попадает в диапазон ячейки
                        if cell_info['start'] <= target_position < cell_info['end']:
                            target_cell = cell_info
                            print(f"🔧 [TABLE] ✅ Найден текст в ячейке [{cell_info['row']}][{cell_info['col']}] на позиции {target_position}")
                            break
            
            if not target_cell:
                print(f"🔧 [TABLE] ❌ Текст '{original_value}' не найден в таблице")
                return False
            
            # Заменяем в найденной ячейке
            cell = target_cell['cell']
            for para_idx, paragraph in enumerate(cell.paragraphs):
                paragraph_text = getattr(paragraph, 'text', '') or ''
                # Проверяем с нормализацией текста
                if original_value and (original_value in paragraph_text or 
                                      normalized_original in self._normalize_text(paragraph_text)):
                    print(f"🔧 [TABLE] Замена в параграфе {para_idx} ячейки [{target_cell['row']}][{target_cell['col']}]")
                    cell_replacement_made = self._replace_in_paragraph(
                        paragraph, original_value, replacement_value, {}
                    )
                    if cell_replacement_made:
                        replacement_made = True
                        print(f"🔧 [TABLE] ✅ Замена выполнена успешно!")
                        return True
                    else:
                        print(f"🔧 [TABLE] ❌ Замена не удалась")
            
            return replacement_made
            
        except Exception as e:
            print(f"🔧 [TABLE] ❌ Ошибка замены в таблице: {str(e)}")
            import traceback
            print(f"🔧 [TABLE] Traceback: {traceback.format_exc()}")
            return False
    
    def _generate_replacement_value(self, original_value: str, category: str, existing_uuid: str = None) -> str:
        """
        🎯 ЦЕНТРАЛИЗОВАННАЯ генерация замещающего значения
        
        Args:
            original_value: Исходное значение
            category: Категория найденных данных
            existing_uuid: Существующий UUID из анализа (для обратной совместимости)
            
        Returns:
            UUID для замещения
        """
        
        # ВСЕГДА используем централизованный детерминистический UUID
        # Специальная логика для contract_number и information_system ОТКЛЮЧЕНА
        return self.uuid_mapper.get_uuid_for_text(original_value, category)
    
    def _generate_contract_number_replacement(self, original_number: str) -> str:
        """Генерация замещающего значения для номеров контрактов с сохранением структуры"""
        
        # Получаем базовый UUID для консистентности
        base_uuid = self.uuid_mapper.get_uuid_for_text(original_number, 'contract_number')
        short_id = base_uuid.replace('-', '')[:8].upper()
        
        # Сохраняем структуру оригинального номера
        if '/' in original_number:
            parts = original_number.split('/')
            if len(parts) == 2:
                return f"{short_id[:2]}/{parts[1]}"
        elif '-' in original_number:
            parts = original_number.split('-')
            if len(parts) >= 2:
                return f"{short_id[:2]}-{'-'.join(parts[1:])}"
        
        # Для простых номеров - полная замена
        return short_id
    
    def _generate_information_system_replacement(self, original_value: str) -> str:
        """Генерация замещающего значения для информационных систем"""
        
        # Если есть анонимизированный текст с плейсхолдером
        if '[SYSTEM_ID]' in original_value:
            # Заменяем плейсхолдер на короткий UUID
            base_uuid = self.uuid_mapper.get_uuid_for_text(original_value, 'information_system')
            short_id = base_uuid.replace('-', '')[:8].upper()
            return original_value.replace('[SYSTEM_ID]', short_id)
        
        # Стандартная замена
        return self.uuid_mapper.get_uuid_for_text(original_value, 'information_system')
    
    def generate_replacement_report(self, replacements: List[Dict]) -> Dict[str, Any]:
        """
        Генерация отчета о произведенных заменах
        
        Args:
            replacements: Список замен
            
        Returns:
            Отчет о заменах
        """
        report = {
            'total_replacements': len(replacements),
            'categories': {},
            'confidence_stats': {
                'high': 0,  # > 0.8
                'medium': 0,  # 0.5 - 0.8  
                'low': 0    # < 0.5
            }
        }
        
        for replacement in replacements:
            # Подсчет по категориям
            category = replacement.get('category', 'unknown')
            if category not in report['categories']:
                report['categories'][category] = 0
            report['categories'][category] += 1
            
            # Статистика уверенности
            confidence = replacement.get('confidence', 1.0)
            if confidence > 0.8:
                report['confidence_stats']['high'] += 1
            elif confidence > 0.5:
                report['confidence_stats']['medium'] += 1
            else:
                report['confidence_stats']['low'] += 1
        
        return report
    
    def _normalize_replacements_with_centralized_uuids(self, replacements: List[Dict]) -> List[Dict]:
        """
        🎯 Нормализация замен с централизованной генерацией UUID
        
        Args:
            replacements: Исходный список замен
            
        Returns:
            Нормализованный список с консистентными UUID
        """
        normalized = []
        
        for replacement in replacements:
            original_value = replacement.get('original_value', '')
            category = replacement.get('category', 'data')
            
            if not original_value:
                normalized.append(replacement)
                continue
                
            # Создаем копию замены
            normalized_replacement = replacement.copy()
            
            # 🎯 ОБРАБАТЫВАЕМ АНОНИМИЗИРОВАННЫЙ ТЕКСТ (если есть)
            anonymized_text = replacement.get('anonymized_text')
            if anonymized_text:
                # Если есть готовый анонимизированный текст, используем его как основу
                centralized_uuid = self._generate_replacement_value(anonymized_text, category, None)
                normalized_replacement['uuid'] = centralized_uuid
            else:
                # 🎯 ГЕНЕРИРУЕМ ЦЕНТРАЛИЗОВАННЫЙ UUID
                centralized_uuid = self.uuid_mapper.get_uuid_for_text(original_value, category)
                normalized_replacement['uuid'] = centralized_uuid
            
            normalized.append(normalized_replacement)
        
        return normalized
    
    def apply_complete_anonymization(self, docx_path: str, output_path: str, replacements: List[Dict]) -> Dict[str, Any]:
        """
        🎯 ПОЛНАЯ АНОНИМИЗАЦИЯ: Комплексная анонимизация с обработкой метаданных
        
        Обрабатывает:
        1. Основной текст документа (параграфы, таблицы)
        2. Заголовки и колонтитулы (статический текст + SDT)
        3. Метаданные документа (docProps/core.xml, app.xml, custom.xml)
        
        Args:
            docx_path: Путь к исходному DOCX файлу
            output_path: Путь для сохранения анонимизированного файла
            replacements: Нормализованный список замен
            
        Returns:
            Полная статистика анонимизации
        """
        print(f"🎆 [COMPLETE_ANONYMIZATION] Начало полной анонимизации")
        print(f"📄 Input: {docx_path}")
        print(f"📄 Output: {output_path}")
        print(f"🎯 Замен для обработки: {len(replacements)}")
        
        total_stats = {
            'total_replacements': 0,
            'categories': {},
            'blocks_processed': 0,
            'headers_footers_processed': 0,
            'metadata_replacements': 0,
            'replacement_details': [],
            'phases': {
                'document_content': {},
                'headers_footers': {},
                'metadata': {}
            }
        }
        
        try:
            from docx import Document
            import tempfile
            import os
            
            # Этап 1: Обработка основного содержимого документа
            print(f"🎆 [PHASE 1] Обработка основного содержимого документа")
            
            # Загружаем документ
            doc = Document(docx_path)
            
            # Нормализуем замены
            normalized_replacements = self._normalize_replacements_with_centralized_uuids(replacements)
            
            # Применяем стандартную анонимизацию к документу
            content_stats = self.apply_replacements_to_document(doc, normalized_replacements)
            total_stats['phases']['document_content'] = content_stats
            
            # Сохраняем промежуточный результат
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                doc.save(temp_file.name)
                intermediate_docx = temp_file.name
            
            print(f"🎆 [PHASE 1] ✅ Завершено. Замен: {content_stats['total_replacements']}")
            
            # Этап 2: Обработка метаданных
            print(f"🎆 [PHASE 2] Обработка метаданных документа")
            
            metadata_handler = DocxMetadataHandler(intermediate_docx)
            
            # Извлекаем метаданные
            metadata = metadata_handler.extract_metadata()
            
            # Ищем чувствительные данные в метаданных
            sensitive_metadata = metadata_handler.find_sensitive_metadata(normalized_replacements)
            
            # Анонимизируем метаданные и сохраняем финальный результат
            metadata_success = metadata_handler.anonymize_metadata_in_docx(
                intermediate_docx, output_path, sensitive_metadata
            )
            
            if metadata_success:
                total_stats['metadata_replacements'] = len(sensitive_metadata)
                total_stats['phases']['metadata'] = {
                    'sensitive_found': len(sensitive_metadata),
                    'success': True
                }
                print(f"🎆 [PHASE 2] ✅ Завершено. Замен в метаданных: {len(sensitive_metadata)}")
            else:
                print(f"🎆 [PHASE 2] ⚠️ Ошибка при обработке метаданных")
                # В случае ошибки копируем промежуточный результат
                import shutil
                shutil.copy2(intermediate_docx, output_path)
                total_stats['phases']['metadata'] = {
                    'sensitive_found': len(sensitive_metadata),
                    'success': False,
                    'error': 'Ошибка при анонимизации метаданных'
                }
            
            # Очищаем временный файл
            if os.path.exists(intermediate_docx):
                os.remove(intermediate_docx)
            
            # Агрегируем статистику
            total_stats['total_replacements'] = (
                content_stats['total_replacements'] + 
                total_stats['metadata_replacements']
            )
            total_stats['categories'] = content_stats['categories']
            total_stats['blocks_processed'] = content_stats['blocks_processed']
            total_stats['headers_footers_processed'] = content_stats.get('headers_footers_processed', 0)
            total_stats['replacement_details'] = content_stats['replacement_details']
            
            # Добавляем метаданные в детали
            for metadata_item in sensitive_metadata:
                total_stats['replacement_details'].append({
                    'uuid': metadata_item.get('uuid'),
                    'category': metadata_item.get('category'),
                    'original_value': metadata_item.get('original_value'),
                    'source': 'metadata',
                    'metadata_section': metadata_item.get('metadata_section'),
                    'metadata_property': metadata_item.get('metadata_property'),
                    'success': True
                })
            
            print(f"🎆 [COMPLETE_ANONYMIZATION] ✅ ПОЛНАЯ АНОНИМИЗАЦИЯ ЗАВЕРШЕНА")
            print(f"📈 Общий итог:")
            print(f"  🔢 Всего замен: {total_stats['total_replacements']}")
            print(f"  📄 Замен в документе: {content_stats['total_replacements']}")
            print(f"  📄 Замен в метаданных: {total_stats['metadata_replacements']}")
            print(f"  📝 Обработано блоков: {total_stats['blocks_processed']}")
            print(f"  📝 Headers/Footers: {total_stats['headers_footers_processed']}")
            
            return total_stats
            
        except Exception as e:
            print(f"🎆 [COMPLETE_ANONYMIZATION] ❌ Ошибка при полной анонимизации: {str(e)}")
            import traceback
            print(f"🎆 [COMPLETE_ANONYMIZATION] Traceback: {traceback.format_exc()}")
            
            # Возвращаем пустую статистику
            return {
                'total_replacements': 0,
                'categories': {},
                'blocks_processed': 0,
                'headers_footers_processed': 0,
                'metadata_replacements': 0,
                'replacement_details': [],
                'error': str(e)
            }
    
    def _apply_replacements_to_headers_footers(self, doc, replacements: List[Dict]) -> Dict[str, Any]:
        """
        🎯 ДОПОЛНИТЕЛЬНАЯ ОБРАБОТКА: Применение замен к заголовкам и колонтитулам документа
        Вызывается ПОСЛЕ основной обработки для обеспечения полного покрытия
        
        Args:
            doc: Документ DOCX
            replacements: Нормализованный список замен с консистентными UUID
            
        Returns:
            Статистика замен в headers/footers
        """
        stats = {
            'total_replacements': 0,
            'categories': {},
            'headers_footers_processed': 0,
            'replacement_details': []
        }
        
        try:
            # Фильтруем замены для headers/footers (по типу блока)
            header_footer_replacements = []
            for replacement in replacements:
                block_id = replacement.get('block_id', '')
                if any(block_type in block_id for block_type in ['header_', 'footer_']):
                    header_footer_replacements.append(replacement)
            
            print(f"🔧 [HEADERS_FOOTERS] Найдено замен для headers/footers: {len(header_footer_replacements)}")
            
            if not header_footer_replacements:
                print(f"🔧 [HEADERS_FOOTERS] ⚠️ Нет замен для headers/footers")
                return stats
            
            # Обрабатываем каждую секцию документа
            for section_idx, section in enumerate(doc.sections):
                print(f"🔧 [HEADERS_FOOTERS] Обрабатываем секцию {section_idx}")
                
                # Обрабатываем header секции
                if section.header:
                    header_stats = self._apply_replacements_to_header_footer(
                        section.header, header_footer_replacements, section_idx, 'header'
                    )
                    stats['total_replacements'] += header_stats['replacements_made']
                    stats['replacement_details'].extend(header_stats['details'])
                    
                    # Подсчет по категориям
                    for replacement in header_footer_replacements:
                        if replacement.get('block_id', '').startswith(f'header_{section_idx}'):
                            category = replacement.get('category', 'unknown')
                            if category not in stats['categories']:
                                stats['categories'][category] = 0
                            stats['categories'][category] += 1
                    
                    if header_stats['replacements_made'] > 0:
                        stats['headers_footers_processed'] += 1
                
                # Обрабатываем footer секции
                if section.footer:
                    footer_stats = self._apply_replacements_to_header_footer(
                        section.footer, header_footer_replacements, section_idx, 'footer'
                    )
                    stats['total_replacements'] += footer_stats['replacements_made']
                    stats['replacement_details'].extend(footer_stats['details'])
                    
                    # Подсчет по категориям
                    for replacement in header_footer_replacements:
                        if replacement.get('block_id', '').startswith(f'footer_{section_idx}'):
                            category = replacement.get('category', 'unknown')
                            if category not in stats['categories']:
                                stats['categories'][category] = 0
                            stats['categories'][category] += 1
                    
                    if footer_stats['replacements_made'] > 0:
                        stats['headers_footers_processed'] += 1
            
            print(f"🔧 [HEADERS_FOOTERS] ✅ Обработка завершена. Всего замен: {stats['total_replacements']}")
            return stats
            
        except Exception as e:
            print(f"🔧 [HEADERS_FOOTERS] ❌ Ошибка при обработке headers/footers: {str(e)}")
            import traceback
            print(f"🔧 [HEADERS_FOOTERS] Traceback: {traceback.format_exc()}")
            return stats
    
    def _apply_replacements_to_header_footer(self, container, replacements: List[Dict], 
                                           section_idx: int, container_type: str) -> Dict[str, Any]:
        """
        Применение замен к конкретному header или footer
        
        Args:
            container: Header или Footer объект
            replacements: Список замен
            section_idx: Индекс секции
            container_type: 'header' или 'footer'
            
        Returns:
            Статистика замен для данного контейнера
        """
        container_stats = {
            'replacements_made': 0,
            'details': []
        }
        
        try:
            print(f"🔧 [{container_type.upper()}] Обработка {container_type} секции {section_idx}")
            
            # Фильтруем замены для данного конкретного контейнера
            relevant_replacements = []
            for replacement in replacements:
                block_id = replacement.get('block_id', '')
                if block_id.startswith(f'{container_type}_{section_idx}'):
                    relevant_replacements.append(replacement)
            
            print(f"🔧 [{container_type.upper()}] Найдено релевантных замен: {len(relevant_replacements)}")
            
            if not relevant_replacements:
                return container_stats
            
            # Сортируем замены по позиции (в обратном порядке)
            relevant_replacements.sort(key=lambda x: x.get('position', {}).get('start', 0), reverse=True)
            
            # Применяем замены к параграфам в контейнере
            for paragraph in container.paragraphs:
                paragraph_text = getattr(paragraph, 'text', '') or ''
                print(f"🔧 [{container_type.upper()}] Параграф: '{paragraph_text[:100]}{'...' if len(paragraph_text) > 100 else ''}'")
                
                for replacement in relevant_replacements:
                    original_value = replacement.get('original_value', '')
                    
                    if original_value and original_value in paragraph_text:
                        print(f"🔧 [{container_type.upper()}] ✅ Найден текст для замены: '{original_value}'")
                        
                        # Применяем замену используя существующий метод
                        success = self._replace_in_paragraph(
                            paragraph, 
                            original_value, 
                            replacement.get('uuid', ''),
                            replacement.get('position', {})
                        )
                        
                        if success:
                            container_stats['replacements_made'] += 1
                            container_stats['details'].append({
                                'uuid': replacement.get('uuid'),
                                'category': replacement.get('category'),
                                'original_value': original_value,
                                'container_type': container_type,
                                'section_idx': section_idx,
                                'success': True
                            })
                            print(f"🔧 [{container_type.upper()}] ✅ Замена выполнена успешно")
                        else:
                            container_stats['details'].append({
                                'uuid': replacement.get('uuid'),
                                'category': replacement.get('category'),
                                'original_value': original_value,
                                'container_type': container_type,
                                'section_idx': section_idx,
                                'success': False,
                                'error': 'Не удалось выполнить замену в параграфе'
                            })
                            print(f"🔧 [{container_type.upper()}] ❌ Замена не выполнена")
            
            # 🎯 КРИТИЧНО: Обработка SDT элементов (Structured Document Tags)
            # Эти элементы часто содержат динамические данные в заголовках
            sdt_stats = self._apply_replacements_to_sdt_elements(container, relevant_replacements, container_type, section_idx)
            container_stats['replacements_made'] += sdt_stats['replacements_made']
            container_stats['details'].extend(sdt_stats['details'])
            
            print(f"🔧 [{container_type.upper()}] Итого замен в контейнере: {container_stats['replacements_made']}")
            return container_stats
            
        except Exception as e:
            print(f"🔧 [{container_type.upper()}] ❌ Ошибка при обработке контейнера: {str(e)}")
            return container_stats
    
    def _apply_replacements_to_sdt_elements(self, container, replacements: List[Dict], 
                                          container_type: str, section_idx: int) -> Dict[str, Any]:
        """
        🎯 КРИТИЧНО: Обработка SDT (Structured Document Tags) элементов в заголовках/колонтитулах
        SDT элементы часто содержат динамический контент, который нужно анонимизировать
        
        Args:
            container: Header или Footer объект
            replacements: Список замен для данного контейнера
            container_type: 'header' или 'footer'
            section_idx: Индекс секции
            
        Returns:
            Статистика замен в SDT элементах
        """
        sdt_stats = {
            'replacements_made': 0,
            'details': []
        }
        
        try:
            print(f"🔧 [SDT-{container_type.upper()}] Поиск SDT элементов в {container_type} секции {section_idx}")
            
            # Ищем SDT элементы в XML структуре контейнера
            if hasattr(container, '_element'):
                # Используем try/except для обработки разных типов элементов
                try:
                    # Пытаемся использовать xpath с namespaces (для lxml.etree._Element)
                    sdt_elements = container._element.xpath('.//w:sdt', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                except TypeError:
                    # Если ошибка - используем xpath без namespaces (для BaseOxmlElement)
                    sdt_elements = container._element.xpath('.//w:sdt')
                
                print(f"🔧 [SDT-{container_type.upper()}] Найдено SDT элементов: {len(sdt_elements)}")
                
                for sdt_idx, sdt_element in enumerate(sdt_elements):
                    # Извлекаем текст из SDT элемента
                    try:
                        # Пытаемся использовать xpath с namespaces
                        text_elements = sdt_element.xpath('.//w:t', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                    except TypeError:
                        # Если ошибка - используем xpath без namespaces
                        text_elements = sdt_element.xpath('.//w:t')
                    
                    for text_element in text_elements:
                        current_text = text_element.text or ''
                        print(f"🔧 [SDT-{container_type.upper()}] SDT текст: '{current_text}'")
                        
                        # Ищем соответствующие замены
                        for replacement in replacements:
                            original_value = replacement.get('original_value', '')
                            
                            if original_value and original_value in current_text:
                                print(f"🔧 [SDT-{container_type.upper()}] ✅ Найден текст для замены в SDT: '{original_value}'")
                                
                                # Выполняем замену в SDT элементе
                                new_text = current_text.replace(original_value, replacement.get('uuid', ''), 1)
                                text_element.text = new_text
                                
                                sdt_stats['replacements_made'] += 1
                                sdt_stats['details'].append({
                                    'uuid': replacement.get('uuid'),
                                    'category': replacement.get('category'),
                                    'original_value': original_value,
                                    'container_type': f'{container_type}_sdt',
                                    'section_idx': section_idx,
                                    'sdt_idx': sdt_idx,
                                    'success': True
                                })
                                
                                print(f"🔧 [SDT-{container_type.upper()}] ✅ Замена в SDT выполнена: '{current_text}' → '{new_text}'")
                                break  # Заменяем только первое вхождение
            
            print(f"🔧 [SDT-{container_type.upper()}] Замен в SDT элементах: {sdt_stats['replacements_made']}")
            return sdt_stats
            
        except Exception as e:
            print(f"🔧 [SDT-{container_type.upper()}] ❌ Ошибка при обработке SDT элементов: {str(e)}")
            return sdt_stats