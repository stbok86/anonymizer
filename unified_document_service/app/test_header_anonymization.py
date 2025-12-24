import os
import sys
sys.path.append(os.path.dirname(os.path.abspath(__file__)))
from full_anonymizer import FullAnonymizer

def test_header_anonymization():
    docx_path = r'C:\Projects\Anonymizer\unified_document_service\test_docs\222.docx'
    patterns_path = r'C:\Projects\Anonymizer\unified_document_service\patterns\sensitive_patterns.xlsx'
    output_path = r'C:\Projects\Anonymizer\unified_document_service\test_docs\222_anon.docx'
    excel_path = r'C:\Projects\Anonymizer\unified_document_service\test_docs\222_anon_report.xlsx'
    
    # 1. Инициализация анонимизатора
    anonymizer = FullAnonymizer(patterns_path=patterns_path)
    
    # 2. Загрузка документа один раз
    from docx import Document
    doc = Document(docx_path)
    # 3. Анализ документа (поиск чувствительных данных)
    blocks = anonymizer.block_builder.build_blocks(doc)
    print("[DEBUG] BLOCKS: Выводим все блоки, тип и текст:")
    for i, block in enumerate(blocks):
        print(f"  {i+1}. block_id={block.get('block_id')}, type={block.get('type')}, text='{block.get('text')}'")
        if 'sensitive_patterns' in block:
            print(f"    [DEBUG] sensitive_patterns: {block['sensitive_patterns']}")
    matches = []
    for block in blocks:
        if 'sensitive_patterns' in block:
            for pattern in block['sensitive_patterns']:
                print(f"[DEBUG] MATCH: block_id={block['block_id']}, original='{pattern['original_value']}', category={pattern['category']}")
                matches.append({
                    'original_value': pattern['original_value'],
                    'category': pattern['category'],
                    'block_id': block['block_id'],
                    'position': pattern.get('position', {}),
                })
    print(f"Найдено совпадений: {len(matches)}")

    # 4. Анонимизация документа (применяем к тому же объекту doc)
    anonymizer.formatter.apply_replacements_to_document(doc, matches)
    doc.save(output_path)
    print(f"Анонимизированный файл сохранён: {output_path}")
    
    # 5. Проверка результата (поиск исходного значения в header)
    doc_check = Document(output_path)
    found = False
    for section in doc_check.sections:
        header = section.header
        for paragraph in header.paragraphs:
            if 'Холдинг «ФЕРМСТРАТ»' in paragraph.text:
                found = True
    if found:
        print("❌ Исходное значение найдено в header после анонимизации!")
    else:
        print("✅ Исходное значение успешно заменено в header!")

if __name__ == "__main__":
    test_header_anonymization()
