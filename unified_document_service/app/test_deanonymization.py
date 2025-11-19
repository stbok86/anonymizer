#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Тест функционала деанонимизации
"""

import pandas as pd
from docx import Document
import tempfile
import os
import sys

# Добавляем путь к модулю
sys.path.append(os.path.dirname(__file__))

def create_test_files():
    """Создает тестовые файлы для проверки деанонимизации"""
    
    print("🔧 Создание тестовых файлов...")
    
    # Создаем тестовый анонимизированный документ
    doc = Document()
    
    # Добавляем заголовок
    doc.add_heading('Тестовый анонимизированный документ', 0)
    
    # Добавляем параграфы с UUID
    doc.add_paragraph('Электронная почта: a1b2c3d4-e5f6-7890-abcd-ef1234567890')
    doc.add_paragraph('Телефон: b2c3d4e5-f6g7-8901-bcde-f23456789012')
    doc.add_paragraph('ИНН: c3d4e5f6-g7h8-9012-cdef-345678901234')
    
    # Добавляем таблицу
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = 'Поле'
    table.cell(0, 1).text = 'Значение'
    table.cell(1, 0).text = 'Email'
    table.cell(1, 1).text = 'd4e5f6g7-h8i9-0123-defg-456789012345'
    table.cell(2, 0).text = 'Паспорт'
    table.cell(2, 1).text = 'e5f6g7h8-i9j0-1234-efgh-567890123456'
    
    # Сохраняем документ
    doc_path = tempfile.mktemp(suffix='_test_anonymized.docx')
    doc.save(doc_path)
    print(f"✅ Создан тестовый документ: {doc_path}")
    
    # Создаем таблицу соответствий
    replacement_data = {
        'uuid': [
            'a1b2c3d4-e5f6-7890-abcd-ef1234567890',
            'b2c3d4e5-f6g7-8901-bcde-f23456789012', 
            'c3d4e5f6-g7h8-9012-cdef-345678901234',
            'd4e5f6g7-h8i9-0123-defg-456789012345',
            'e5f6g7h8-i9j0-1234-efgh-567890123456'
        ],
        'original_value': [
            'admin@company.ru',
            '+7 (999) 123-45-67',
            '7701234567',
            'user@example.com',
            '1234 567890'
        ],
        'category': [
            'email',
            'phone',
            'inn',
            'email', 
            'passport'
        ],
        'confidence': [0.95, 0.98, 0.92, 0.97, 0.94]
    }
    
    df = pd.DataFrame(replacement_data)
    table_path = tempfile.mktemp(suffix='_test_replacements.xlsx')
    df.to_excel(table_path, index=False)
    print(f"✅ Создана таблица замен: {table_path}")
    print(f"📊 Записей в таблице: {len(df)}")
    
    return doc_path, table_path


def test_deanonymizer_module():
    """Тестирует модуль DocumentDeanonymizer"""
    
    try:
        from document_deanonymizer import DocumentDeanonymizer
        
        print("=" * 80)
        print("🧪 ТЕСТ МОДУЛЯ DOCUMENT DEANONYMIZER")
        print("=" * 80)
        
        # Создаем тестовые файлы
        doc_path, table_path = create_test_files()
        
        # Инициализируем деанонимизатор
        deanonymizer = DocumentDeanonymizer()
        
        # Выполняем деанонимизацию
        print("\n🔄 Запуск деанонимизации...")
        result = deanonymizer.deanonymize_document(doc_path, table_path)
        
        # Проверяем результат
        if result['success']:
            print("🎉 ДЕАНОНИМИЗАЦИЯ УСПЕШНА!")
            
            stats = result['statistics']
            print(f"\n📊 СТАТИСТИКА:")
            print(f"   • UUID найдено в документе: {stats['total_uuids_found']}")
            print(f"   • Всего замен: {stats['total_replacements']}")
            print(f"   • Успешных замен: {stats['successful_replacements']}")
            print(f"   • Неудачных замен: {stats['failed_replacements']}")
            
            if stats['total_replacements'] > 0:
                success_rate = (stats['successful_replacements'] / stats['total_replacements']) * 100
                print(f"   • Процент успеха: {success_rate:.1f}%")
            
            # Проверяем результирующий файл
            output_path = result['output_path']
            if os.path.exists(output_path):
                print(f"\n📄 Деанонимизированный документ создан: {output_path}")
                
                # Читаем результат для проверки
                result_doc = Document(output_path)
                print(f"\n📋 СОДЕРЖИМОЕ ДЕАНОНИМИЗИРОВАННОГО ДОКУМЕНТА:")
                
                for i, paragraph in enumerate(result_doc.paragraphs[:5], 1):
                    if paragraph.text.strip():
                        print(f"   {i}. {paragraph.text}")
                
                # Очищаем временный файл
                os.unlink(output_path)
                print(f"🧹 Очищен временный файл результата")
            
            # Проверяем отчет
            if 'report_path' in result and os.path.exists(result['report_path']):
                print(f"📊 Отчет создан: {result['report_path']}")
                os.unlink(result['report_path'])
                print(f"🧹 Очищен временный отчет")
            
        else:
            print("❌ ОШИБКА ДЕАНОНИМИЗАЦИИ!")
            print(f"Ошибка: {result.get('error', 'Неизвестная ошибка')}")
            return False
        
        # Очищаем тестовые файлы
        for file_path in [doc_path, table_path]:
            if os.path.exists(file_path):
                os.unlink(file_path)
        print(f"🧹 Очищены тестовые файлы")
        
        print(f"\n✅ ТЕСТ ЗАВЕРШЕН УСПЕШНО!")
        return True
        
    except ImportError as e:
        print(f"❌ Не удалось импортировать DocumentDeanonymizer: {e}")
        return False
    except Exception as e:
        print(f"❌ Ошибка при тестировании: {e}")
        return False


def test_api_endpoint():
    """Тестирует API endpoint через requests"""
    
    try:
        import requests
        
        print("\n" + "=" * 80)
        print("🌐 ТЕСТ API ENDPOINT")
        print("=" * 80)
        
        # Проверяем доступность сервисов
        gateway_url = "http://localhost:8002"
        unified_url = "http://localhost:8003"
        
        print(f"🔍 Проверка доступности сервисов...")
        
        # Gateway
        try:
            response = requests.get(f"{gateway_url}/health", timeout=5)
            print(f"✅ Gateway ({gateway_url}): {response.status_code}")
        except requests.exceptions.ConnectionError:
            print(f"❌ Gateway недоступен: {gateway_url}")
            print(f"💡 Запустите Gateway: cd gateway && python app/main.py")
            return False
        
        # Unified Service
        try:
            response = requests.get(f"{unified_url}/health", timeout=5)
            print(f"✅ Unified Service ({unified_url}): {response.status_code}")
        except requests.exceptions.ConnectionError:
            print(f"❌ Unified Service недоступен: {unified_url}")
            print(f"💡 Запустите Unified Service: cd unified_document_service && python app/main.py")
            return False
        
        # Создаем тестовые файлы
        doc_path, table_path = create_test_files()
        
        print(f"\n🚀 Отправка запроса на деанонимизацию...")
        
        # Отправляем запрос
        files = {
            'document': ('test_doc.docx', open(doc_path, 'rb'), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'),
            'replacement_table': ('test_table.xlsx', open(table_path, 'rb'), 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        }
        
        response = requests.post(
            f"{gateway_url}/deanonymize",
            files=files,
            timeout=60
        )
        
        # Закрываем файлы
        for file_tuple in files.values():
            file_tuple[1].close()
        
        if response.status_code == 200:
            result = response.json()
            print(f"🎉 API ТЕСТ УСПЕШЕН!")
            print(f"📊 Ответ: {result.get('message', 'Деанонимизация завершена')}")
            
            if 'statistics' in result:
                stats = result['statistics']
                print(f"📊 Статистика:")
                print(f"   • Успешных замен: {stats.get('successful_replacements', 0)}")
                print(f"   • Всего замен: {stats.get('total_replacements', 0)}")
        else:
            print(f"❌ API ОШИБКА: {response.status_code}")
            print(f"Ответ: {response.text}")
            return False
        
        # Очищаем файлы
        for file_path in [doc_path, table_path]:
            if os.path.exists(file_path):
                os.unlink(file_path)
        
        return True
        
    except ImportError:
        print("❌ Модуль requests не установлен")
        return False
    except Exception as e:
        print(f"❌ Ошибка API теста: {e}")
        return False


if __name__ == "__main__":
    print("🔓 ТЕСТИРОВАНИЕ ФУНКЦИОНАЛА ДЕАНОНИМИЗАЦИИ")
    print("=" * 80)
    
    # Тестируем модуль
    module_success = test_deanonymizer_module()
    
    # Тестируем API (если модуль работает)
    if module_success:
        api_success = test_api_endpoint()
        
        if module_success and api_success:
            print(f"\n🎉 ВСЕ ТЕСТЫ ПРОШЛИ УСПЕШНО!")
        else:
            print(f"\n⚠️ Есть проблемы с API, но модуль работает")
    else:
        print(f"\n❌ МОДУЛЬ НЕ РАБОТАЕТ")
    
    print(f"\n💡 Для полного тестирования запустите:")
    print(f"   1. Gateway: cd gateway && python app/main.py")
    print(f"   2. Unified Service: cd unified_document_service && python app/main.py") 
    print(f"   3. Frontend: cd frontend && streamlit run streamlit_app.py")