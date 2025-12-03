import requests
import json

# Отправляем запрос на анонимизацию через API
url = "http://localhost:8009/anonymize"

payload = {
    "file_path": "test_docs/test_01_1_4_SD33.docx",
    "output_path": "test_docs/test_api_output.docx",
    "excel_report_path": "test_docs/test_api_report.xlsx"
}

print("Отправляем запрос на анонимизацию через API...")
response = requests.post(url, json=payload, timeout=120)

print(f"\nСтатус: {response.status_code}")

if response.status_code == 200:
    result = response.json()
    print(f"\nУспех: {result.get('status')}")
    print(f"Всего замен: {result.get('statistics', {}).get('total_replacements', 0)}")
    
    # Проверяем, есть ли в логах что-то про МИНИСТЕРСТВО
    print("\n" + "="*80)
    print("Проверяем результат:")
    print("="*80)
    
    from docx import Document
    doc = Document('test_docs/test_api_output.docx')
    
    # Проверяем параграф 0
    para0 = doc.paragraphs[0].text if len(doc.paragraphs) > 0 else ""
    print(f"\nПараграф 0: '{para0}'")
    
    if "МИНИСТЕРСТВО" in para0.upper():
        print("❌ МИНИСТЕРСТВО НЕ ЗАМЕНЕНО!")
    else:
        print("✅ МИНИСТЕРСТВО заменено")
        
    # Проверяем Excel
    import openpyxl
    wb = openpyxl.load_workbook('test_docs/test_api_report.xlsx')
    ws = wb.active
    
    print("\n" + "="*80)
    print("Проверяем Excel таблицу:")
    print("="*80)
    
    ministry_found = False
    for row in ws.iter_rows(min_row=2, values_only=True):
        if row[0] and "министер" in str(row[0]).lower():
            print(f"Найдена запись: '{row[0]}'")
            ministry_found = True
    
    if not ministry_found:
        print("❌ НЕТ записи про министерство в таблице")
else:
    print(f"Ошибка: {response.text}")
