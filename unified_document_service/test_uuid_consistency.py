"""
Тест проверки консистентности UUID между документом и таблицей замен
"""
import os
import sys
import requests
import base64
import pandas as pd
from docx import Document
import re

# Настройка кодировки
sys.stdout.reconfigure(encoding='utf-8')
os.environ['PYTHONIOENCODING'] = 'utf-8'

API_URL = "http://localhost:8009"

def test_uuid_consistency():
    """Проверяет, что UUID в документе совпадают с UUID в таблице замен"""
    
    print("=" * 80)
    print("🧪 ТЕСТ КОНСИСТЕНТНОСТИ UUID")
    print("=" * 80)
    
    # Путь к тестовому документу
    test_doc_path = r'C:\Projects\Anonymizer\unified_document_service\test_docs\test_01_1_4_SD33.docx'
    
    if not os.path.exists(test_doc_path):
        print(f"❌ Тестовый документ не найден: {test_doc_path}")
        return False
    
    print(f"\n📄 Тестовый документ: {test_doc_path}")
    
    # Шаг 1: Отправляем документ на полную анонимизацию
    print("\n� Шаг 1: Отправка документа на анонимизацию...")
    
    with open(test_doc_path, 'rb') as f:
        files = {'file': (os.path.basename(test_doc_path), f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        data = {
            'patterns_file': 'patterns/sensitive_patterns.xlsx',
            'generate_excel_report': 'true',
            'generate_json_ledger': 'true'
        }
        
        response = requests.post(f"{API_URL}/anonymize_full", files=files, data=data, timeout=120)
    
    if response.status_code != 200:
        print(f"❌ Ошибка API: {response.status_code}")
        print(response.text)
        return False
    
    result = response.json()
    print(f"✅ Анонимизация выполнена")
    print(f"   Замен: {result.get('statistics', {}).get('total_replacements', 0)}")
    
    # Шаг 2: Извлекаем файлы
    print("\n🔄 Шаг 2: Извлечение файлов...")
    
    if 'files_base64' not in result:
        print("❌ Нет файлов в ответе")
        return False
    
    # Декодируем анонимизированный документ
    if 'anonymized_document_base64' not in result['files_base64']:
        print("❌ Нет анонимизированного документа")
        return False
    
    doc_data = base64.b64decode(result['files_base64']['anonymized_document_base64'])
    doc_path = 'test_anonymized.docx'
    with open(doc_path, 'wb') as f:
        f.write(doc_data)
    print(f"✅ Документ сохранен: {doc_path}")
    
    # Декодируем Excel таблицу замен
    if 'excel_report_base64' not in result['files_base64']:
        print("❌ Нет Excel таблицы замен")
        return False
    
    excel_data = base64.b64decode(result['files_base64']['excel_report_base64'])
    excel_path = 'test_replacements.xlsx'
    with open(excel_path, 'wb') as f:
        f.write(excel_data)
    print(f"✅ Excel таблица сохранена: {excel_path}")
    
    # Шаг 3: Извлекаем UUID из документа
    print("\n� Шаг 3: Анализ UUID в документе...")
    
    doc = Document(doc_path)
    doc_uuids = set()
    
    # UUID паттерн (стандартный формат)
    uuid_pattern = re.compile(r'\b[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}\b', re.IGNORECASE)
    
    for para in doc.paragraphs:
        found_uuids = uuid_pattern.findall(para.text)
        doc_uuids.update(found_uuids)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                found_uuids = uuid_pattern.findall(cell.text)
                doc_uuids.update(found_uuids)
    
    print(f"✅ Найдено {len(doc_uuids)} уникальных UUID в документе")
    
    # Шаг 4: Извлекаем UUID из Excel
    print("\n🔄 Шаг 4: Анализ UUID в таблице замен...")
    
    df = pd.read_excel(excel_path)
    print(f"📊 Колонки в Excel: {list(df.columns)}")
    
    # Находим колонку с UUID (может называться по-разному)
    uuid_column = None
    for col in df.columns:
        if 'замена' in col.lower() or 'uuid' in col.lower() or 'идентификатор' in col.lower():
            uuid_column = col
            break
    
    if uuid_column is None:
        print(f"❌ Не найдена колонка с UUID в Excel. Доступные колонки: {list(df.columns)}")
        return False
    
    excel_uuids = set()
    for val in df[uuid_column]:
        if pd.notna(val):
            val_str = str(val)
            found_uuids = uuid_pattern.findall(val_str)
            excel_uuids.update(found_uuids)
    
    print(f"✅ Найдено {len(excel_uuids)} уникальных UUID в таблице замен")
    
    # Шаг 5: Сравнение UUID
    print("\n� Шаг 5: Сравнение UUID...")
    
    # UUID только в документе
    only_in_doc = doc_uuids - excel_uuids
    # UUID только в Excel
    only_in_excel = excel_uuids - doc_uuids
    # UUID в обоих
    in_both = doc_uuids & excel_uuids
    
    print(f"\n📊 РЕЗУЛЬТАТЫ:")
    print(f"   UUID в обоих местах:      {len(in_both)}")
    print(f"   UUID только в документе:  {len(only_in_doc)}")
    print(f"   UUID только в таблице:    {len(only_in_excel)}")
    
    # Выводим примеры
    if in_both:
        print(f"\n✅ Примеры UUID в обоих местах (первые 3):")
        for uuid_val in list(in_both)[:3]:
            print(f"   - {uuid_val}")
    
    if only_in_doc:
        print(f"\n⚠️ UUID только в документе (первые 5):")
        for uuid_val in list(only_in_doc)[:5]:
            print(f"   - {uuid_val}")
    
    if only_in_excel:
        print(f"\n⚠️ UUID только в таблице замен (первые 5):")
        for uuid_val in list(only_in_excel)[:5]:
            print(f"   - {uuid_val}")
    
    # Проверка детерминизма - повторная анонимизация
    print("\n🔄 Шаг 6: Проверка детерминизма (повторная анонимизация)...")
    
    with open(test_doc_path, 'rb') as f:
        files = {'file': (os.path.basename(test_doc_path), f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        data = {
            'patterns_file': 'patterns/sensitive_patterns.xlsx',
            'generate_excel_report': 'true'
        }
        
        response2 = requests.post(f"{API_URL}/anonymize_full", files=files, data=data, timeout=120)
    
    if response2.status_code == 200:
        result2 = response2.json()
        excel_data2 = base64.b64decode(result2['files_base64']['excel_report_base64'])
        excel_path2 = 'test_replacements_2.xlsx'
        with open(excel_path2, 'wb') as f:
            f.write(excel_data2)
        
        df2 = pd.read_excel(excel_path2)
        excel_uuids_2 = set()
        for val in df2[uuid_column]:
            if pd.notna(val):
                val_str = str(val)
                found_uuids = uuid_pattern.findall(val_str)
                excel_uuids_2.update(found_uuids)
        
        if excel_uuids == excel_uuids_2:
            print("✅ UUID детерминистичны (одинаковые при повторной анонимизации)")
        else:
            print("❌ UUID НЕ детерминистичны!")
            print(f"   Разница: {len(excel_uuids ^ excel_uuids_2)} UUID отличаются")
    
    # Итоговая оценка
    print("\n" + "=" * 80)
    if len(only_in_doc) == 0 and len(only_in_excel) == 0:
        print("✅ ТЕСТ ПРОЙДЕН: UUID полностью консистентны")
        return True
    else:
        print("⚠️ ТЕСТ НЕ ПРОЙДЕН: Есть несовпадения UUID")
        print(f"   Несовпадений в документе: {len(only_in_doc)}")
        print(f"   Несовпадений в таблице:   {len(only_in_excel)}")
        return False


if __name__ == "__main__":
    success = test_uuid_consistency()
    sys.exit(0 if success else 1)
