#!/usr/bin/env python3
"""
АНАЛИЗ ИСХОДНОГО ДОКУМЕНТА
=========================

Анализируем исходный документ для понимания причин дублирования UUID
"""

import os
from docx import Document

def analyze_original_document():
    """Анализирует исходный документ"""
    
    original_doc_path = "unified_document_service/test_docs/test_01_1_4_S.docx"
    
    if not os.path.exists(original_doc_path):
        print(f"❌ Исходный файл не найден: {original_doc_path}")
        return
    
    print("🔍 АНАЛИЗ ИСХОДНОГО ДОКУМЕНТА")
    print("=" * 50)
    
    # Загружаем исходный документ
    doc = Document(original_doc_path)
    
    # Анализируем table_2 детально
    print(f"📋 АНАЛИЗ TABLE_2 В ИСХОДНОМ ДОКУМЕНТЕ:")
    print("-" * 40)
    
    if len(doc.tables) > 2:
        table_2 = doc.tables[2]
        print(f"Размер таблицы: {len(table_2.rows)} строк x {len(table_2.columns)} столбцов")
        
        # Ищем строки, которые содержат "14 августа 2023"
        target_date = "14 августа 2023"
        
        for row_idx, row in enumerate(table_2.rows):
            for cell_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                if target_date in text:
                    print(f"\n🎯 НАЙДЕН '{target_date}' в строке {row_idx}, ячейке {cell_idx}:")
                    print(f"   Полный текст: '{text}'")
                    
                    # Показываем контекст - соседние строки
                    print(f"   Контекст:")
                    for context_row in range(max(0, row_idx-1), min(len(table_2.rows), row_idx+2)):
                        if context_row != row_idx:
                            context_text = table_2.rows[context_row].cells[cell_idx].text.strip()
                            print(f"     Строка {context_row}: '{context_text[:100]}{'...' if len(context_text) > 100 else ''}'")
    else:
        print("❌ table_2 не найдена")
    
    # Также проанализируем весь документ на предмет "14 августа 2023"
    print(f"\n🔍 ПОИСК '{target_date}' ВО ВСЕМ ДОКУМЕНТЕ:")
    print("-" * 50)
    
    occurrences = 0
    
    # Параграфы
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text
        if target_date in text:
            occurrences += 1
            print(f"📄 Параграф {i}: '{text.strip()[:200]}{'...' if len(text.strip()) > 200 else ''}'")
    
    # Таблицы
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                text = cell.text
                if target_date in text:
                    occurrences += 1
                    print(f"📋 Таблица {table_idx}, строка {row_idx}, ячейка {cell_idx}: '{text.strip()[:200]}{'...' if len(text.strip()) > 200 else ''}'")
    
    print(f"\n📊 ИТОГО найдено '{target_date}': {occurrences} раз(а)")

if __name__ == "__main__":
    analyze_original_document()