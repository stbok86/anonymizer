#!/usr/bin/env python3
"""
Поиск документов с датой "14 августа 2023"
"""

import os
from docx import Document
import glob

def search_for_date_in_docs():
    """
    Ищем документы содержащие дату "14 августа 2023"
    """
    search_text = "14 августа 2023"
    base_dir = r"C:\Projects\Anonymizer"
    
    # Ищем все docx файлы рекурсивно
    docx_files = []
    for root, dirs, files in os.walk(base_dir):
        for file in files:
            if file.endswith('.docx'):
                full_path = os.path.join(root, file)
                docx_files.append(full_path)
    
    print(f"Найдено {len(docx_files)} файлов .docx")
    
    found_files = []
    
    for file_path in docx_files:
        try:
            print(f"\nПроверяем: {file_path}")
            doc = Document(file_path)
            
            # Проверяем параграфы
            found_in_paras = 0
            for para in doc.paragraphs:
                if search_text in para.text:
                    found_in_paras += 1
            
            # Проверяем таблицы
            found_in_tables = 0
            table_details = []
            for table_idx, table in enumerate(doc.tables):
                table_found = 0
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        if search_text in cell.text:
                            table_found += 1
                            table_details.append(f"table_{table_idx} ячейка [{row_idx}][{cell_idx}]")
                            found_in_tables += 1
            
            if found_in_paras > 0 or found_in_tables > 0:
                print(f"✅ НАЙДЕНО в {file_path}")
                print(f"   Параграфы: {found_in_paras}")
                print(f"   Таблицы: {found_in_tables}")
                if table_details:
                    print(f"   Детали таблиц: {table_details}")
                found_files.append({
                    'path': file_path,
                    'paras': found_in_paras,
                    'tables': found_in_tables,
                    'table_details': table_details
                })
                
        except Exception as e:
            print(f"❌ Ошибка при чтении {file_path}: {str(e)}")
    
    print(f"\n{'='*50}")
    print(f"ИТОГИ ПОИСКА:")
    print(f"Найдено файлов с датой '{search_text}': {len(found_files)}")
    
    for result in found_files:
        print(f"\n📄 {result['path']}")
        print(f"   📝 В параграфах: {result['paras']}")
        print(f"   🗂️ В таблицах: {result['tables']}")
        if result['table_details']:
            for detail in result['table_details']:
                print(f"      {detail}")
    
    return found_files

if __name__ == "__main__":
    search_for_date_in_docs()