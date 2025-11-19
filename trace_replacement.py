#!/usr/bin/env python3
"""
ТРАССИРОВКА КОНКРЕТНОЙ ЗАМЕНЫ
============================

Отслеживаем что происходит с элементом paragraph_82 
"Общество с ограниченной ответственностью «КАМА Технологии»"
при нажатии "Подтвердить анонимизацию"
"""

import os
import sys
import json
import tempfile
import shutil

unified_service_path = os.path.join(os.path.dirname(__file__), 'unified_document_service', 'app')
sys.path.append(unified_service_path)

from docx import Document
from full_anonymizer import FullAnonymizer
from block_builder import BlockBuilder

def trace_specific_replacement():
    """
    Трассировка конкретной замены paragraph_82
    """
    print("🎯 ТРАССИРОВКА ЗАМЕНЫ paragraph_82")
    print("=" * 80)
    
    try:
        from docx import Document
        from full_anonymizer import FullAnonymizer
        print("✅ Модули импортированы")
    except ImportError as e:
        print(f"❌ Ошибка импорта: {e}")
        return
    
    doc_path = "unified_document_service/test_docs/test_01_1_4_S.docx"
    
    if not os.path.exists(doc_path):
        print(f"❌ Файл не найден: {doc_path}")
        return
    
    # Создаем временную копию для тестирования
    temp_path = "temp_test_document.docx"
    shutil.copy2(doc_path, temp_path)
    output_path = "temp_test_output.docx"
    
    print(f"📄 Исходный файл: {doc_path}")
    print(f"🔄 Тестовый файл: {temp_path}")
    print(f"📤 Выходной файл: {output_path}")
    print()
    
    # Симулируем данные как они приходят из интерфейса
    selected_items = [{
        'block_id': 'paragraph_82',
        'original_value': 'Общество с ограниченной ответственностью «КАМА Технологии»',
        'uuid': '12345678-1234-1234-1234-123456789abc',
        'position': {'start': 0, 'end': 58},
        'category': 'organization',
        'confidence': 0.95,
        'approved': True
    }]
    
    print("📋 СИМУЛИРУЕМ ДАННЫЕ ИЗ ИНТЕРФЕЙСА:")
    print(f"   Block ID: {selected_items[0]['block_id']}")
    print(f"   Original: '{selected_items[0]['original_value']}'")
    print(f"   UUID: {selected_items[0]['uuid']}")
    print()
    
    # ЭТАП 1: Проверяем что элемент существует в документе
    print("🔍 ЭТАП 1: ПРОВЕРКА СУЩЕСТВОВАНИЯ ЭЛЕМЕНТА")
    print("-" * 45)
    
    doc = Document(temp_path)
    
    # Ищем paragraph_82 через прямой доступ
    target_para = None
    if len(doc.paragraphs) > 82:
        target_para = doc.paragraphs[82]
        print(f"✅ paragraph_82 найден")
        print(f"   Текст: '{target_para.text}'")
        print(f"   Длина: {len(target_para.text)} символов")
        
        target_text = selected_items[0]['original_value']
        if target_text in target_para.text:
            print("✅ Искомый текст НАЙДЕН в параграфе!")
        else:
            print("❌ Искомый текст НЕ найден в параграфе!")
            print(f"   Ищем: '{target_text}'")
            print(f"   В тексте: '{target_para.text}'")
    else:
        print(f"❌ paragraph_82 не найден (всего параграфов: {len(doc.paragraphs)})")
        return
    
    # ЭТАП 2: Выполняем замену через FullAnonymizer
    print("\\n⚡ ЭТАП 2: ВЫПОЛНЕНИЕ ЗАМЕНЫ ЧЕРЕЗ FullAnonymizer")
    print("-" * 50)
    
    anonymizer = FullAnonymizer()
    
    # Включаем детальное логирование
    print("🔧 Вызываем anonymize_selected_items()...")
    
    result = anonymizer.anonymize_selected_items(
        input_path=temp_path,
        output_path=output_path,
        selected_items=selected_items
    )
    
    print(f"📊 Результат anonymize_selected_items: {result.get('status', 'unknown')}")
    print(f"📊 Сообщение: {result.get('message', 'нет сообщения')}")
    print(f"📊 Замен применено: {result.get('replacements_applied', 'неизвестно')}")
    
    # ЭТАП 3: Проверяем результат
    print("\\n🔍 ЭТАП 3: ПРОВЕРКА РЕЗУЛЬТАТА ЗАМЕНЫ")
    print("-" * 35)
    
    if os.path.exists(output_path):
        print("✅ Выходной файл создан")
        
        # Открываем результирующий документ
        result_doc = Document(output_path)
        
        if len(result_doc.paragraphs) > 82:
            result_para = result_doc.paragraphs[82]
            result_text = result_para.text
            
            print(f"📄 Текст после замены: '{result_text}'")
            
            uuid_text = selected_items[0]['uuid']
            original_text = selected_items[0]['original_value']
            
            if uuid_text in result_text:
                print("🎉 ЗАМЕНА ПРИМЕНЕНА УСПЕШНО!")
            elif original_text in result_text:
                print("❌ ЗАМЕНА НЕ ПРИМЕНЕНА - остался исходный текст!")
            else:
                print("⚠️  Неопределенный результат - текст изменился, но не на UUID")
            
            # Сравнение символ в символ
            print("\\n🔬 ДЕТАЛЬНОЕ СРАВНЕНИЕ:")
            print(f"   ДО:     '{target_para.text}'")
            print(f"   ПОСЛЕ:  '{result_text}'")
            print(f"   Равны:  {target_para.text == result_text}")
            
        else:
            print("❌ paragraph_82 не найден в результирующем документе")
    else:
        print("❌ Выходной файл НЕ создан!")
    
    # ЭТАП 4: Анализ логов FullAnonymizer (если есть)
    print("\\n📋 ЭТАП 4: АНАЛИЗ ПРОЦЕССА ЗАМЕНЫ")
    print("-" * 35)
    
    # Повторяем замену с более детальным анализом
    print("🔧 Повторная замена с детальным анализом...")
    
    # Заново загружаем исходный документ
    doc2 = Document(temp_path)
    
    # Получаем блоки
    blocks = anonymizer.block_builder.build_blocks(doc2)
    
    # Ищем наш блок
    target_block = None
    for block in blocks:
        if block['block_id'] == 'paragraph_82':
            target_block = block
            break
    
    if target_block:
        print(f"✅ Блок paragraph_82 найден в BlockBuilder")
        print(f"   Содержимое: '{target_block.get('content', '')}'")
        print(f"   Тип: {target_block.get('block_type', 'unknown')}")
        print(f"   Element: {type(target_block.get('element'))}")
        
        # Проверяем element
        element = target_block.get('element')
        if element and hasattr(element, 'text'):
            print(f"   Element.text: '{element.text}'")
            
            # Проверяем совпадение
            target_text = selected_items[0]['original_value']
            if target_text in element.text:
                print("✅ Текст найден в element!")
            else:
                print("❌ Текст НЕ найден в element!")
        else:
            print("❌ Element не имеет атрибута text!")
    else:
        print("❌ Блок paragraph_82 НЕ найден в BlockBuilder!")
    
    # Очистка
    try:
        os.remove(temp_path)
        os.remove(output_path)
    except:
        pass
    
    print("\\n" + "=" * 80)
    print("🎯 ТРАССИРОВКА ЗАВЕРШЕНА")

if __name__ == "__main__":
    trace_specific_replacement()